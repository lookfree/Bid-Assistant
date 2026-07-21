from __future__ import annotations

import asyncio
import io
import uuid
from datetime import date
from typing import Literal

from docx import Document
from fastapi import APIRouter
from pydantic import BaseModel, Field

from agent.agents.bidding_agent.render.pdf import docx_to_pdf
from agent.parsing.storage_read import storage

# spec315b 契约 3：POST /render/checklist 同步无状态渲染——不进 thread、不涉计费，
# App 把模板+状态合成后灌进来，agent 出 .docx 落 MinIO 返回 key（预签名在 App API 做）。
# /render/risk-report 同范式：废标体检报告导出（体检 review 步已计费，导报告免费）。

router = APIRouter()

_STATUS_LABEL = {"pass": "✓ 通过", "risk": "⚠ 风险", "pending": "待办"}
_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PDF_CT = "application/pdf"


class ChecklistItem(BaseModel):
    text: str
    status: Literal["pass", "risk", "pending"] = "pending"
    owner: str = ""
    note: str = ""
    # 知识库命中：可空展示字符串（如「已具备 · 营业执照」），直接作为一列输出，null 显示 —
    library_hit: str | None = None


class ChecklistGroup(BaseModel):
    id: str
    title: str
    items: list[ChecklistItem]


class ChecklistBody(BaseModel):
    title: str
    project_name: str | None = None
    groups: list[ChecklistGroup]


def render_checklist_docx(body: ChecklistBody) -> bytes:
    """终极审核表 .docx：标题+项目名+日期，每组一张五列表格
    （检查项/状态/责任人/备注/知识库），末尾签字/日期栏。
    与 render/docx.py 同风格：python-docx 直出，确定性无 LLM。"""
    doc = Document()
    doc.add_heading(body.title, level=0)
    if body.project_name:
        doc.add_paragraph(f"项目名称：{body.project_name}")
    doc.add_paragraph(f"导出日期：{date.today().isoformat()}")
    for g in body.groups:
        doc.add_heading(g.title, level=1)
        t = doc.add_table(rows=len(g.items) + 1, cols=5)
        t.style = "Table Grid"
        for j, h in enumerate(("检查项", "状态", "责任人", "备注", "知识库")):
            t.rows[0].cells[j].text = h
        for i, item in enumerate(g.items, start=1):
            row = (item.text, _STATUS_LABEL[item.status], item.owner, item.note,
                   item.library_hit or "—")
            for j, v in enumerate(row):
                t.rows[i].cells[j].text = v
    doc.add_paragraph("")
    doc.add_paragraph("审核人（签字）：____________________    日期：____________")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


@router.post("/render/checklist")
async def render_checklist(body: ChecklistBody):
    """渲染 → 上传 MinIO artifacts/checklist/<uuid>.docx → {key}。"""
    data = render_checklist_docx(body)
    key = f"artifacts/checklist/{uuid.uuid4()}.docx"
    await storage.put_bytes(key, data, content_type=_DOCX_CT)
    return {"key": key}


class ReportItem(BaseModel):
    level: str = ""                                # 风险等级标签（高/中，前端展示同款自由文案）
    title: str
    chapter: str = ""                              # 所在标书章节
    advice: str = ""                               # 整改建议


class RiskReportBody(BaseModel):
    title: str = "废标体检报告"
    project_name: str | None = None
    score: int | None = None
    high: int = 0
    mid: int = 0
    passed: int = 0
    items: list[ReportItem] = Field(default_factory=list)
    passed_items: list[str] = Field(default_factory=list)
    format: Literal["docx", "pdf"] = "docx"


def render_risk_report_docx(body: RiskReportBody) -> bytes:
    """废标体检报告 .docx：标题+项目/日期+评分总览，风险项四列表格（等级/风险项/章节/建议），
    已通过检查项列表，尾部 AI 辅助生成声明。与审核表同风格：python-docx 直出，确定性无 LLM。"""
    doc = Document()
    doc.add_heading(body.title, level=0)
    if body.project_name:
        doc.add_paragraph(f"项目名称：{body.project_name}")
    doc.add_paragraph(f"导出日期：{date.today().isoformat()}")
    if body.score is not None:
        doc.add_paragraph(
            f"体检评分：{body.score} 分　·　高风险 {body.high} 项　中风险 {body.mid} 项　已通过 {body.passed} 项")
    if body.items:
        doc.add_heading("风险项", level=1)
        t = doc.add_table(rows=len(body.items) + 1, cols=4)
        t.style = "Table Grid"
        for j, h in enumerate(("风险等级", "风险项", "所在章节", "整改建议")):
            t.rows[0].cells[j].text = h
        for i, item in enumerate(body.items, start=1):
            for j, v in enumerate((item.level or "—", item.title, item.chapter or "—", item.advice or "—")):
                t.rows[i].cells[j].text = v
    if body.passed_items:
        doc.add_heading("已通过检查项", level=1)
        for p in body.passed_items:
            doc.add_paragraph(f"✓ {p}")
    doc.add_paragraph("")
    doc.add_paragraph("本报告由 AI 辅助生成，仅供投标文件编制参考，请结合招标文件原文人工复核确认后使用。")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


class ReadReportBody(BaseModel):
    """标书分析报告（读标结论全量落 docx）。字段与 read result 同构但宽松（dict 直传），
    旧读标结果缺字段不至 422。"""
    title: str = "标书分析报告"
    project_name: str | None = None
    project_meta: dict = Field(default_factory=dict)
    categories: list[dict] = Field(default_factory=list)
    scoring: list[dict] = Field(default_factory=list)
    risk_summary: list[str] = Field(default_factory=list)
    required_structure: list[dict] = Field(default_factory=list)
    packages: list[dict] = Field(default_factory=list)


_META_LABELS = (("name", "项目名称"), ("code", "招标编号"), ("buyer", "采购人"),
                ("budget", "预算/最高限价"), ("deadline", "投标截止"), ("evaluation_method", "评标办法"))


def _item_flags(it: dict) -> str:
    """条目标志列：★不可偏离 / 废标风险 / 未明确（missing），无标志显示 —。"""
    flags = []
    if it.get("star"):
        flags.append("★不可偏离")
    if it.get("risk"):
        flags.append("废标风险")
    if it.get("status") == "missing":
        flags.append("未明确")
    return "、".join(flags) or "—"


def _read_report_categories(doc: Document, categories: list[dict]) -> None:
    """分类解读：每类一节，三列表格（条目/内容/标志）。"""
    for c in categories:
        items = c.get("items") or []
        doc.add_heading(str(c.get("title") or "分类"), level=1)
        if not items:
            doc.add_paragraph("（无条目）")
            continue
        t = doc.add_table(rows=len(items) + 1, cols=3)
        t.style = "Table Grid"
        for j, h in enumerate(("条目", "内容", "标志")):
            t.rows[0].cells[j].text = h
        for i, it in enumerate(items, start=1):
            for j, v in enumerate((str(it.get("title") or ""), str(it.get("value") or ""), _item_flags(it))):
                t.rows[i].cells[j].text = v


def _read_report_extras(doc: Document, body: ReadReportBody) -> None:
    """评分表 / 废标红线 / 投标文件构成 / 包件划分（有数据才出对应节）。"""
    if body.scoring:
        doc.add_heading("评分表", level=1)
        t = doc.add_table(rows=len(body.scoring) + 1, cols=3)
        t.style = "Table Grid"
        for j, h in enumerate(("评分项", "分值", "类别")):
            t.rows[0].cells[j].text = h
        for i, s in enumerate(body.scoring, start=1):
            star = "★" if s.get("star") else ""
            for j, v in enumerate((f"{star}{s.get('name') or ''}", str(s.get("score") or ""), str(s.get("category") or ""))):
                t.rows[i].cells[j].text = v
    if body.risk_summary:
        doc.add_heading("废标红线汇总", level=1)
        for r in body.risk_summary:
            doc.add_paragraph(f"⚠ {r}")
    if body.required_structure:
        doc.add_heading("投标文件构成清单", level=1)
        for s in body.required_structure:
            req = "必备" if s.get("required", True) else "可选"
            notes = f"（{s.get('notes')}）" if s.get("notes") else ""
            doc.add_paragraph(f"· [{req}] {s.get('title') or ''}{notes}")
    if body.packages:
        doc.add_heading("包件划分", level=1)
        for p in body.packages:
            budget = f"，预算 {p.get('budget')}" if p.get("budget") else ""
            notes = f"（{p.get('notes')}）" if p.get("notes") else ""
            doc.add_paragraph(f"· {p.get('name') or ''}{budget}{notes}")


def render_read_report_docx(body: ReadReportBody) -> bytes:
    """标书分析报告 .docx：项目要素 → 分类解读表格 → 评分表/红线/构成/包件 → AI 声明。
    确定性无 LLM（读标结论已存在，渲染免费），与审核表/体检报告同风格。"""
    doc = Document()
    doc.add_heading(body.title, level=0)
    if body.project_name:
        doc.add_paragraph(f"项目：{body.project_name}")
    doc.add_paragraph(f"导出日期：{date.today().isoformat()}")
    for key, label in _META_LABELS:
        v = body.project_meta.get(key)
        if v:
            doc.add_paragraph(f"{label}：{v}")
    _read_report_categories(doc, body.categories)
    _read_report_extras(doc, body)
    doc.add_paragraph("")
    doc.add_paragraph("本报告由 AI 辅助生成，仅供投标文件编制参考，请结合招标文件原文人工复核确认后使用。")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


@router.post("/render/read-report")
async def render_read_report(body: ReadReportBody):
    """渲染标书分析报告 → MinIO artifacts/report/<uuid>.docx → {key}。免计费（读标已收费）。"""
    data = render_read_report_docx(body)
    key = f"artifacts/report/{uuid.uuid4()}.docx"
    await storage.put_bytes(key, data, content_type=_DOCX_CT)
    return {"key": key}


@router.post("/render/risk-report")
async def render_risk_report(body: RiskReportBody):
    """渲染体检报告 → MinIO artifacts/report/<uuid>.docx|pdf → {key, format}。
    format=pdf 走 best-effort LibreOffice 转换（与 spec323 export 同链路）：失败回落 docx，
    format 字段如实返回实际产物格式——前端据此提示「PDF 转换失败，已导出 Word」。"""
    data = render_risk_report_docx(body)
    fmt = "docx"
    if body.format == "pdf":
        pdf = await asyncio.to_thread(docx_to_pdf, data)
        if pdf is not None:
            data, fmt = pdf, "pdf"
    key = f"artifacts/report/{uuid.uuid4()}.{fmt}"
    await storage.put_bytes(key, data, content_type=_PDF_CT if fmt == "pdf" else _DOCX_CT)
    return {"key": key, "format": fmt}
