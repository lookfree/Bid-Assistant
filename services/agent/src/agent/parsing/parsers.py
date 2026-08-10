from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import replace

from agent.parsing.docx_sections import (
    Block, heading_style_levels, is_bold_line, paragraph_level, split_docx_blocks,
)
from agent.parsing.types import ParsedDoc, UnsupportedDocument

# 章节标题启发式：第N章/第N节/第N篇/第N部分，或「一、二、」式顶层编号（标题一般较短）。
_HEADING = re.compile(r"^(第\s*[一二三四五六七八九十百零〇\d]+\s*[章节篇部分]|[一二三四五六七八九十]+\s*[、．.])")


def _is_heading(t: str) -> bool:
    return len(t) <= 40 and bool(_HEADING.match(t))


def _heading_level(t: str) -> int:
    """1 = 第N章/节/篇/部分；2 = 「一、」式顶层编号。只用于左栏层级渲染，判错至多是字号不对。"""
    return 1 if t.lstrip().startswith("第") else 2


def _split_clauses(lines: list[str]) -> tuple[list[dict], list[dict]]:
    """按章节标题分节、节内非空段落顺序编号 → ([{id: sec-N-cN, text}], [{sec, title, level}])。
    启发式（不 OCR/精排版）：无法识别章节时整体退化为 sec-1，供读标/提纲引用作 clause_ids 定位。
    标题**另存一份**、不进 clauses：混进去会挤掉条款序号，clause_id 口径一变，定位与引用全线受影响。"""
    clauses: list[dict] = []
    headings: list[dict] = []
    sec_n = 1
    sec_id = "sec-1"
    c_n = 0
    seen_heading = False
    for raw in lines:
        t = raw.strip()
        if not t:
            continue
        if _is_heading(t):
            if seen_heading or clauses:      # 前面已有章节或正文，才递增（首个标题保持 sec-1）
                sec_n += 1
            sec_id = f"sec-{sec_n}"
            headings.append({"sec": sec_id, "title": t, "level": _heading_level(t)})
            seen_heading = True
            c_n = 0
            continue                          # 标题本身不作为条款（但已记进 headings）
        c_n += 1
        clauses.append({"id": f"{sec_id}-c{c_n}", "text": t})
    return clauses, headings


def _docx_lines_in_order(d) -> tuple[list[Block], int]:
    """按文档顺序遍历正文块（段落与表格交错）→ (正文块, 正文内嵌图片张数)。表格行以 \t 连接单元格。

    每个块随手记下 Word 自己标的大纲层级与加粗（见 parsing/docx_sections.py）：章节切分要用它，
    而层级只有在这里手里还攥着段落元素时问得到，事后拿着纯文本行再问已经问不到了。
    2026-07-22 生产实测根因：招标文件的格式模板（授权委托书/应答一览表等）几乎都排在**表格**里，
    旧实现条款分句只喂段落 → 格式章只剩标题占节号、模板正文整段缺失（sec 空洞），
    内容生成拿不到招标模板原文只能自创格式。

    顺带数图片（w:drawing = 现代 DrawingML，w:pict = 旧 Word 的 VML）：正文里贴的证照/盖章
    扫描图一个字都提不出来，不数出来的话审查会把印在图上的材料判成「缺少」（见
    ParsedDoc.embedded_images）。**只走 body**，页眉页脚的 logo 天然不在其中。"""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = d.element.body
    styles = heading_style_levels(d)
    blocks: list[Block] = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, d)
            if p.text.strip():
                blocks.append(Block(p.text, level=paragraph_level(child, styles),
                                    bold=is_bold_line(p)))
        elif child.tag == qn("w:tbl"):
            for r in Table(child, d).rows:
                line = "\t".join(c.text for c in r.cells)
                if line.strip():
                    blocks.append(Block(line, table=True))
    images = sum(1 for _ in body.iter(qn("w:drawing"))) + sum(1 for _ in body.iter(qn("w:pict")))
    return blocks, images


def parse_docx(data: bytes) -> ParsedDoc:
    """解析 .docx 字节 → 段落文本 + 表格 + 条款 id（条款按文档顺序含表格行，见 _docx_lines_in_order）。
    章节切分走 docx 专用口径（Word 大纲层级优先，见 parsing/docx_sections.py）——
    pdf/xlsx 仍用 _split_clauses 的启发式。"""
    from docx import Document
    d = Document(io.BytesIO(data))
    blocks, images = _docx_lines_in_order(d)
    tables: list[list[list[str]]] = [[[c.text for c in r.cells] for r in t.rows] for t in d.tables]
    clauses, headings = split_docx_blocks(blocks)
    return ParsedDoc(text="\n".join(b.text for b in blocks), kind="docx", tables=tables,
                     embedded_images=images, clauses=clauses, headings=headings)


# 一页可见文字少于这么多字就当扫描图片页：正文页随便一段都远超，扫描页至多剩页眉页码。
_MIN_VISIBLE_CHARS = 20


# 混合页门槛：页里贴了图、可见文字又少于这么多字 → 也当扫描图片页。
# 2026-08-10 实测形态：「法定代表人授权委托书扫描件如下：」（21 字）+ 整版证照图。可见字数过得了
# 上面 20 字的门槛，实质内容却一个字都提不出来——不认成图片页就既不 OCR 也不进「无法核验」统计，
# 直接退回治理之前的误判（把印在图里的材料判成「缺少」）。
# 100 字是**保守**门槛：正常正文页随便一段都远超它，带 logo 页眉/落款章的文本页（几百字）不会误伤。
_MIXED_PAGE_MAX_CHARS = 100


def visible_len(text: str) -> int:
    """「可见文字」字数：空白/换行不算字，否则满页空行的扫描页会被当成有内容。
    页面扫描判定、OCR 识别结果的判定、read 的 docx 拒绝闸共用它，口径单点在此。"""
    return len("".join(text.split()))


def is_image_page(text: str) -> bool:
    """这一页是不是提不出可见文字（= 扫描图片页）。
    阈值**单点在此**——页数统计、OCR 选页与识别结果的可读判定共用同一判据，改一处全线生效。"""
    return visible_len(text) < _MIN_VISIBLE_CHARS


def _has_image_xobject(page) -> bool:
    """这一页的资源里有没有位图（/Resources /XObject 下 /Subtype = /Image）。

    **只对可见字数落在灰区（20–100 字）的页调用**：解析 XObject 要把间接对象读出来，
    而位图的间接对象一读就是整条图片流进 pypdf 的对象缓存——整份文件逐页做，139 页的扫描件
    就是几百 MB 白吃。灰区页本来就少，短路之后这条路几乎不走。
    任何异常都当「没有图」：这是加严判定的辅助信号，坏 xref 不该把整份文件的解析拖垮。"""
    try:
        res = page.get("/Resources")
        xobjects = res.get_object().get("/XObject") if res is not None else None
        if xobjects is None:
            return False
        return any(v.get_object().get("/Subtype") == "/Image"
                   for v in xobjects.get_object().values())
    except Exception:  # noqa: BLE001 见 docstring：判不出来就当没有图，绝不抛给解析入口
        return False


def _scanned_flags(pages, texts: list[str]) -> list[bool]:
    """逐页「模型看不看得见这页」：纯字数门槛 → 灰区页再查有没有贴图（见 _has_image_xobject）。"""
    flags: list[bool] = []
    for pg, t in zip(pages, texts):
        n = visible_len(t)
        flags.append(n < _MIN_VISIBLE_CHARS
                     or (n < _MIXED_PAGE_MAX_CHARS and _has_image_xobject(pg)))
    return flags


def scanned_page_indices(doc: ParsedDoc) -> list[int]:
    """doc 里扫描图片页的页序号（0 基）。OCR 只对这些页做——文本页绝不重复识别。
    优先用解析时定下的逐页判定（含混合页），没有那份信息时退回纯字数判据。"""
    if doc.image_page_flags:
        return [i for i, f in enumerate(doc.image_page_flags) if f]
    return [i for i, t in enumerate(doc.page_texts) if is_image_page(t)]


def parse_pdf(data: bytes) -> ParsedDoc:
    """解析 .pdf 字节 → 逐页文本(拼接) + 页数 + 扫描图片页数 + 条款 id。
    提不出文字的页在这里只计数、不猜内容；能不能识别出来是下一步的事（parsing/ocr.py 配置驱动，
    未配置 OCR 服务时一切照旧）——下游据此说「无法核验」，而不是当成「没有这份材料」。"""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = [(pg.extract_text() or "") for pg in reader.pages]
    flags = _scanned_flags(reader.pages, pages)
    text = "\n".join(pages)
    clauses, headings = _split_clauses(text.split("\n"))
    return ParsedDoc(text=text, kind="pdf", pages=len(reader.pages), image_pages=sum(flags),
                     page_texts=pages, image_page_flags=flags, clauses=clauses, headings=headings)


# 识别文本插回正文时的页首标记：既告诉模型这段字是从扫描页认出来的（可能带识别误差），
# 也让它知道这一页**确实有内容**，不再当成看不见。
_OCR_PAGE_MARK = "[第{n}页·扫描件识别]"


def _recognized(ocr_text: str, page_text: str) -> bool:
    """这一页的识别结果算不算「已看见」（可以从 image_pages / 注记里扣掉）。两道门槛取高的那个：

    · **够读**（_MIN_VISIBLE_CHARS）：糊掉的章、被当成字符的花纹常常只出一两个字，
      那种页用户实际上还是什么都看不见；
    · **不少于原页本来就有的可见字数**：混合页（20–99 字真实文本 + 整版贴图）只认出一枚印章时，
      那 22 个字并不代表整版图看见了——实质内容照旧提不出来，「无法核验」的注记必须留着。
      纯扫描页原文≈0，max 取到的仍是前一条，行为与只有前一条时逐字节一致。
    """
    return visible_len(ocr_text) >= max(_MIN_VISIBLE_CHARS, visible_len(page_text))


def splice_ocr_pages(doc: ParsedDoc, ocr_texts: dict[int, str]) -> ParsedDoc:
    """OCR 文本按页插回原位置 → 新的 ParsedDoc（{页序号: 识别文字}，0 基）。

    条款与标题**必须重算**：识别出来的字要和正文一样参与章节切分和预算，只改 text 不改 clauses
    的话，下游按 clauses 聚章（nodes/common.py::parse_bid_docs），认出来的内容一个字都进不了
    审查材料。image_pages 扣掉已识别的页——剩下的才是真正还看不见的页（全部识别成功 → 注记消失）。

    「已识别」的判据见 _recognized；不够格的页**既不扣 image_pages，也一个字都不拼回**：
    拼回去的话，一份全扫描的废件会凭那几个「章」字切出条款、聚出非空 chapters，
    绕过 review 的「解析不出正文 → run 失败 + 全额退款」闸——用户为一份什么都看不见的文件
    付了一次审查的钱。

    拼回的方式按页分两种：纯扫描页（原文提不出可见文字）**替换**；混合页（原页本来就有可见文字）
    **追加**——原页那几十字是精确可选文本（投标人名称、法定代表人、日期、电话），审查要拿它逐字
    比对，换成一段近似识别就成了拿错字去比（实测「有限」→「有眼」）。追加会让标题类文字重复一遍，
    比丢掉原文划算得多。
    """
    readable = {i: t for i, t in ocr_texts.items() if _recognized(t, doc.page_texts[i])}
    if not readable:
        return doc
    pages = list(doc.page_texts)
    for i, text in readable.items():
        mark = _OCR_PAGE_MARK.format(n=i + 1)
        pages[i] = f"{mark}\n{text}" if is_image_page(pages[i]) else f"{pages[i]}\n{mark}\n{text}"
    full = "\n".join(pages)
    clauses, headings = _split_clauses(full.split("\n"))
    # 逐页判定同步扣掉已识别的页：留着旧值的话，谁再拿这份 doc 问「哪几页看不见」都会拿到
    # 一个与 image_pages 对不上的答案（识别过的页仍在名单里）。
    flags = [f and i not in readable for i, f in enumerate(doc.image_page_flags)]
    return replace(doc, text=full, page_texts=pages, clauses=clauses, headings=headings,
                   image_page_flags=flags,
                   image_pages=max(0, doc.image_pages - len(readable)))


def parse_xlsx(data: bytes) -> ParsedDoc:
    """解析 .xlsx 字节 → 各表非空行文本 + 表格结构 + 条款 id。"""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    tables: list[list[list[str]]] = []
    lines: list[str] = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                rows.append(cells)
                lines.append("\t".join(cells))
        if rows:
            tables.append(rows)
    clauses, headings = _split_clauses(lines)
    return ParsedDoc(text="\n".join(lines), kind="xlsx", tables=tables,
                     clauses=clauses, headings=headings)


_LEGACY_TARGET = {"doc": "docx", "xls": "xlsx"}


def _convert_legacy(data: bytes, ext: str) -> tuple[bytes, str]:
    """经 LibreOffice headless 把旧格式 .doc/.xls 转成 .docx/.xlsx 字节（spec320）。
    soffice 缺失或转换失败/超时 → 抛 UnsupportedDocument，调用方（多文件读标）按文件降级跳过，不崩整体。"""
    target_ext = _LEGACY_TARGET[ext]
    if shutil.which("soffice") is None:
        raise UnsupportedDocument(f"缺少 soffice，无法转换 .{ext}")
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, f"input.{ext}")
        with open(src, "wb") as f:
            f.write(data)
        try:
            # 每次转换独立 UserInstallation profile：默认 profile 有单实例锁，
            # 多个 .doc/.xls 并发转换会互相拿不到锁而静默失败（评审 Important 项）。
            profile = os.path.join(tmp, "lo-profile")
            subprocess.run(
                ["soffice", "--headless", f"-env:UserInstallation=file://{profile}",
                 "--convert-to", target_ext, "--outdir", tmp, src],
                timeout=60, check=True, capture_output=True,
            )
        except (subprocess.SubprocessError, OSError) as e:
            raise UnsupportedDocument(f".{ext} 转换失败: {e}") from e
        out_path = os.path.join(tmp, f"input.{target_ext}")
        if not os.path.exists(out_path):
            raise UnsupportedDocument(f".{ext} 转换未产出文件")
        with open(out_path, "rb") as f:
            return f.read(), target_ext


def _parse_doc(data: bytes) -> ParsedDoc:
    converted, _ = _convert_legacy(data, "doc")
    return parse_docx(converted)


def _parse_xls(data: bytes) -> ParsedDoc:
    converted, _ = _convert_legacy(data, "xls")
    return parse_xlsx(converted)


_DISPATCH = {"docx": parse_docx, "pdf": parse_pdf, "xlsx": parse_xlsx,
             "doc": _parse_doc, "xls": _parse_xls}


# 已知的文档透明加密（DLP）软件封装头：文件被整体加密成密文，扩展名不变、内容已不是原格式。
# 2026-08-05 生产实测：一份 7.6MB 的 .pdf 头是 %TSD-Header-###%、无 %%EOF，pypdf 只报
# 「流意外结束」，读标据此以为是瞬时问题而降级到工具兜底，最后把文件问题报成了模型问题。
_ENCRYPTED_WRAPPERS: tuple[bytes, ...] = (b"%TSD-Header-###%",)

# 扩展名 → 必须出现在头部的魔数。只列魔数无歧义的格式：doc/xls 故意不列——.doc 里装 RTF、
# 装 docx 都是历史常见写法，LibreOffice 照样能转，在这里强判会误伤本来能用的文件。
_REQUIRED_MAGIC: dict[str, bytes] = {
    "pdf": b"%PDF-",
    "docx": b"PK\x03\x04",
    "xlsx": b"PK\x03\x04",
}

# 取样长度：PDF 规范容忍文件头前有少量前导字节（pypdf 也容忍），故在头部一段范围内找而非要求偏移 0。
_MAGIC_SAMPLE_BYTES = 1024


def _check_magic(data: bytes, ext: str) -> None:
    """内容与扩展名不符就直接抛出可读原因。上传入口（App API）已按同一套规则拦过一道，
    这里是纵深防御：覆盖上传校验上线前已存的老文件，以及其它入口进来的文件。"""
    head = data[:_MAGIC_SAMPLE_BYTES]
    for sig in _ENCRYPTED_WRAPPERS:
        if head.startswith(sig):
            raise UnsupportedDocument(
                "文件已被文档加密软件封装成密文（扩展名未变，内容已不是原格式）。"
                "请在加密软件中对该文件走解密/外发流程，导出明文后重新上传。"
                "注意：在装有加密客户端的电脑上能正常打开，不代表文件本身是明文")
    magic = _REQUIRED_MAGIC.get(ext)
    if magic and magic not in head:
        raise UnsupportedDocument(
            f"文件内容与扩展名 .{ext} 不符（未找到该格式的文件头），"
            f"可能被加密软件封装、下载/上传未完成，或扩展名被改过")


def parse_bytes(data: bytes, filename: str) -> ParsedDoc:
    """按文件扩展名分发到对应解析器；不支持的类型、内容与扩展名不符都抛 UnsupportedDocument。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    fn = _DISPATCH.get(ext)
    if not fn:
        raise UnsupportedDocument(f"不支持的文档类型: .{ext}")
    _check_magic(data, ext)
    return fn(data)
