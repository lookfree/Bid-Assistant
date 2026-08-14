from __future__ import annotations
import base64
import copy
import io
import re
from typing import Callable
from bs4 import BeautifulSoup
from docx import Document
from agent.agents.bidding_agent.render.sanitize import normalize_chapter_html, strip_document_shell, strip_template_disclaimers
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt, RGBColor

_CONTAINERS = ("div", "section", "article", "body")

# H1-H5 → 磅值（章/节/小节/细分/明细五级，见 _apply_bid_styles）。五级都必须入表：
# 未配置的 Word 内建标题样式会继承主题蓝/西文字体（三级提纲落地时 h4 首次可达，五级后 h5 同理）。
_HEADING_SIZES = {
    "Heading 1": Pt(16), "Heading 2": Pt(14), "Heading 3": Pt(12), "Heading 4": Pt(12), "Heading 5": Pt(12),
}


def _apply_bid_styles(doc: Document) -> None:
    """标书排版惯例（一次性设在 Document 的样式上，覆盖 python-docx 默认模板）：
    正文宋体小四(12pt)；一/二/三级标题黑体加粗黑色——Word 默认标题走主题色蓝，
    投标文件要求严肃的黑白配色，不能保留默认蓝。
    注：服务端镜像目前只装了 fonts-noto-cjk（没有宋体/黑体字体文件），LibreOffice
    转 PDF 时找不到这两个字体名会退回 Noto CJK 渲染；用户在 Word 里打开 .docx 本身
    是原生渲染，不受影响。"""
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for style_name, size in _HEADING_SIZES.items():
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def _emit_placeholder_image(doc: Document, el, fetch_object: Callable[[str], bytes | None] | None) -> None:
    """附录占位图取字节落图（2026-08-09 资质附录系统章节 Plan A①）：`<img data-object-key>`
    无 src 无字节，经 fetch_object 按 key 现取；fetch_object 缺省/取不到/取字节抛错（MinIO
    404/网络抖动）/取到但 add_picture 解码失败——统一落「（图片加载失败：alt）」占位行，
    best-effort，与既有 data: 坏图分支同语义，不中断整本渲染。"""
    key = el.get("data-object-key", "")
    alt = el.get("alt") or key
    data = None
    if fetch_object is not None:
        try:
            data = fetch_object(key)
        except Exception:  # noqa: BLE001 取字节失败——占位保文
            data = None
    if data is None:
        doc.add_paragraph(f"（图片加载失败：{alt}）")
        return
    try:
        doc.add_picture(io.BytesIO(data), width=Inches(6))
    except Exception:  # noqa: BLE001 坏字节/非受支持图片格式
        doc.add_paragraph(f"（图片加载失败：{alt}）")


_ALIGN_MAP = {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
_ALIGN_RE = re.compile(r"text-align\s*:\s*(center|right)")


def _apply_align(paragraph, el) -> None:
    """行内 style 的 text-align（center/right）落到段落对齐。表单抬头（「响   应   函」）、
    落款、日期栏靠这个才能与招标模板同样居中/靠右——此前渲染层完全无视对齐，编辑器里
    居中的文字导出也统统变左对齐（2026-08-13 用户实测「响应函少了居中的标题」）。
    居中/靠右段不吃首行缩进：自定义格式把缩进设在 Normal 上，带缩进的"居中"是歪的。"""
    m = _ALIGN_RE.search(el.get("style") or "")
    if m:
        paragraph.alignment = _ALIGN_MAP[m.group(1)]
        paragraph.paragraph_format.first_line_indent = Pt(0)


def _emit_el(doc: Document, el, fetch_object: Callable[[str], bytes | None] | None = None) -> None:
    """单个 HTML 元素 → docx：h1/h2→Heading2、h3/h4→Heading3、p→段落、ul/li→项目符号、
    table→表格；容器标签（div 等）递归展开，防止整块被 get_text 压扁成一段。
    fetch_object：附录占位图取字节回调，随递归原样透传给子节点。"""
    name = getattr(el, "name", None)
    if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        # 章内层级绝对映射（评审二轮:相对归一会让同一 <h4> 在不同章落不同 Word 级）。
        # 写手契约与提纲五级一一对应：二级节 h3 / 三级小节 h4 / 四级细分 h5 / 五级明细 h6，
        # 章标题占 Heading 1，故正文各级整体下移一位 → Word 2/3/4/5。
        # 旧文档整章全 <h3> 同样得到 章(1)→节(2) 的层级（平级问题就此修复）；
        # h1/h2 是跑偏时的防御位（按节待遇），绝不落成正文段落。
        _apply_align(doc.add_heading(
            el.get_text(strip=True),
            level={"h1": 2, "h2": 2, "h3": 2, "h4": 3, "h5": 4, "h6": 5}.get(name, 5),
        ), el)
    elif name == "p":
        _apply_align(doc.add_paragraph(el.get_text(strip=True)), el)
        for img in el.find_all("img"):  # 光标处插图常嵌在段落里，只取文字会把图整个丢掉
            _emit_el(doc, img, fetch_object)
    elif name == "ul":
        for li in el.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
    elif name == "ol":
        # 有序列表（spec329 编辑器可产出）：不接会整段压扁成一行无编号文本
        for li in el.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text(strip=True), style="List Number")
    elif name == "table":
        _emit_table(doc, el)
    elif name == "img":
        # 用户在编辑器插入的图片（data URL 内嵌，spec 无外链图）：解码落图；坏图跳过不阻断整本渲染
        src = el.get("src", "")
        if src.startswith("data:image/"):
            try:
                doc.add_picture(io.BytesIO(base64.b64decode(src.split(",", 1)[1])), width=Inches(5.5))
            except Exception:  # noqa: BLE001 base64 破损/格式不支持——丢图保文
                pass
        elif el.get("data-object-key"):  # 附录占位图（无 data: src，只带 MinIO key）
            _emit_placeholder_image(doc, el, fetch_object)
    elif name in _CONTAINERS:
        for child in el.children:
            _emit_el(doc, child, fetch_object)
    elif text := el.get_text(strip=True):
        doc.add_paragraph(text)


# Word 表格的实用列上限：模型笔误写出 colspan="999" 时用来夹紧，不让一个笔误撑出上千列。
_MAX_COLS = 64


def _span(raw, limit: int) -> int:
    """HTML 的 colspan/rowspan 取值 → 合法跨度。

    模型偶尔写出非数字（2026-08-06 生产实测 `rowspan="wer"`）或离谱的大数。直接 int() 会让
    整个导出以一句 Python 异常收场——用户看到的是 `invalid literal for int() with base 10`，
    连点九次导出、每次 0.2 秒就崩，既不知道问题在哪也无从自救。
    解析不了/非正数 → 1（不合并；宁可表格少一次合并，也不能整本标书导不出来）；
    超出上限 → 夹到上限（那正是「合并到底」的本意，比丢掉合并更接近模型想表达的东西）。
    """
    try:
        n = int(str(raw).strip())
    except (TypeError, ValueError):
        return 1
    if n < 1:
        return 1
    return min(n, limit)


def _emit_table(doc: Document, el) -> None:
    """HTML 表 → docx 表,支持 colspan/rowspan（spec329 合并单元格,审查修正）：
    先按占位矩阵展开网格定位（合并格占多个格位）,再用 python-docx cell.merge 合并——
    旧实现按 td 个数顺位填格,一行里有合并格时后续列整体左移错位。"""
    rows = el.find_all("tr")
    if not rows:
        return
    grid: list[list[dict | None]] = []  # 每格: {"text":…, "anchor":(r,c)} 或 None
    for ri, r in enumerate(rows):
        while len(grid) <= ri:
            grid.append([])
        ci = 0
        for cell in r.find_all(["td", "th"]):
            while ci < len(grid[ri]) and grid[ri][ci] is not None:
                ci += 1  # 跳过上方 rowspan 占掉的格位
            # 跨度上限：列取实用上限，行取「本行往下还剩几行」——rowspan 再大也不可能超过表格本身
            cs = _span(cell.get("colspan"), _MAX_COLS)
            rs = _span(cell.get("rowspan"), max(1, len(rows) - ri))
            for dr in range(rs):
                while len(grid) <= ri + dr:
                    grid.append([])
                row = grid[ri + dr]
                while len(row) < ci + cs:
                    row.append(None)
                for dc in range(cs):
                    row[ci + dc] = {"text": cell.get_text(strip=True) if (dr == 0 and dc == 0) else "",
                                    "anchor": (ri, ci)}
            ci += cs
    cols = max(len(r) for r in grid)
    t = doc.add_table(rows=len(grid), cols=cols)
    t.style = "Table Grid"   # 网格线：偏差表/报价表没有边框不可读（e2e PDF 实测）
    merged: set[tuple[int, int]] = set()
    for ri, row in enumerate(grid):
        for ci in range(cols):
            g = row[ci] if ci < len(row) else None
            if g is None:
                continue
            ar, ac = g["anchor"]
            if (ar, ac) == (ri, ci):
                cell = t.rows[ri].cells[ci]
                cell.text = g["text"]
                # 单元格不吃正文首行缩进（spec330 格式把缩进设在 Normal 上,窄列缩 2 字符换行灾难）
                cell.paragraphs[0].paragraph_format.first_line_indent = Pt(0)
            elif (ar, ac) not in merged:
                t.cell(ar, ac).merge(t.cell(ri, ci))
                merged.add((ar, ac))
            else:
                t.cell(ar, ac).merge(t.cell(ri, ci))


def _emit_html(doc: Document, html: str, fetch_object: Callable[[str], bytes | None] | None = None) -> None:
    """HTML 最小映射到 docx。复杂样式（行内富文本等）为后续加固项。"""
    soup = BeautifulSoup(html or "", "html.parser")
    for el in soup.children:
        _emit_el(doc, el, fetch_object)


def _cover_line(doc: Document, text: str, size: int) -> None:
    """封面居中一行：统一走 run 设字号，标题行调用方另设加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)  # 封面居中行不吃 Normal 首行缩进（spec330）
    run = p.add_run(text)
    run.font.size = Pt(size)


# project_meta 键名归一：读标 schema 里是自由 dict,历史数据用中文键(项目名称/采购编号/采购人),
# 渲染读英文键(name/code/buyer)导致封面/页眉落兜底(e2e 实测)。取值时按别名依次找。
_META_ALIASES = {
    "name": ("name", "项目名称", "项目名"),
    "code": ("code", "采购编号", "招标编号", "项目编号"),
    "buyer": ("buyer", "采购人", "招标人", "采购单位"),
}


def _norm_meta(meta: dict) -> dict:
    out = dict(meta)
    for key, aliases in _META_ALIASES.items():
        if not out.get(key):
            val = next((meta[a] for a in aliases if meta.get(a)), None)
            if val:
                out[key] = val
    return out


def _style_cover(doc: Document, meta: dict, package: dict | None = None) -> None:
    """封面：居中大标题（项目名）+ 信息块（采购人/编号/日期占位）+ 投标人盖章占位。
    package 存在（选包，spec324）⇒ 项目名下加「包件：《name》」一行；未选包时逐字节不变。"""
    _cover_line(doc, meta.get("name", "投标文件"), 26)
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph()
    if package and package.get("name"):
        _cover_line(doc, f"包件：《{package['name']}》", 14)
    if meta.get("buyer"):
        _cover_line(doc, f"采购人：{meta['buyer']}", 14)
    if meta.get("code"):
        _cover_line(doc, f"招标编号：{meta['code']}", 14)
    _cover_line(doc, f"日期：{meta.get('date', '____年__月__日')}", 14)
    doc.add_paragraph()
    _cover_line(doc, "投标人：____________________（盖章）", 14)
    doc.add_page_break()


def _fld_char(kind: str):
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), kind)
    return el


def _fld_instr(instr_text: str):
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")   # 保住 instr 里的空格（TOC 开关之间的分隔）
    instr.text = instr_text
    return instr


def _add_field(paragraph, instr_text: str) -> None:
    """在段落里插入一个 Word 域（fldChar begin/instrText/separate/end 四件套 OXML）；
    页脚 PAGE 域用这个整段版本；TOC 域要把缓存条目夹在 separate 与 end 之间，
    用同一套 _fld_char/_fld_instr 零件自行拼开口版（评审 2026-08-13：域构造只许一份实现）。"""
    r = paragraph.add_run()._r
    for node in (_fld_char("begin"), _fld_instr(instr_text), _fld_char("separate"), _fld_char("end")):
        r.append(node)


def _add_toc_field(doc: Document):
    """真目录域 + 可回填的缓存区：TOC \\o "1-4" \\h \\z \\u。

    域拆成两个段落——begin/instrText/separate 一段、end 一段，正文渲染完由
    _fill_toc_entries 把实际产出的标题静态写进两者**之间**（Word 域的缓存结果区）。
    为什么不能只留空域赌「打开时自动更新」（2026-08-11 的做法）：Word 只是弹一次确认，
    WPS 和导出的 PDF 根本不理 updateFields，用户看到的仍是空白目录页（2026-08-12
    用户二次实测「还是没有目录」）。缓存条目是普通文档内容，任何查看器直接可见；
    在 Word 里更新域会把缓存区整体重建成带页码的目录，两头都不吃亏。
    返回持有 fldChar end 的段落，供回填时定位插入点。"""
    doc.add_heading("目录", level=1)
    field_p = doc.add_paragraph()
    # 目录只收到四级：五级明细（① 值班安排）进目录会把目录撑得比正文还碎，
    # 评标专家反而找不到重点——五级仍在正文里有层级，只是不进目录。
    run = field_p.add_run()._r
    for node in (_fld_char("begin"), _fld_instr('TOC \\o "1-4" \\h \\z \\u'), _fld_char("separate")):
        run.append(node)
    end_p = doc.add_paragraph()
    end_p.add_run()._r.append(_fld_char("end"))
    _set_update_fields_on_open(doc)
    doc.add_page_break()
    return end_p


_TOC_CACHE_LEVELS = 3      # 静态缓存条目只收三级：没有页码的四级明细堆一页纯属噪音
_TOC_INDENT_CM = 0.74      # 每深一级缩进一档（与 Word 默认 TOC 样式的视觉一致）


def _fill_toc_entries(doc: Document, toc_end_p) -> None:
    """把正文里实际渲染出的标题回填进目录域缓存区（fldChar separate 与 end 之间）。
    条目取自成品文档本身（Heading 1..N 段落），而不是另算一份提纲——目录必须和正文
    逐字一致，两处各算一遍就会慢慢长歪（附录/签章页这类系统章也要进目录）。"""
    passed_end = False
    entries: list[tuple[int, str]] = []
    for p in doc.paragraphs:
        if p._p is toc_end_p._p:
            passed_end = True
            continue
        if not passed_end:
            continue   # 域之前只有封面和「目录」标题自己——目录不给自己列条目
        m = re.fullmatch(r"Heading (\d)", p.style.name or "")
        if m and int(m.group(1)) <= _TOC_CACHE_LEVELS and p.text.strip():
            entries.append((int(m.group(1)), p.text.strip()))
    for level, text in entries:
        entry = toc_end_p.insert_paragraph_before(text)
        entry.paragraph_format.left_indent = Cm(_TOC_INDENT_CM * (level - 1))
        entry.paragraph_format.first_line_indent = Pt(0)   # 不吃 Normal/自定义格式的首行缩进
        for run in entry.runs:
            run.font.size = Pt(10.5)


def _set_update_fields_on_open(doc: Document) -> None:
    """让 Word 打开文档时自动更新所有域（settings.xml 的 w:updateFields）。

    没有这一句，目录域打开后是**空的**——页码只有 Word 的排版引擎算得出，域本身不含内容，
    用户看到的就是「导出的文档没有目录」（2026-08-11 用户实测反馈）。此前靠正文里写一句
    「请按 F9 更新目录」提示用户手动更新，等于把我们的实现细节转嫁给用户，还得留一行废话在
    标书正文里（那是要交给评委的文件）。页脚的 PAGE 页码域同理受益。
    Word 打开时会弹一次「是否更新域」确认，选是即可；WPS 与 LibreOffice 同样认这个开关。
    """
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    flag = OxmlElement("w:updateFields")
    flag.set(qn("w:val"), "true")
    settings.append(flag)


def _add_page_number_footer(doc: Document, project_name: str) -> None:
    """默认节：页眉写项目名、页脚居中 PAGE 域页码（逐页连续编码，招标方常见硬要求）。"""
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = project_name
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.first_line_indent = Pt(0)  # 页脚页码不吃 Normal 首行缩进（spec330）
    _add_field(footer_p, "PAGE")


def _add_ai_notice(doc: Document) -> None:
    """文档末尾生成说明：备案要求的显式标识，导出环节自动写入（用户定稿时可自行删除）。"""
    doc.add_paragraph()
    p = doc.add_paragraph("本内容由智启元投标助手生成合成类算法辅助生成，仅供投标文件编制参考，请结合招标文件原文和企业实际情况复核确认后使用。")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# spec330 输出格式：GB 字号 → 磅值;默认参数=用户 2026-07-23 提供的口径。
# fmt=None（不传）→ 维持现行样式,与既有导出一致;传 fmt（含空 dict）→ 以默认值起底逐项覆盖。
_GB_PT = {"三号": 16, "四号": 14, "小四": 12, "五号": 10.5}
_FMT_DEFAULT = {
    "margin_cm": {"top": 2.2, "bottom": 2.2, "left": 2.3, "right": 2.3},
    "heading_font": "宋体", "heading_size": "四号", "heading_bold": True,
    "body_font": "宋体", "body_size": "小四", "body_indent_chars": 2,
    "line_spacing": 1.5,  # 1 / 1.5 / "fixed22"（固定 22 磅）
}


def _set_line_spacing(pf, spacing) -> None:
    if spacing == "fixed22":
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(22)
    else:
        pf.line_spacing = float(spacing)


def _apply_custom_format(doc: Document, fmt: dict) -> None:
    """按用户输出格式覆盖样式（spec330）：A4 纵向 + 页边距 + 正文/标题字体字号缩进行距。
    只在显式传 fmt 时调用;逐项以 _FMT_DEFAULT 起底,用户改哪项覆盖哪项。"""
    f = {**_FMT_DEFAULT, **{k: v for k, v in fmt.items() if v is not None}}
    m = {**_FMT_DEFAULT["margin_cm"], **(f.get("margin_cm") or {})}
    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(21), Cm(29.7)  # A4 纵向
        sec.top_margin, sec.bottom_margin = Cm(float(m["top"])), Cm(float(m["bottom"]))
        sec.left_margin, sec.right_margin = Cm(float(m["left"])), Cm(float(m["right"]))
    body_pt = _GB_PT.get(f["body_size"], 12)
    normal = doc.styles["Normal"]
    normal.font.name = f["body_font"]
    normal.font.size = Pt(body_pt)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), f["body_font"])
    # 首行缩进 N 字符 = N × 字号;行距设在 Normal 段落格式上,全文（含标题继承前的基准）统一。
    # 缩进溢入表格/封面/页脚的问题由各发射点显式置零解决（_emit_table 单元格、_cover_line、页脚）。
    normal.paragraph_format.first_line_indent = Pt(body_pt * int(f["body_indent_chars"]))
    _set_line_spacing(normal.paragraph_format, f["line_spacing"])
    head_pt = _GB_PT.get(f["heading_size"], 14)
    for style_name in _HEADING_SIZES:
        style = doc.styles[style_name]
        style.font.name = f["heading_font"]
        style.font.size = Pt(head_pt)
        style.font.bold = bool(f["heading_bold"])
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), f["heading_font"])
        style.paragraph_format.first_line_indent = Pt(0)  # 标题首行缩进 0 字符、左对齐
        _set_line_spacing(style.paragraph_format, f["line_spacing"])


# 复印章证照锚定（2026-08-14 授权书截图立案）：引导行组名 → 锚定人名词——
# **与线上就位共用 cert_placement.id_person_words 一份**（两处各养必然漂移）;
# 非身份证类（营业执照等）没有粘贴框,不锚。


# 图 alt 的条目标题前缀（cert_placement._image_alt 写成「标题|ocr 摘要」）
_ALT_TITLE_RE = re.compile(r'alt="([^"|]{1,30})[|"]')


def _parse_cert_groups(tail: str) -> list[tuple[str, str, list[str]]]:
    """证照尾巴 → [(组名, 引导行 HTML, [图块…])]。一条引导行统辖其后所有无引导图块——
    证照条目正反面=1 行引导+N 张图(评审四轮 F1:按块劈开会把反面孤儿在章末)。
    **引导行缺失时从图 alt 前缀(条目标题)反查组名**(评审六轮 F1,致命回归):线上框位
    替换后的稿子没有引导行(用户口径:更干净),导出锚定全靠这一步找回标签——丢了标签,
    图退回章尾,入框特性整体倒退。alt 同标题的连续无引导块并成一组(正反面同组)。"""
    from agent.agents.bidding_agent.nodes.cert_placement import _group_of
    from agent.agents.bidding_agent.render.form_copier import CERT_BLOCK_RE

    groups: list[tuple[str, str, list[str]]] = []
    for m in CERT_BLOCK_RE.finditer(tail):
        lead, label = m.group(1) or "", m.group(2) or ""
        body = m.group(0)[len(lead):].lstrip() if lead else m.group(0)
        if not lead:
            am = _ALT_TITLE_RE.search(body)
            label = (_group_of(am.group(1)) or "") if am else ""
        if lead or not groups or (label and label != groups[-1][0]):
            groups.append((label, lead, [body]))
        else:
            groups[-1][2].append(body)
    return groups


def _cell_text(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t")))


def _find_cert_anchor(nodes: list, label: str):
    """本章嫁接节点里找该组的粘贴框 → (稳定键, 落点元素, 模式) 或 None。
    段落框=("after")图插段后;表格里的框格=("cell")图直接进格——整表后插会离框一页远
    (评审四轮 F3)。键用 (节点序, 格序)——lxml 代理 id 跨迭代不稳定,绝不当键。
    判据:人名词＋「身份证」＋「粘贴」同段/同格（.//w:t 连文本框内壁一起收）。"""
    from agent.agents.bidding_agent.nodes.cert_placement import id_person_words

    words = id_person_words(label)
    if not words:
        return None
    for i, el in enumerate(nodes):
        if el.tag == qn("w:tbl"):
            for j, tc in enumerate(el.iter(qn("w:tc"))):
                text = _cell_text(tc)
                if "身份证" in text and "粘贴" in text and any(w in text for w in words):
                    return ((i, j), tc, "cell")
            continue
        text = _cell_text(el)
        if "身份证" in text and "粘贴" in text and any(w in text for w in words):
            return ((i, -1), el, "after")
    return None


from agent.agents.bidding_agent.render.form_copier import MC_NS as _MC_NS

_A_EXT = "{http://schemas.openxmlformats.org/drawingml/2006/main}ext"
# wordprocessingDrawing 命名空间单份（评审六轮 F7：extent/positionH 各自手抄必漂移）
_WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
_WP_EXTENT = _WP_NS + "extent"


def _box_x(choice) -> tuple[str, int] | None:
    """框的水平位置 →（基准 relativeFrom, 偏移 EMU）；拿不到返回 None。
    基准必须一并带出（评审六轮 F3）：page 基准的 914400 与 column 基准的 0 没有可比性，
    裸数值排序会把版面右框排到左边。"""
    pos = next(iter(choice.iter(_WP_NS + "positionH")), None)
    if pos is None:
        return None
    po = next(iter(pos.iter(_WP_NS + "posOffset")), None)
    try:
        off = int((po.text or "").strip()) if po is not None else None
    except ValueError:
        return None
    return (pos.get("relativeFrom") or "", off) if off is not None else None


def _anchor_boxes(anchor) -> list[dict]:
    """锚元素里的粘贴框。**按水平位置左→右排序**（2026-08-14 用户终验：正面照要进
    左框，而 XML 里右框常排在前——文档序≠版面序）；任一框拿不到位置则保持文档序。
    每框收 Choice 内容区、Fallback 内容区（评审五轮 C2：LibreOffice 转 PDF 可能读
    降级层，只装 Choice 的话 PDF 预览里扫描件整体消失——两层同装镜像）、框尺寸与
    框上文字（按人配框用）。"""
    out = []
    for ac in anchor.iter(_MC_NS + "AlternateContent"):
        choice = next((c for c in ac if c.tag == _MC_NS + "Choice"), None)
        if choice is None:
            continue
        tx = next(iter(choice.iter(qn("w:txbxContent"))), None)
        if tx is None:
            continue
        fb = next((c for c in ac if c.tag == _MC_NS + "Fallback"), None)
        fb_tx = next(iter(fb.iter(qn("w:txbxContent"))), None) if fb is not None else None
        ext = next(iter(choice.iter(_A_EXT)), None)
        out.append({"tx": tx, "fb": fb_tx, "x": _box_x(choice),
                    "cx": int(ext.get("cx") or 0) if ext is not None else 0,
                    "cy": int(ext.get("cy") or 0) if ext is not None else 0,
                    "text": "".join(t.text or "" for t in choice.iter(qn("w:t")))})
    xs = [b["x"] for b in out]
    # 排序前提：全体框都有位置**且同一基准**（评审六轮 F3）——基准不同坐标不可比，保持文档序
    if xs and all(x is not None for x in xs) and len({x[0] for x in xs}) == 1:
        out.sort(key=lambda b: b["x"][1])
    return out


def _boxable_units(new_els: list) -> list:
    """与框位一一配对的单元：图段，以及**取图失败的占位段**——失败也要占住本框位，
    否则反面扫描件会顶进正面的框（评审五轮 C3）。"""
    units = []
    for e in new_els:
        if next(iter(e.iter(qn("w:drawing"))), None) is not None:
            units.append(e)
        elif "".join(t.text or "" for t in e.iter(qn("w:t"))).startswith("（图片加载失败"):
            units.append(e)
    return units


_WPS_NS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}"


def _center_unit(unit) -> None:
    """图段水平居中（2026-08-15 用户实测：图贴左上角不像"贴上去的"）。"""
    ppr = unit.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        unit.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), "center")


def _anchor_middle(tx) -> None:
    """框内容垂直居中：wps:bodyPr anchor="ctr"（只动收图的框；bodyPr 在 wsp 子序列
    里位于 txbx 之后，SubElement 追加即合法）。Fallback 层无对应属性，不动。"""
    txbx = tx.getparent()
    wsp = txbx.getparent() if txbx is not None else None
    if wsp is None or not str(wsp.tag).endswith("}wsp"):
        return
    body = wsp.find(_WPS_NS + "bodyPr")
    if body is None:
        from lxml import etree
        body = etree.SubElement(wsp, _WPS_NS + "bodyPr")
    body.set("anchor", "ctr")


def _fill_boxes(boxes: list[dict], used: set, label: str, units: list) -> list:
    """单元逐一装框 → 已装单元。**本人词匹配的空框优先**，无匹配才取下一空框——
    两组共用一个锚（合并框段落）时占用跟踪+按框上标签配人，绝不把两人的证叠进
    同一框（评审五轮 C1）。Choice/Fallback 双层同装镜像（C2）。"""
    from agent.agents.bidding_agent.nodes.cert_placement import id_person_words

    words = id_person_words(label)
    placed = []
    for unit in units:
        cand = ([i for i, b in enumerate(boxes) if i not in used
                 and any(w in b["text"] for w in words)]
                or [i for i in range(len(boxes)) if i not in used])
        if not cand:
            break
        i = cand[0]
        used.add(i)
        _shrink_drawing(unit, boxes[i]["cx"], boxes[i]["cy"])
        _center_unit(unit)
        # 收到**图**的框清空说明文字（2026-08-14 用户终验：扫描件本就该盖住「粘贴处」
        # 提示，留着更乱）；取图失败的占位段不清——空框失去标签只剩一行报错更糟。
        has_img = next(iter(unit.iter(qn("w:drawing"))), None) is not None
        if has_img:
            _anchor_middle(boxes[i]["tx"])
        for layer in (boxes[i]["tx"], boxes[i]["fb"]):
            if layer is None:
                continue
            if has_img:
                for old in list(layer):
                    layer.remove(old)
            layer.append(copy.deepcopy(unit) if layer is boxes[i]["fb"] else unit)
        placed.append(unit)
    return placed


def _shrink_drawing(img_el, box_cx: int, box_cy: int) -> None:
    """图片段等比适配框内（宽 ≤84%、高 ≤72%，允许放大）。上限从 96/92 收紧
    （2026-08-15 用户实测：92% 高仍溢出框底——a:ext 是外框尺寸，文本框内边距
    （默认左右 0.1"/上下 0.05"）与圆角都要吃掉可用空间）。
    框尺寸缺失/图无尺寸则不动，宁可原样也不写出 0 尺寸的图。"""
    if not box_cx or not box_cy:
        return
    exts = [e for e in img_el.iter() if e.tag in (_WP_EXTENT, _A_EXT) and e.get("cx")]
    if not exts:
        return
    cx, cy = int(exts[0].get("cx") or 0), int(exts[0].get("cy") or 0)
    if not cx or not cy:
        return
    scale = min(box_cx * 0.84 / cx, box_cy * 0.72 / cy)
    ncx, ncy = int(cx * scale), int(cy * scale)
    for e in exts:
        e.set("cx", str(ncx))
        e.set("cy", str(ncy))


def _graft_cert_tail(doc: Document, nodes: list, tail: str, fetch_object) -> None:
    """复印章证照块分组落位：整组（引导行+全部图）一起锚到粘贴框——段落框插正后方、
    表格框进格内;独占锚点时引导行省略（框本身就是标签）,多组共享一个锚点时**保留**
    组名引导行并按尾巴顺序排列（评审四轮 F2:去标+乱序=身份证张冠李戴）;
    锚不上的组保持章末原样。"""
    groups = _parse_cert_groups(tail)
    found = [(label, lead, blocks, _find_cert_anchor(nodes, label))
             for label, lead, blocks in groups]
    per_key: dict = {}
    for *_, anchor in found:
        if anchor:
            per_key[anchor[0]] = per_key.get(anchor[0], 0) + 1
    body = doc.element.body
    cursors: dict = {}
    box_used: dict = {}
    for label, lead, blocks, anchor in found:
        keep_lead = anchor is None or per_key[anchor[0]] > 1
        # 新元素按**位置**圈定,绝不用 id() 集合——lxml 元素是代理对象,同一节点两次迭代
        # 拿到的代理 id 不同,集合差把全书当"新元素"搬走(2026-08-14 首版实测全书错乱)。
        # _emit_html 只在 body 末尾(sectPr 之前)追加,位置切片是精确的。
        n_before = len(body)
        _emit_html(doc, ((lead + "\n") if keep_lead and lead else "") + "\n".join(blocks),
                   fetch_object)
        if anchor is None:
            continue
        kids = list(body)
        end = len(kids) - 1 if kids and kids[-1].tag == qn("w:sectPr") else len(kids)
        start = n_before - 1 if end != len(kids) else n_before
        key, target, mode = anchor
        # 图**进框内**（2026-08-14 深夜实测：图在框后的文档流里会被排到下一页,框空着）：
        # 锚元素里有 wps 文本框时,单元逐一装进各框（占用跟踪+本人词配框,五轮 C1;
        # 表格格里的框同样装,五轮 C4）;框不够装的、引导行/空段,照旧跟在锚点后。
        new_els = kids[start:end]
        boxes = _anchor_boxes(target)
        boxed = _fill_boxes(boxes, box_used.setdefault(key, set()), label,
                            _boxable_units(new_els)) if boxes else []
        rest = [el for el in new_els if not any(el is b for b in boxed)]
        if mode == "cell":
            for el in rest:
                target.append(el)
            continue
        ins = cursors.get(key, target)
        for el in rest:
            ins.addnext(el)
            ins = el
        cursors[key] = ins


def render_docx(outline: dict, chapters: dict, *, meta: dict | None = None,
                 package: dict | None = None,
                 fetch_object: Callable[[str], bytes | None] | None = None,
                 fmt: dict | None = None, scope: str = "full",
                 copier_nodes: dict[str, list] | None = None) -> bytes:
    """完整标书 .docx：封面 + 真目录域页 + 按 outline 顺序各章正文（含「资格证明文件」附录，
    2026-08-09 起前置为生成期系统章节，随 outline/chapters 与其余章节一并渲染，不再是独立
    追加步骤）+ 签章页 + AI 生成提示（spec326 算法备案，恒定追加，见 _add_ai_notice）。
    确定性，无 LLM。package（选包，spec324）存在时封面项目名下加一行包件名。
    fetch_object（2026-08-09 Plan A①）：章节 HTML 里 `<img data-object-key>` 占位图（无 src
    无字节）经它按 key 现取字节落图；取不到/未传/取到坏字节 → 占位行「（图片加载失败：…）」，
    与既有 data: 坏图分支同语义，best-effort，不阻断整本渲染。
    scope（分册，spec 2026-08-08-export-scope）："full"（默认，逐字节兼容旧调用）/"tech"/"business"：
    封面项目名与页脚文档名追加「·技术标部分」/「·商务标部分」，章标题不再带（技术标）/（商务标）
    尾巴——整册同组，逐章带尾巴是噪音。章节过滤是调用方职责，渲染器只管拿到的数据。"""
    meta = _norm_meta(meta or {})
    _SCOPE_SUFFIX = {"tech": "·技术标部分", "business": "·商务标部分"}
    suffix = _SCOPE_SUFFIX.get(scope, "")
    if suffix:
        meta = {**meta, "name": f"{meta.get('name', '投标文件')}{suffix}"}
    doc = Document()
    _apply_bid_styles(doc)
    if fmt is not None:  # spec330 输出格式：显式配置才覆盖,缺省与既有导出一致
        _apply_custom_format(doc, fmt)
    _style_cover(doc, meta, package)
    toc_end_p = _add_toc_field(doc)
    _add_page_number_footer(doc, meta.get("name", "投标文件"))
    # 章节正文：按 outline 顺序（缺正文出占位，不报错）。每章另起一页——评标翻阅按章定位，
    # 章接章挤在同一页找不到边界（用户要求）；首章不加，否则目录后会多出一整页空白。
    for i, ch in enumerate(outline.get("chapters", [])):
        group = "技术标" if ch.get("group") == "tech" else "商务标"
        if i:
            doc.add_page_break()
        tag = f"（{group}）" if scope == "full" else ""
        doc.add_heading(f"{ch.get('no', '')} {ch.get('title', '')}{tag}", level=1)
        # 复印机章（spec 2026-08-14）：招标 docx 原样 XML 节点直接嫁接——版式是复制不是重建，
        # 该章的 HTML 近似版只供编辑器预览，导出以原格式为准（fill_blanks 已在节点上填过空）。
        copied = (copier_nodes or {}).get(ch.get("id", ""))
        if copied is not None:
            from agent.agents.bidding_agent.render.form_copier import graft_nodes
            graft_nodes(doc, copied["nodes"] if isinstance(copied, dict) else copied)
            # 已就位证照图随复印章落位（2026-08-14 授权书实测两阶段）：不落,全书唯一一份
            # 执照/身份证随被替换的 HTML 消失;只缀章末,身份证图离招标「粘贴处」框一页远
            tail = copied.get("tail") if isinstance(copied, dict) else ""
            if tail:
                _graft_cert_tail(doc, copied["nodes"], tail, fetch_object)
            continue
        # 防御清洗：库存章节可能带完整文档壳（<head><style>...），不剥会把样式文本吐进正文；
        # 再与提纲对齐（剥内嵌旧章标题 + 小节编号跟随当前章号）——标书必须按用户设置后的提纲出
        body = strip_document_shell(chapters.get(ch.get("id", ""), ""))
        # 免责语渲染时也清（T6 回放实证）：旧提示词年代生成的存量章节里还带着
        # 「本表格式与招标文件模板可能存在差异」，只在生成/改写时清,存量项目一导出就原样漏出
        body = strip_template_disclaimers(body)
        body = normalize_chapter_html(body, ch.get("no", ""), ch.get("title", ""), ch.get("id", ""))
        if body:
            _emit_html(doc, body, fetch_object)
        else:
            doc.add_paragraph("（本章正文待生成）")
    # 签章页
    doc.add_page_break()
    doc.add_heading("投标人承诺与签章", level=1)
    doc.add_paragraph("法定代表人/授权代表（签字）：____________   日期：__________")
    _add_ai_notice(doc)
    # 正文齐了才回填目录缓存——条目取自成品文档的标题段落，与正文逐字一致
    _fill_toc_entries(doc, toc_end_p)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
