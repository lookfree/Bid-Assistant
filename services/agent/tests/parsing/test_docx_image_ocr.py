"""docx 正文内嵌图片的 OCR（扫描件治理第三步）。

前两步治的是 PDF 扫描页；线下标书更常见的形态是 **.docx 正文里贴着证照图**。
2026-08-10 生产实测（康恒环境那单三份 .docx，内嵌图 156/106/5 张，商务文件正文只有 11950 字）：
营业执照、银行资信证明、审计报告全在图里，解析结果一个字都不留，审查便把**实际提供了的**材料
判成「未提供」的高风险。这一步让内嵌图走与扫描页**同一条**链路：同一个进程级并发闸、
同一条 run 级 deadline、同一套单请求总帽与熔断，识别文字按图片在正文里的位置插回。
"""
import io
import time
import zipfile

import httpx
import pytest

from agent.parsing import ocr as ocr_mod
from agent.parsing.parsers import parse_bytes

_KEY = "uploads/u/x/商务文件-安几科技.docx"
# 内嵌图的来源注记（与扫描页的页首注记同族，统一以「【系统注记」开头——裸标记会被审查模型
# 当成用户文件里的编辑残留报成风险，见 parsing/types.SYSTEM_NOTE_PREFIX）。**带图序号**。
_MARK = "【系统注记·图片识别 第{n}张】"
_ANY_MARK = "【系统注记·图片识别 第"        # 数张数/判有无：不关心是第几张时用它


def _png(w: int = 900, h: int = 700, color: str = "white") -> bytes:
    """一张位图。尺寸/颜色各不相同才会被 python-docx 存成**不同**的媒体部件
    （它按内容哈希去重），矢量化改写那一步要按文件名精确挑一张。"""
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (w, h), color).save(buf, format="PNG")
    return buf.getvalue()


def _docx(*items) -> bytes:
    """按给定顺序拼 docx：str = 段落，bytes = 贴在此处的内嵌图片。"""
    from docx import Document

    d = Document()
    for it in items:
        if isinstance(it, str):
            d.add_paragraph(it)
        else:
            d.add_picture(io.BytesIO(it))
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _as_vector(data: bytes, media: str) -> bytes:
    """把 docx 里指定的一张媒体图换成矢量图（.emf）：Word 里粘贴 Visio/流程图就是这个形态。
    python-docx 造不出来（add_picture 只认位图头），只能在 zip 层改名 + 改关系与内容类型。"""
    src = zipfile.ZipFile(io.BytesIO(data))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as z:
        for item in src.infolist():
            body, name = src.read(item.filename), item.filename
            if name == f"word/media/{media}":
                name = name.replace(".png", ".emf")
            elif name == "word/_rels/document.xml.rels":
                body = body.replace(media.encode(), media.replace(".png", ".emf").encode())
            elif name == "[Content_Types].xml":
                body = body.replace(b'<Default Extension="png"',
                                    b'<Default Extension="emf" ContentType="image/x-emf"/>'
                                    b'<Default Extension="png"')
            z.writestr(name, body)
    return out.getvalue()


@pytest.fixture
def docx_env(monkeypatch, ocr_stub):
    """(桩, 跑一遍解析并识别内嵌图的协程)。桩与 HTTP 打法见 conftest.ocr_stub。"""

    async def run(docx: bytes, key: str = _KEY, **kw):
        def _read(k):
            ocr_stub.reads.append(k)
            return docx

        monkeypatch.setattr(ocr_mod, "read_bytes", _read)
        doc = parse_bytes(docx, key.rsplit("/", 1)[-1])
        return doc, await ocr_mod.ocr_docx_images(doc, key, **kw)

    return ocr_stub, run


async def test_embedded_images_are_recognized_and_spliced_back_where_they_sit(docx_env):
    """三张内嵌位图：每张都送 OCR，识别文字按**图片在正文中的位置**插回并带可读标记，
    全部识别成功后「看不见的内嵌图」归零——审查那条注记随之消失。
    识别出来的字必须重新参与条款切分，否则下游按 clauses 聚章时一个字都进不了审查材料。"""
    stub, run = docx_env
    before, after = await run(_docx("第一章 资格证明", "营业执照扫描件如下：", _png(),
                                    "以上为营业执照。", _png(880, 660, "ivory"),
                                    _png(1000, 800, "azure")))
    assert before.embedded_images == 3          # 第一步的判定：三张图一个字都提不出来
    assert len(stub.requests) == 3
    assert all(r["mode"] == "lines" and r["max_chars"] == 5000 for r in stub.requests)
    assert after.embedded_images == 0

    lines = after.text.split("\n")
    i = lines.index("营业执照扫描件如下：")
    assert lines[i + 1] == _MARK.format(n=1)                                  # 就插在这张图原来的位置
    assert "识别文字1" in lines[i + 2]
    assert lines.index("以上为营业执照。") == i + 3               # 后文原样跟在识别文字之后
    assert any("识别文字1" in c["text"] for c in after.clauses)   # 参与条款切分


async def test_failed_images_stay_counted_as_invisible(docx_env):
    """部分识别失败：成功的拼回正文，失败的**仍计入 embedded_images**——
    审查对那几张照旧说「无法核验」，而不是当成已经看过（诚实注记不缩水）。"""
    stub, run = docx_env
    ok = "识别成功·营业执照正副本·统一社会信用代码 91310000MA1K3XXXXX"   # 够读（≥20 字）
    stub.reply = lambda n, body: (httpx.Response(200, json={"text": ok})
                                  if n == 1 else httpx.Response(500, text="boom"))
    _, after = await run(_docx("正文", _png(), _png(880, 660, "ivory")))
    assert len(stub.requests) == 2
    assert "识别成功" in after.text and after.text.count(_ANY_MARK) == 1
    assert after.embedded_images == 1


def _as_alternate_content(data: bytes) -> bytes:
    """把 docx 里那张图改写成 Word 给「文本框/形状里的图」用的 mc:AlternateContent 写法：
    mc:Choice 里放现代的 w:drawing、mc:Fallback 里放旧的 w:pict，**两者是同一张图**。"""
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from lxml import etree

    from agent.parsing.parsers import _image_rid

    d = Document(io.BytesIO(data))
    drawing = next(el for el in d.element.body.iter() if el.tag == qn("w:drawing"))
    rid, run = _image_rid(drawing), drawing.getparent()
    ns = ('xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
          'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
          'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
          'xmlns:v="urn:schemas-microsoft-com:vml"')
    run.replace(drawing, parse_xml(
        f"<mc:AlternateContent {ns}>"
        f'<mc:Choice Requires="wps">{etree.tostring(drawing).decode()}</mc:Choice>'
        f'<mc:Fallback><w:pict><v:shape><v:imagedata r:id="{rid}"/></v:shape></w:pict>'
        f"</mc:Fallback></mc:AlternateContent>"))
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


async def test_an_image_written_as_alternate_content_counts_once(docx_env):
    """Word 给文本框/形状里的图常写成 mc:Choice(w:drawing) + mc:Fallback(w:pict)——
    **同一张图的两种写法**。两支都数的话张数翻倍、OCR 请求翻倍，同一张证照的识别文字
    还会插进正文两遍（用户看到的是重复的营业执照文字）。"""
    stub, run = docx_env
    before, after = await run(_as_alternate_content(_docx("正文", _png())))
    assert before.embedded_images == 1          # 不是 2
    assert len(stub.requests) == 1
    assert after.text.count(_ANY_MARK) == 1


async def test_the_same_image_is_recognized_once_and_reused(docx_env):
    """同一张图（公章/抬头图/页脚二维码）贴在几十处：只识别一次，结果回填到每一处。

    逐处各发一次请求，就是把 300 张的单文件额度和 20 分钟里的位子让给同一张图——
    真正要紧的证照被挤到帽子外面。python-docx 那边本来就只存一份媒体部件。"""
    stub, run = docx_env
    stamp = _png()
    _before, after = await run(_docx("第一处", stamp, "第二处", stamp, "第三处", stamp))
    assert len(stub.requests) == 1              # 只识别一次
    assert after.text.count(_ANY_MARK) == 3         # 三处都拼回了
    assert after.embedded_images == 0


async def test_a_file_whose_images_are_all_skipped_sends_nothing(docx_env):
    """筛完一张都不够格（全是矢量图/装饰小图）：一次请求都不发、不构造客户端，
    也不打「识别 0/0 张」的假完成日志。张数照第一步口径原样保留。"""
    stub, run = docx_env
    _before, after = await run(_docx("正文", _png(60, 60), _png(80, 40, "ivory")))
    assert stub.requests == [] and stub.clients == 0
    assert after.embedded_images == 2 and _ANY_MARK not in after.text


async def test_recognized_lines_never_become_chapter_titles(docx_env):
    """识别文字是**正文**，一个字都不许参与章节切分的标题判定。

    证照 OCR 出来的行大量长成「1、法定代表人：张三」「一、企业基本情况」——正是启发式回退
    （文档自己一个大纲层级都没有时走的那条路）眼里的章节标题。放行的三重伤害都实测过：
    ① 被判成标题的那一行**被丢出 clauses**，等于识别内容一个字进不了审查材料；
    ② 图片来源注记留在上一节、识别正文被切到下一节，模型看不出这段字来自同一张图；
    ③ 图**后面的原文**被重挂到一个由 OCR 噪声命名的假节下——原文档结构被识别误差改写。"""
    stub, run = docx_env
    lines = ["1、法定代表人：张三", "统一社会信用代码 91310000MA1K3XXXXX", "营业执照正副本齐全有效"]
    stub.reply = lambda n, body: httpx.Response(200, json={"text": "\n".join(lines)})
    # 全是普通段落的文档（一个大纲层级都没有）→ 切分走启发式回退，正是出事的那条路
    _, after = await run(_docx("承诺函", "以下是营业执照：", _png(), "以上为营业执照。"))

    assert after.headings == []                       # 识别文字没有造出任何章节标题
    texts = [c["text"] for c in after.clauses]
    for line in lines:
        assert line in texts, f"识别出来的「{line}」没进 clauses（= 进不了审查材料）"
    secs = {c["id"].rsplit("-c", 1)[0] for c in after.clauses}
    assert secs == {"sec-1"}                          # 图前、识别文字、图后原文仍是同一节


async def test_images_that_only_yield_a_few_characters_are_not_spliced_back(docx_env):
    """只认出几个花纹字的图不算「看见」：既不拼回正文，也不从张数里扣。

    门槛与扫描页共用 `_recognized`（比对基准是空串）。拼回去的话，一份全是图的废件会凭那
    几个字切出条款、聚出非空 chapters，绕过 review「解析不出正文 → run 失败 + 全额退款」的闸
    ——用户为一份什么都看不见的文件付了一次审查的钱。"""
    stub, run = docx_env
    stub.reply = lambda n, body: httpx.Response(200, json={"text": "章 ※ 图"})
    before, after = await run(_docx("正文", _png()))
    assert len(stub.requests) == 1                    # 图确实送去识别了
    assert after.embedded_images == 1                 # 认不出来 → 张数不扣，注记照旧
    assert _ANY_MARK not in after.text and after.text == before.text


async def test_unconfigured_ocr_changes_nothing(monkeypatch, ocr_stub):
    """OCR_BASE_URL 未配置 = 这套环境没部署识别服务：零 HTTP、零取字节，输出与改前逐字节一致。

    绊线走**调用记录**而不是在桩里抛异常：ocr_docx_images 最外层有一道 except 兜底，
    桩抛的 AssertionError 会被它吞掉再 return doc，断言结果和「压根没调用」一模一样。"""
    calls: list[str] = []
    docx = _docx("第一章 资格证明", "营业执照扫描件如下：", _png())
    real_client = httpx.AsyncClient
    monkeypatch.setattr(ocr_mod.settings, "ocr_base_url", None)
    monkeypatch.setattr(ocr_mod, "read_bytes", lambda key: (calls.append("read_bytes"), docx)[1])
    monkeypatch.setattr(ocr_mod.httpx, "AsyncClient",
                        lambda **kw: (calls.append("http"), real_client(**kw))[1])
    doc = parse_bytes(docx, "商务文件.docx")
    after = await ocr_mod.ocr_docx_images(doc, _KEY)
    assert calls == []                      # 判据被挪走 → 这里必红（不再被兜底吞掉）
    assert after is doc and after.embedded_images == 1 and after.text == doc.text


async def test_vector_and_tiny_images_are_skipped_but_still_counted(docx_env):
    """矢量图（emf/wmf，OCR 服务解不出来）与页眉 logo 那种小图一律不发请求：
    156 张图的文件里装饰图占大半，每张都识别就是把 20 分钟的预算烧在没有信息量的图上。
    跳过的图**仍计入 embedded_images**——它们照旧是「看不见」的东西，注记不该因此缩水。"""
    stub, run = docx_env
    docx = _as_vector(_docx("正文", _png(), _png(60, 60, "ivory"), _png(1000, 800, "azure")),
                      "image3.png")          # 第三张（azure）改成矢量图
    before, after = await run(docx)
    assert before.embedded_images == 3
    assert len(stub.requests) == 1           # 只有那张够大的位图被送去识别
    assert after.embedded_images == 2        # 小图 + 矢量图仍算看不见
    assert after.text.count(_ANY_MARK) == 1


async def test_progress_is_reported_over_the_images_actually_sent(docx_env):
    """长识别期间前端横幅不能一动不动：进度按**真正送出去的**张数播报（跳过的图不进分母）。"""
    stub, run = docx_env
    frames: list[tuple[int, int]] = []

    async def _on_progress(done, total):
        frames.append((done, total))

    await run(_docx("正文", _png(), _png(880, 660, "ivory"), _png(60, 60, "azure")),
              on_progress=_on_progress)
    assert len(stub.requests) == 2
    assert frames[-1] == (2, 2)


async def test_exhausted_budget_skips_the_file_without_even_fetching_it(docx_env):
    """预算是**一次审查**的：前面的文件把它用光后，这份连字节都不再取
    （几百 MB 的下载 + 上百张图的解码，白干一遍代价太大），张数照第一步口径原样保留。"""
    stub, run = docx_env
    _, after = await run(_docx("正文", _png()), deadline=time.monotonic() - 1)
    assert stub.requests == [] and stub.reads == []
    assert after.embedded_images == 1


async def test_both_ocr_paths_share_one_run_level_deadline(monkeypatch, ocr_stub):
    """PDF 扫描页与 docx 内嵌图**共享同一条 deadline**，不是各开一份 20 分钟。

    各开一份的话，一次审查最多 10 份文件 = 最坏 200 分钟——用户在一个已预扣积分的步上
    干等几小时，而心跳泵还一直说这个 run 活着。

    判据是**两条识别循环真正收到的是同一个时刻**：只断言 new_deadline 起了一次表是不够的
    ——内嵌图那条路若无视传进去的 deadline 自己 new_deadline()，起表次数照样是 1（反向变异
    M2 实证），护栏看着在、其实不在。"""
    import agent.agents.bidding_agent.nodes.common as common_mod
    from agent.agents.bidding_agent.nodes.common import parse_bid_docs
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()                                    # 一页扫描件（提不出文字）
    scanned_pdf = bytes(pdf.output())
    docx = _docx("第一章 资格证明", _png())
    blobs = {"a.pdf": scanned_pdf, "b.docx": docx}
    starts: list[float] = []

    def _deadline():
        starts.append(time.monotonic())
        return time.monotonic() + ocr_mod._TOTAL_BUDGET_S

    used: list[float] = []
    real_stream = ocr_mod._ocr_stream

    async def _spy(fetch, indices, on_progress, deadline, what):
        used.append(deadline)                     # 这一条识别循环实际吃的是哪条预算
        return await real_stream(fetch, indices, on_progress, deadline, what)

    monkeypatch.setattr(ocr_mod, "_ocr_stream", _spy)
    monkeypatch.setattr(common_mod, "ocr_deadline", _deadline)
    monkeypatch.setattr(common_mod, "read_and_parse",
                        lambda key: parse_bytes(blobs[key], key))
    monkeypatch.setattr(ocr_mod, "read_bytes", lambda key: blobs[key])
    chapters, scanned = await parse_bid_docs(["a.pdf", "b.docx"])
    assert len(ocr_stub.requests) == 2                # 两条链路都真发了请求
    assert len(starts) == 1                           # 预算只起了一次表
    assert len(used) == 2 and used[0] == used[1]      # 两条链路吃的是同一条预算
    # 都识别出来了 → 留下「已识别」统计（供审查提示词解释正文里的识别文字段落，
    # 2026-08-13 实测：统计消失时模型对着识别文字照样判"内容不可见"）
    assert scanned == [{"name": "a.pdf", "pages": 1, "image_pages": 0, "recognized_pages": 1},
                       {"name": "b.docx", "embedded_images": 0, "recognized_images": 1}]
    assert "识别文字" in "".join(chapters.values())    # 识别文字真进了审查材料


async def test_a_legacy_doc_is_converted_once_and_its_images_recognized(monkeypatch, ocr_stub):
    """.doc 内嵌图打通（2026-08-13 实测：《响应文件.doc》11 张证照图只出"内容不可见"注记）：
    识别侧就地转换一次（计入共享预算），图识别、拼回、注记归零——与 .docx 同一口径。"""
    docx = _docx("第一章 资格证明", "执照如下：", _png())
    ole = b"\xd0\xcf\x11\xe0OLE-LEGACY-BYTES"
    seen = {}

    def fake_convert(data, ext):
        seen["ext"], seen["data"] = ext, data
        return docx, "docx"

    monkeypatch.setattr("agent.parsing.parsers._convert_legacy", fake_convert)
    monkeypatch.setattr(ocr_mod, "read_bytes", lambda k: ole)
    doc = parse_bytes(docx, "x.docx")           # .doc 解析产物本就是转换后的 docx 形状
    after = await ocr_mod.ocr_docx_images(doc, "uploads/u/响应文件.doc")
    assert seen == {"ext": "doc", "data": ole}, "识别侧没有对 .doc 做就地转换"
    assert len(ocr_stub.requests) == 1
    assert after.embedded_images == 0, "识别成功后注记没归零"
    assert "识别文字1" in after.text


async def test_a_legacy_doc_that_fails_to_convert_keeps_the_honest_note(monkeypatch, ocr_stub):
    """.doc 转换失败（soffice 缺失/超时）→ 原样返回，注记保持「无法核验」，审查步不炸。"""
    from agent.parsing.types import UnsupportedDocument

    def boom(data, ext):
        raise UnsupportedDocument("soffice 不可用")

    monkeypatch.setattr("agent.parsing.parsers._convert_legacy", boom)
    monkeypatch.setattr(ocr_mod, "read_bytes", lambda k: b"\xd0\xcf\x11\xe0OLE")
    doc = parse_bytes(_docx("第一章", "图：", _png()), "x.docx")
    after = await ocr_mod.ocr_docx_images(doc, "uploads/u/响应文件.doc")
    assert after.embedded_images == doc.embedded_images == 1
    assert not ocr_stub.requests


def test_needs_ocr_counts_a_doc_with_embedded_images(monkeypatch):
    """预算起表口径与发请求同点：.doc 带内嵌图也要起表——否则表起晚了，
    识别做到一半被「预算已用光」拦腰截断。"""
    monkeypatch.setattr(ocr_mod.settings, "ocr_base_url", "http://ocr:8100")
    doc = parse_bytes(_docx("第一章", "图：", _png()), "x.docx")
    assert ocr_mod.needs_ocr(doc, "uploads/u/响应文件.doc")
    assert not ocr_mod.needs_ocr(doc, "uploads/u/响应文件.wps")


async def test_exhausted_budget_after_download_skips_the_doc_conversion(monkeypatch, ocr_stub):
    """下载刚把预算吃光 → 不再跑 60s 转换与整本重解析（评审 2026-08-13：入口检查过了、
    下载耗尽余额、最重的活反而照跑，一张图都发不出去）。"""
    clock = {"now": 0.0}
    monkeypatch.setattr(ocr_mod.time, "monotonic", lambda: clock["now"])

    def slow_read(k):
        clock["now"] = 100.0   # 下载期间预算耗尽
        return b"\xd0\xcf\x11\xe0OLE"

    def must_not_convert(data, ext):
        raise AssertionError("预算已尽仍在跑 .doc 转换")

    monkeypatch.setattr(ocr_mod, "read_bytes", slow_read)
    monkeypatch.setattr("agent.parsing.parsers._convert_legacy", must_not_convert)
    doc = parse_bytes(_docx("第一章", "图：", _png()), "x.docx")
    after = await ocr_mod.ocr_docx_images(doc, "uploads/u/响应文件.doc", deadline=50.0)
    assert after.embedded_images == doc.embedded_images
    assert not ocr_stub.requests
