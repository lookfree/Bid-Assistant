"""docx 章节切分（2026-08-10 生产实测缺陷）。

用户反馈：审查在**文档里明明写着**的内容上报「未响应」（技术偏离表第 3 项「身份集成」、
第 4 项「终端接入」）。取证结论不是截断、不是模型：一份 9 万字的技术文件解析出
**0 条标题、1 个节**，整本连续文本喂给模型。而那份文件在 Word 导航窗格里有清晰的多级标题
（`1.技术偏离表 / 1.1 总体技术规范偏离表 / 1.1.2 核心架构要求偏离表 / 2.项目概况`）——
作者标了，我们没认。这一组用例锁住「先用 Word 自己的结构、再退启发式」的判据。
"""
import io

import pytest
from docx import Document
from docx.oxml import parse_xml

from agent.parsing.parsers import parse_docx

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _bytes(d) -> bytes:
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _outline_style(d, name: str, level: int):
    """自带 w:outlineLvl 的自定义段落样式：真实标书里作者常不用内置「标题 1」，
    而是自建「正文标题一」这类样式并设大纲级别，Word 导航窗格照样成树。"""
    from docx.enum.style import WD_STYLE_TYPE
    st = d.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    st.element.get_or_add_pPr().append(
        parse_xml(f'<w:outlineLvl xmlns:w="{W}" w:val="{level - 1}"/>'))
    return st


def _titles(parsed) -> list[str]:
    return [h["title"] for h in parsed.headings]


def _sec_of(parsed, needle: str) -> str:
    """含某段文字的条款所在节号。"""
    for c in parsed.clauses:
        if needle in c["text"]:
            return c["id"].rsplit("-c", 1)[0]
    raise AssertionError(f"条款里找不到 {needle}")


class TestWordOwnStructure:
    def test_heading_styles_become_sections(self):
        """内置标题样式（Heading 1..9）就是作者标注的层级，必须原样成节。"""
        d = Document()
        d.add_heading("1.技术偏离表", level=1)
        d.add_paragraph("本表逐条应答招标技术要求。")
        d.add_heading("1.1 总体技术规范偏离表", level=2)
        d.add_paragraph("总体规范全部响应。")
        d.add_heading("1.1.2 核心架构要求偏离表", level=3)
        d.add_paragraph("3 身份集成：支持对接统一身份认证平台。")
        d.add_paragraph("4 终端接入：支持 Windows/macOS/信创终端接入。")

        parsed = parse_docx(_bytes(d))
        assert _titles(parsed) == ["1.技术偏离表", "1.1 总体技术规范偏离表", "1.1.2 核心架构要求偏离表"]
        assert [h["sec"] for h in parsed.headings] == ["sec-1", "sec-2", "sec-3"]
        assert [h["level"] for h in parsed.headings] == [1, 2, 3]
        # 用户实例里"找不到"的两项必须落在它们自己那一节里
        assert _sec_of(parsed, "身份集成") == "sec-3"
        assert _sec_of(parsed, "终端接入") == "sec-3"

    def test_custom_style_with_outline_level_is_a_heading(self):
        """自定义样式只要带大纲级别就算标题——真实标书里这比内置样式还常见。"""
        d = Document()
        _outline_style(d, "正文标题一", 1)
        d.add_paragraph("2.项目概况", style="正文标题一")
        d.add_paragraph("项目位于上海。")
        assert _titles(parse_docx(_bytes(d))) == ["2.项目概况"]

    def test_direct_outline_level_on_the_paragraph_wins(self):
        """作者直接在段落上标的大纲级别优先于样式（他就是想让这一段成为标题）。"""
        d = Document()
        p = d.add_paragraph("3.实施方案")
        p._p.get_or_add_pPr().append(parse_xml(f'<w:outlineLvl xmlns:w="{W}" w:val="1"/>'))
        d.add_paragraph("分三阶段实施。")
        parsed = parse_docx(_bytes(d))
        assert _titles(parsed) == ["3.实施方案"] and parsed.headings[0]["level"] == 2


class TestFallbackHeuristics:
    def test_numbered_paragraphs_without_styles_still_split(self):
        """一个大纲级别都没有的文档才走启发式：多级阿拉伯编号要认出来并给对层级。"""
        d = Document()
        for t in ("1.技术偏离表", "本表逐条应答。", "1.1 总体技术规范偏离表", "全部响应。",
                  "第二章 项目概况", "项目位于上海。", "三、商务条款", "付款方式：验收后 30 日。"):
            d.add_paragraph(t)
        parsed = parse_docx(_bytes(d))
        assert _titles(parsed) == ["1.技术偏离表", "1.1 总体技术规范偏离表", "第二章 项目概况", "三、商务条款"]
        assert [h["level"] for h in parsed.headings] == [1, 2, 1, 2]
        assert _sec_of(parsed, "付款方式") == "sec-4"

    def test_a_numeric_line_that_is_not_a_title_is_not_a_heading(self):
        """「100.00 元」这类以数字开头的正文不能成标题——编号后面必须真有标题文字。"""
        d = Document()
        d.add_paragraph("投标总价")
        d.add_paragraph("100.00")
        assert parse_docx(_bytes(d)).headings == []

    def test_long_numbered_sentence_is_body_text(self):
        """编号开头但一长句的是条款正文，不是标题（标题一般很短）。"""
        d = Document()
        d.add_paragraph("1.1 投标人应当在投标截止时间前完成投标文件的上传，逾期上传的投标文件"
                        "采购人不予受理，由此产生的一切后果由投标人自行承担。")
        assert parse_docx(_bytes(d)).headings == []

    def test_a_numbered_list_item_that_ends_in_punctuation_is_not_a_heading(self):
        """真实标书实测：编号列表项短得过得了长度门槛，靠「标题不以句读收尾」把它挡回去。"""
        d = Document()
        d.add_paragraph("1、提供投标须知规定的全部投标文件：正本1份，副本4份。")
        d.add_paragraph("2、保证遵守招标文件中有关规定和收费标准。")
        assert parse_docx(_bytes(d)).headings == []

    def test_bold_cover_page_characters_are_not_headings(self):
        """封面上「报 / 名 / 材 / 料」是四个各自成段的加粗大字。18 份真实标书实测：
        认加粗短行当标题只会凭空造出一堆一个字的节，故加粗一律不作数。"""
        d = Document()
        for ch in "报名材料":
            d.add_paragraph().add_run(ch).bold = True
        d.add_paragraph().add_run("日期：       年      月       日").bold = True
        assert parse_docx(_bytes(d)).headings == []


class TestWhichSignalWins:
    """「作者标了大纲层级」是**覆盖度**判定，不是「有没有」判定。"""

    def test_a_single_styled_cover_title_does_not_veto_the_numbering_heuristic(self):
        """封面用内置「标题 1」、正文章节手打「第一章」是中文标书很常见的排版。

        按「有一段带层级就算标了」判，整本就为了那一行封面放弃编号启发式——实测 11 节塌成
        1 节，正是本次要治的那个生产故障形态换了扇门进来。够不上覆盖度就两种信号并用：
        段落自己标了层级的用它，没标的再猜。"""
        d = Document()
        d.add_heading("投标文件", level=1)          # 封面：唯一一个带样式的标题
        for i in "一二三四五六七八九十":
            d.add_paragraph(f"第{i}章 章节标题")     # 手打编号：Word 眼里就是普通正文
            d.add_paragraph("本章正文若干。")
        parsed = parse_docx(_bytes(d))
        assert _titles(parsed)[0] == "投标文件"     # 封面那条仍按它自己标的层级算
        assert len(parsed.headings) == 11
        assert _sec_of(parsed, "本章正文若干。") == "sec-2"

    def test_a_handful_of_styled_titles_loses_to_a_document_full_of_numbering(self):
        """封面/目录那几行用了样式、正文几十章全是手打编号：几行样式标题压不过整本编号。"""
        d = Document()
        for t in ("投标文件", "目  录", "投标函", "法定代表人授权书"):
            d.add_heading(t, level=1)
        for i in range(20):
            d.add_paragraph(f"第{i + 1}章 手打章节标题")
            d.add_paragraph("本章正文若干。")
        assert len(parse_docx(_bytes(d)).headings) == 24

    def test_words_own_outline_wins_over_numbered_body_lines(self):
        """反过来：作者标够了大纲层级，正文里长得像编号标题的行（目录行、清单项）一律不作数。
        真实语料实测，这些行大半是目录（「一、磋商响应函\t5」）和日期落款，认了就是几十个假节。"""
        d = Document()
        for i in (1, 2, 3):
            d.add_heading(f"{i}.第{i}章", level=1)
            d.add_paragraph("1.1 投标文件组成")      # 启发式眼里的标题，实为目录式正文
            d.add_paragraph("正文若干")
        assert _titles(parse_docx(_bytes(d))) == ["1.第1章", "2.第2章", "3.第3章"]


class TestTablesAndDegradation:
    def test_table_rows_are_never_headings(self):
        """偏离表首行「招标文件的要求」是表头、行里的「1 身份集成」是条目——都不是章节。
        表格里的段落一律不参与标题判定，否则一张偏离表就能切出几十个假节。"""
        d = Document()
        d.add_heading("1.1.2 核心架构要求偏离表", level=3)
        t = d.add_table(rows=3, cols=2)
        t.rows[0].cells[0].text = "招标文件的要求"
        t.rows[0].cells[1].text = "应答情况"
        t.rows[1].cells[0].text = "3 身份集成：对接统一身份认证平台"
        t.rows[1].cells[1].text = "完全响应"
        t.rows[2].cells[0].text = "4 终端接入：支持信创终端"
        t.rows[2].cells[1].text = "完全响应"

        parsed = parse_docx(_bytes(d))
        assert _titles(parsed) == ["1.1.2 核心架构要求偏离表"]
        assert _sec_of(parsed, "身份集成") == "sec-1" and _sec_of(parsed, "终端接入") == "sec-1"

    def test_table_rows_are_never_headings_in_the_fallback_path_either(self):
        """**无样式**文档里的偏离表：表格守卫在启发式那条路上同样要生效。

        上面那条用例用 add_heading 造文档 → 走 Word 大纲层级那条路，而那条路的表格块 level
        天然是 None，`_fallback_level` 里的表格守卫**根本走不到**（反向变异实证：把那道守卫
        删掉，全量用例照样全绿）。真实偏离表首列恰恰是「1.1 xxx」这种编号，放行就是几十个假节。"""
        d = Document()
        d.add_paragraph("投标人对招标技术要求逐条应答如下：")
        t = d.add_table(rows=3, cols=2)
        t.rows[0].cells[0].text = "招标文件的要求"
        t.rows[0].cells[1].text = "应答情况"
        t.rows[1].cells[0].text = "1.1 身份集成对接统一认证"
        t.rows[1].cells[1].text = "完全响应"
        t.rows[2].cells[0].text = "一、终端接入支持信创"
        t.rows[2].cells[1].text = "完全响应"

        parsed = parse_docx(_bytes(d))
        assert parsed.headings == []
        assert _sec_of(parsed, "身份集成") == "sec-1" and _sec_of(parsed, "终端接入") == "sec-1"

    def test_a_chapter_title_in_its_own_bordered_cell_is_still_a_title(self):
        """中文标书很常把章节标题套进一个**带边框的单格表**里排版。

        「表格行一律不判标题」这条一刀切会把这一族文档打回「整本一节」——改前的
        `_split_clauses` 只看文本、根本不知道有表格，反而处理得了。守卫要收窄到
        **多单元格的数据行**（偏离表首列那种编号才是真噪声）。"""
        d = Document()
        for title, body in (("第一章 采购公告", "采购内容：渗透测试服务。"),
                            ("第二章 评审办法", "综合评分法。")):
            box = d.add_table(rows=1, cols=1)
            box.rows[0].cells[0].text = title
            d.add_paragraph(body)
        parsed = parse_docx(_bytes(d))
        assert _titles(parsed) == ["第一章 采购公告", "第二章 评审办法"]
        assert _sec_of(parsed, "综合评分法") == "sec-2"

    def test_amounts_and_dates_are_not_chapter_titles(self):
        """金额与日期不许成节：幻影节会顺移它之后所有条款的 id，而 clause id 是前端定位与
        偏离表引用的键。真实语料实测「2016 年1月20日」被切成过一个节。
        顿号后面接数字的真标题不能误杀（「2、2012年以来…类似项目案例」）。"""
        d = Document()
        for t in ("100.00元", "2026.08 完成", "2016 年1月20日"):
            d.add_paragraph(t)
        assert parse_docx(_bytes(d)).headings == []
        d2 = Document()
        d2.add_paragraph("2、2012年以来合同金额10万元以上类似项目案例")
        d2.add_paragraph("详见业绩表。")
        assert _titles(parse_docx(_bytes(d2))) == ["2、2012年以来合同金额10万元以上类似项目案例"]

    def test_document_without_any_heading_stays_one_section(self):
        """纯扫描件式/无任何结构的文档退回今天的单节行为，逐字节不变。"""
        d = Document()
        for t in ("承诺函", "我方承诺所提供的全部资料真实有效。", "投标人：某某科技有限公司"):
            d.add_paragraph(t)
        parsed = parse_docx(_bytes(d))
        assert parsed.headings == []
        assert [c["id"] for c in parsed.clauses] == ["sec-1-c1", "sec-1-c2", "sec-1-c3"]

    def test_clause_text_keeps_the_existing_stripped_shape(self):
        """条款文本去首尾空白，与既有 _split_clauses 同口径——前端把审查发现定位回原文时
        拿的就是这段文本去比对，留着首行缩进的全角空格会比对不上。"""
        d = Document()
        d.add_paragraph("　　我方承诺所提供的全部资料真实有效。 ")
        assert parse_docx(_bytes(d)).clauses[0]["text"] == "我方承诺所提供的全部资料真实有效。"

    def test_small_document_keeps_its_short_sections(self):
        """小文档里的短节是**真章节**（「第一章 采购公告」下面就一句话），不许并掉。"""
        d = Document()
        for t in ("第一章 采购公告", "采购内容：渗透测试服务。", "第二章 评审办法", "综合评分法。"):
            d.add_paragraph(t)
        parsed = parse_docx(_bytes(d))
        assert len(parsed.headings) == 2
        assert _sec_of(parsed, "综合评分法") == "sec-2"

    def test_tiny_sections_are_merged_when_the_heuristic_shatters_the_document(self):
        """启发式误把成百上千条编号列表项当标题时，过小的节并入前一节——
        把粒度拉回「几十节」的量级，而不是让每条列表项独占一节。

        **上下界都要断言**：只卡上界的话，一路级联把整篇并成 1 节也算通过——而那正是本次
        要治的故障形态（实测过：合并只看待并入的那一节、不看已经攒了多少，400 条列表项
        坍成 1 节）。攒够 `_MIN_SECTION_CHARS` 就该另起一节。"""
        d = Document()
        d.add_paragraph("第一章 技术需求")
        for i in range(400):
            d.add_paragraph(f"{i + 1}.需求项")
            d.add_paragraph("略。")
        parsed = parse_docx(_bytes(d))
        secs = {c["id"].rsplit("-c", 1)[0] for c in parsed.clauses}
        assert 5 <= len(secs) <= 50, f"粒度失控：{len(secs)} 节"
        assert 5 <= len(parsed.headings) <= 50
        # 内容一个字都不能丢：并节只改归属（标题要么留在 headings，要么转成前一节的条款）
        assert (sum("需求项" in c["text"] for c in parsed.clauses)
                + sum("需求项" in h["title"] for h in parsed.headings)) == 400

    def test_the_authors_own_headings_are_never_merged_away(self):
        """作者标了大纲层级的章**任何模式下都不许被并掉**（docx_sections 模块的不变量）。

        实证形态（真实标书常见）：正文 20 章用了 Heading 1，后面跟一大片**没标样式**的附件
        清单（几百条手打编号）。附件那片把大纲覆盖度压到一半以下 → 走两种信号并用的模式 →
        几百个小节 → 触发合并。合并若不认 `styled`，20 个作者标注的章只剩 2 个（实测）：
        把 Word 的大纲认出来了又扔掉，等于换个入口重造「整本一节」。"""
        d = Document()
        for i in range(20):
            d.add_heading(f"第{i + 1}章 章节", level=1)
            d.add_paragraph("本章概述。")
            for j in range(3):
                d.add_paragraph(f"{j + 1}.要点{j + 1}")
                d.add_paragraph("本要点正文若干字。")
        for k in range(200):                       # 作者没标样式的附件清单
            d.add_paragraph(f"{k + 1}.附件项")
            d.add_paragraph("略。")
        parsed = parse_docx(_bytes(d))
        authored = [h for h in parsed.headings if h["title"].startswith("第")]
        assert len(authored) == 20, f"作者标注的章只剩 {len(authored)} 个"
        assert len({c["id"].rsplit("-c", 1)[0] for c in parsed.clauses}) >= 20


class TestDownstreamContracts:
    def test_multi_file_merge_does_not_overwrite_sections(self):
        """多文件合并的全局节号重排语义不能破：第二份的节号整体后移，标题跟着走。"""
        from agent.parsing.merge import merge_parsed

        a = Document()
        a.add_heading("1.技术偏离表", level=1)
        a.add_paragraph("甲文正文。")
        a.add_heading("2.项目概况", level=1)
        a.add_paragraph("甲文概况。")
        b = Document()
        b.add_heading("1.商务条款", level=1)
        b.add_paragraph("乙文正文。")

        doc_a, doc_b = parse_docx(_bytes(a)), parse_docx(_bytes(b))
        clauses, ranges, headings = merge_parsed([("a.docx", doc_a), ("b.docx", doc_b)])
        assert [h["sec"] for h in headings] == ["sec-1", "sec-2", "sec-3"]
        assert [h["title"] for h in headings] == ["1.技术偏离表", "2.项目概况", "1.商务条款"]
        assert [c["id"] for c in clauses if "乙文" in c["text"]] == ["sec-3-c1"]
        assert ranges[1] == {"name": "b.docx", "sec_from": 3, "sec_to": 3}

    def test_a_file_ending_in_a_title_only_section_does_not_collide_with_the_next(self):
        """甲文以「只有标题、没有正文」的节收尾时，乙文的首节不许撞上它的节号。

        偏移量只看 clauses 的最大节号 → 空标题节整个被忽略 → 乙文首节沿用甲文尾节的号，
        `_clause_source` 取第一个匹配的标题，偏离表「出处」列就会印上**另一份文件**的标题；
        file_ranges 也会把甲文的尾节划进乙文的区间。docx 认出大纲层级之后每份文件几百条标题，
        以空标题节收尾（末尾一个「附件清单」标题后面直接结束）是常态。"""
        from agent.agents.bidding_agent.nodes.content import _clause_source
        from agent.parsing.merge import merge_parsed

        a = Document()
        a.add_heading("1.正文章", level=1)
        a.add_paragraph("甲文正文。")
        a.add_heading("2.附件清单", level=1)      # 只有标题、没有正文的尾节
        b = Document()
        b.add_heading("1.乙文首章", level=1)
        b.add_paragraph("乙文正文。")

        docs = [("a.docx", parse_docx(_bytes(a))), ("b.docx", parse_docx(_bytes(b)))]
        clauses, ranges, headings = merge_parsed(docs)
        assert [h["sec"] for h in headings] == ["sec-1", "sec-2", "sec-3"]
        assert [c["id"] for c in clauses if "乙文" in c["text"]] == ["sec-3-c1"]
        assert ranges == [{"name": "a.docx", "sec_from": 1, "sec_to": 2},
                          {"name": "b.docx", "sec_from": 3, "sec_to": 3}]
        # 出处必须是乙文自己的标题，不能是甲文的尾节标题
        assert _clause_source({"doc_headings": headings}, ["sec-3-c1"]) == "1.乙文首章"

    def test_headings_resolve_a_clause_id_back_to_its_title(self):
        """headings 的用处：把内部条款 id 还原成人看得懂的出处（偏离表「出处」列）。"""
        from agent.agents.bidding_agent.nodes.content import _clause_source

        d = Document()
        d.add_heading("1.技术偏离表", level=1)
        d.add_paragraph("略。")
        d.add_heading("1.1.2 核心架构要求偏离表", level=3)
        d.add_paragraph("3 身份集成：对接统一身份认证平台。")
        parsed = parse_docx(_bytes(d))

        read = {"doc_headings": parsed.headings}
        assert _clause_source(read, [_sec_of(parsed, "身份集成") + "-c1"]) == "1.1.2 核心架构要求偏离表"

    def test_sections_split_the_document_into_chapters_downstream(self):
        """_aggregate 按 sec-N 分组聚章：切出节之后审查才是按章体检，而不是整本一章。"""
        from agent.agents.bidding_agent.nodes.common import _aggregate

        d = Document()
        for i in range(5):
            d.add_heading(f"{i + 1}.第{i + 1}章", level=1)
            d.add_paragraph(f"第{i + 1}章正文。")
        out: dict[str, str] = {}
        _aggregate(parse_docx(_bytes(d)), out)
        assert len(out) == 5 and "第3章正文" in out["sec-3"]


@pytest.mark.parametrize("kind", ["pdf", "xlsx"])
def test_other_formats_are_untouched(docgen, kind):
    """pdf/xlsx 走的仍是原来的启发式切分，行为不变。"""
    from agent.parsing.parsers import parse_bytes
    data = docgen.pdf("Tender PDF Body") if kind == "pdf" else docgen.xlsx()
    doc = parse_bytes(data, f"t.{kind}")
    assert doc.clauses and doc.clauses[0]["id"].startswith("sec-1-")
