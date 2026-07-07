"""spec315b parsing 扩展：docx 图片 dHash / 文档属性抽取（hermetic，内存造文件）。"""
import io

from agent.parsing.media import dhash64, extract_doc_meta, extract_media_hashes, hamming


def _gradient_img(invert: bool = False):
    """确定性灰度渐变图：dHash 每一位都有明确的亮暗方向，反转后逐位取反。"""
    from PIL import Image
    img = Image.new("L", (64, 64))
    img.putdata([(255 - x * 4) if invert else x * 4 for _ in range(64) for x in range(64)])
    return img


def _png_bytes(img) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def test_dhash_stable_and_discriminative():
    """同图哈希相同；单调调亮不改变相邻像素大小关系 → 哈希不变；反转渐变 → 全位翻转。"""
    img = _gradient_img()
    h = dhash64(img)
    assert h == dhash64(_gradient_img())
    brighter = img.point(lambda p: min(255, p + 10))
    assert hamming(h, dhash64(brighter)) <= 6          # 近邻判定阈值内
    assert hamming(h, dhash64(_gradient_img(invert=True))) == 64


def test_hamming_counts_bits():
    assert hamming(0b1010, 0b1010) == 0
    assert hamming(0b1111, 0b0000) == 4


def test_extract_media_hashes_docx():
    """docx 内嵌图片被抽出且哈希与原图一致；非 docx 类型返回空。"""
    from docx import Document
    from docx.shared import Inches
    img = _gradient_img()
    d = Document()
    d.add_paragraph("含图片的投标文件")
    d.add_picture(io.BytesIO(_png_bytes(img)), width=Inches(2))
    buf = io.BytesIO()
    d.save(buf)
    hashes = extract_media_hashes(buf.getvalue(), "docx")
    assert hashes == [dhash64(img)]                    # 嵌入字节原样 → 哈希完全一致
    assert extract_media_hashes(b"", "pdf") == []      # pdf 图片抽取 v1 不做


def test_extract_doc_meta_docx():
    from docx import Document
    d = Document()
    d.core_properties.author = "张三"
    d.core_properties.last_modified_by = "李四"
    d.add_paragraph("正文")
    buf = io.BytesIO()
    d.save(buf)
    meta = extract_doc_meta(buf.getvalue(), "docx")
    assert meta["author"] == "张三"
    assert meta["last_modified_by"] == "李四"
    assert meta["company"] is None                     # python-docx 默认模板无 Company


def test_extract_doc_meta_pdf():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_author("Li Si")
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(0, 10, "bid document")
    meta = extract_doc_meta(bytes(pdf.output()), "pdf")
    assert meta["author"] == "Li Si"
    assert meta["last_modified_by"] is None            # pdf 无对应字段


def test_extract_doc_meta_unknown_kind_all_none():
    meta = extract_doc_meta(b"whatever", "xlsx")
    assert meta == {"author": None, "last_modified_by": None, "company": None, "created": None}
