import re
from agent.parsing.parsers import parse_bytes
from agent.parsing.types import UnsupportedDocument
import pytest


def test_parse_docx(docgen):
    doc = parse_bytes(docgen.docx("招标文件正文", "第二段"), "tender.docx")
    assert doc.kind == "docx" and "招标文件正文" in doc.text and "第二段" in doc.text


def test_parse_xlsx_text_and_tables(docgen):
    doc = parse_bytes(docgen.xlsx(), "score.xlsx")
    assert doc.kind == "xlsx" and "技术标" in doc.text
    assert doc.tables and doc.tables[0][0][0] == "评分项"


def test_parse_pdf(docgen):
    doc = parse_bytes(docgen.pdf("Tender PDF Body"), "t.pdf")
    assert doc.kind == "pdf" and doc.pages == 1 and "Tender PDF Body" in doc.text


def test_pdf_counts_pages_without_visible_text():
    """扫描图片页统计（2026-08-09 生产实测：366 页标书 139 页零文字）：提不出文字的页要数出来，
    审查才敢把"看不见"说成"无法核验"，而不是断言"缺少"。
    用空白页模拟——对文本提取而言，扫描图片页与空白页是同一回事（都提不出字）。"""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(0, 10, "Tender body text that is clearly longer than the threshold")
    pdf.add_page()                      # 无文字：等价于扫描图片页
    doc = parse_bytes(bytes(pdf.output()), "t.pdf")
    assert doc.pages == 2 and doc.image_pages == 1


def test_non_pdf_documents_report_no_image_pages(docgen):
    """docx/xlsx 没有"页"的概念 → 恒为 0（默认值兜底，既有解析路径行为不变）。"""
    assert parse_bytes(docgen.docx("招标文件正文"), "t.docx").image_pages == 0
    assert parse_bytes(docgen.xlsx(), "s.xlsx").image_pages == 0


def _pdf_page(text: str, with_image: bool) -> bytes:
    """造一页 PDF：text = 页上的可见文字；with_image = 页上再贴一张位图（模拟证照/盖章扫描图）。
    英文文本只是因为内置字体不支持中文——判定按「可见字数」走，与语种无关。"""
    import io as _io
    from fpdf import FPDF
    from PIL import Image

    pdf = FPDF()
    pdf.set_font("helvetica", size=12)
    pdf.add_page()
    if text:
        pdf.multi_cell(0, 8, text)
    if with_image:
        buf = _io.BytesIO()
        Image.new("RGB", (200, 120), (200, 200, 200)).save(buf, format="PNG")
        buf.seek(0)
        pdf.image(buf, x=10, y=60, w=100)
    return bytes(pdf.output())


# 去掉空白后 30 字：过得了 20 字的纯文字门槛，实质内容却全在图里（「…扫描件如下：」正是 21 字）
_TITLE_ONLY = "Attached authorization scan copy:"
_BODY = "Technical response body sentence. " * 30      # 去掉空白 900 字：正常正文页


def test_one_line_title_plus_a_scan_image_counts_as_an_image_page():
    """混合页：一行标题 + 整版贴图。可见字数 21 过得了 20 字门槛，模型却一个字实质内容都读不到——
    不识别成图片页就既不 OCR、也不进「无法核验」统计，直接退回改这套东西之前的误判。"""
    from agent.parsing.parsers import scanned_page_indices

    doc = parse_bytes(_pdf_page(_TITLE_ONLY, with_image=True), "t.pdf")
    assert doc.image_pages == 1
    assert scanned_page_indices(doc) == [0]         # 且 OCR 真的会去识别它


def test_a_text_page_with_a_logo_is_never_treated_as_a_scan():
    """保守优先：几百字正文 + 页眉 logo 的正常文本页绝不能被误判成扫描页——
    误判的代价是给它发一次 OCR、并在报告里对一页明明看得见的内容说「无法核验」。"""
    from agent.parsing.parsers import scanned_page_indices

    doc = parse_bytes(_pdf_page(_BODY, with_image=True), "t.pdf")
    assert doc.image_pages == 0 and scanned_page_indices(doc) == []


def test_short_page_without_any_image_keeps_the_original_threshold():
    """页里没有图 ⇒ 判据一如既往只看字数：21 字的页不算扫描页（空页/纯页码页才算）。"""
    from agent.parsing.parsers import scanned_page_indices

    doc = parse_bytes(_pdf_page(_TITLE_ONLY, with_image=False), "t.pdf")
    assert doc.image_pages == 0 and scanned_page_indices(doc) == []


def test_broken_page_resources_never_break_parsing():
    """查图片是加严判定的辅助信号：坏 PDF 的 xref 让它抛错时只当「这页没有图」，
    绝不能把整份文件的解析拖垮（解析失败 = 读标/审查当场失败）。"""
    from agent.parsing.parsers import _has_image_xobject

    class _Bad:
        def get(self, *a, **kw):
            raise ValueError("坏 xref")

    assert _has_image_xobject(_Bad()) is False


def test_clauses_have_stable_ids(docgen):
    doc = parse_bytes(docgen.docx("招标文件正文", "第二段"), "tender.docx")
    assert doc.clauses and re.match(r"^sec-.+-c1$", doc.clauses[0]["id"])
    assert doc.clauses[0]["text"] == "招标文件正文"


def _docx_bytes(d) -> bytes:
    import io
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _png_bytes():
    """最小可用位图（python-docx 要读图头算尺寸，随手造的字节串不行）。"""
    import io
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (8, 8), "white").save(buf, format="PNG")
    buf.seek(0)
    return buf


def test_docx_counts_the_images_embedded_in_the_body():
    """docx 正文里贴的证照/盖章扫描图在解析结果里一个字都不留——与扫描 PDF 同一个病，
    只是 docx 连「有多少页看不见」都没有，审查便把印在图上的材料判成「缺少」。
    数出来（w:drawing 现代图 + w:pict 旧 VML 图），审查才有得诚实交代。
    表格单元格里的图同样要数：授权委托书、报价一览表的盖章图几乎都贴在表格里。"""
    from docx import Document
    from docx.oxml import parse_xml
    from agent.parsing.parsers import parse_docx

    d = Document()
    d.add_paragraph("第一章 资格证明")
    d.add_paragraph("营业执照扫描件如下：")
    d.add_picture(_png_bytes())                                  # 段落里的图
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].paragraphs[0].add_run().add_picture(_png_bytes())   # 表格单元格里的图
    run = d.add_paragraph().add_run()                            # 旧 Word 的 VML 图
    run._r.append(parse_xml(
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    parsed = parse_docx(_docx_bytes(d))
    assert parsed.embedded_images == 3
    assert "营业执照扫描件如下：" in parsed.text     # 正文解析一如既往


def test_docx_without_images_reports_zero(docgen):
    """没贴图的 docx 一切照旧：计数为 0 → 审查不加任何注记，提示词逐字节不变。"""
    assert parse_bytes(docgen.docx("招标文件正文", "第二段"), "t.docx").embedded_images == 0


def test_docx_header_logo_is_not_counted_as_an_invisible_material():
    """页眉页脚的公司 logo 每份文件都有，它不是"看不见的材料"。
    只数 body 就天然排除了它们（页眉页脚不在 body 里）——误伤的话每份 docx 都会挂上注记。"""
    from docx import Document
    from agent.parsing.parsers import parse_docx

    d = Document()
    d.add_paragraph("第一章 技术方案")
    d.sections[0].header.paragraphs[0].add_run().add_picture(_png_bytes())
    assert parse_docx(_docx_bytes(d)).embedded_images == 0


def test_unsupported_type_raises():
    with pytest.raises(UnsupportedDocument):
        parse_bytes(b"x", "a.zip")


def test_docx_clauses_include_table_rows_in_document_order():
    """2026-07-22 生产根因回归：招标模板排在表格里，条款分句必须含表格行且按文档顺序归节——
    否则格式章只剩标题占节号（sec 空洞），内容生成拿不到模板原文。"""
    import io
    from docx import Document
    from agent.parsing.parsers import parse_docx

    d = Document()
    d.add_paragraph("第一章 采购公告")
    d.add_paragraph("采购内容：渗透测试服务。")
    d.add_paragraph("第二章 应答文件格式")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "法定代表人授权委托书"
    t.rows[0].cells[1].text = "致：____（采购人名称）"
    t.rows[1].cells[0].text = "应答人："
    t.rows[1].cells[1].text = "____（盖章）"
    d.add_paragraph("第三章 评审办法")
    d.add_paragraph("综合评分法。")
    buf = io.BytesIO()
    d.save(buf)

    parsed = parse_docx(buf.getvalue())
    by_id = {c["id"]: c["text"] for c in parsed.clauses}
    # 表格行成为第二章的条款（\t 连接单元格），且第三章顺延不错位
    sec2 = [t for i, t in by_id.items() if i.startswith("sec-2-")]
    assert any("授权委托书" in t for t in sec2)
    assert any("盖章" in t for t in sec2)
    sec3 = [t for i, t in by_id.items() if i.startswith("sec-3-")]
    assert sec3 == ["综合评分法。"]
