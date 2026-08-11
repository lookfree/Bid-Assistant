"""送审材料里**我们自己**生成的辅助信息，不许被模型当成投标文件的内容。

2026-08-11 生产实测（康恒环境那单重跑审查），报告里冒出这么一条：
    「[中风险] 投标文件多处出现章节编号(如 sec-xxx)和内嵌图片标记，未作清理，影响文件整洁性和专业性」
用户的 .docx 里根本没有这些东西——`sec-N` 是我们的章节键，识别标记是我们拼回识别文字时加的。
既是拿自己的实现细节冤枉用户，也是把内部实现泄露出去；每一份带图/带扫描页的标书都会中招。

两道防线各测一半：
  · 数据侧 —— 注记统一带「【系统注记」前缀（自述身份、且保住"这段字是从哪张图/哪一页认出来的"），
              提纲/构成清单里的内部条款 id 一律不进载荷，产出里谈论注记的发现确定性丢弃；
  · 提示词侧 —— 系统提示讲清这两类东西是什么、以及识别文字为什么可信度更低。
"""
import asyncio

import pytest

from agent.agents.bidding_agent.nodes.common import html_to_review_text, parse_bid_docs
from agent.agents.bidding_agent.schemas import RiskReport
from agent.parsing.types import SYSTEM_NOTE_PREFIX, ParsedDoc
from agent.runtime.registry import RunContext
from agent.agents.bidding_agent.nodes.review import make_review_node

_RISK_ARGS = {
    "score": 78, "high": 0, "mid": 0, "passed": 0,
    "items": [], "passed_items": [],
}

_OCR_TEXT = "营业执照\n统一社会信用代码 91310000MA1FL0XXXX\n安几科技（上海）有限公司"


def _docx_with_two_images() -> bytes:
    """正文里贴着两张证照图的 .docx（康恒环境那三份商务文件的形态）。"""
    import io

    from PIL import Image
    from docx import Document

    def png(w, h, color):
        buf = io.BytesIO()
        Image.new("RGB", (w, h), color).save(buf, format="PNG")
        return buf.getvalue()

    d = Document()
    d.add_paragraph("第二章 资格证明文件")
    d.add_paragraph("投标人须提供营业执照副本复印件并加盖公章。")
    d.add_picture(io.BytesIO(png(900, 700, "white")))
    d.add_picture(io.BytesIO(png(880, 690, "ivory")))
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _recognised_doc() -> ParsedDoc:
    """走**真实**拼回链路（splice_docx_images）造出识别后的解析结果——识别文字前的注记由
    生产代码自己生成，用例不自己伪造，否则测不到注记长什么样。"""
    from agent.parsing.parsers import docx_body_images, parse_bytes, splice_docx_images

    data = _docx_with_two_images()
    blocks, images = docx_body_images(data)
    return splice_docx_images(parse_bytes(data, "商务文件.docx"), blocks, images,
                              {i: _OCR_TEXT for i in range(len(images))})


def _run_review(gw, state) -> tuple[str, str]:
    """跑一遍审查节点 → (系统提示, 用户消息)。"""
    ctx = RunContext(run_id="r", agent_type="bidding_agent", thread_id="t", gateway=gw)
    asyncio.run(make_review_node(ctx)(state))
    msgs = gw.chats[-1].last_messages
    return msgs[0].content, msgs[1].content


class TestPayload:
    async def test_recognised_text_reaches_review_as_a_self_describing_note(self, monkeypatch):
        """从图里认出来的文字进到送审材料时，前面的注记必须自述身份（「【系统注记」）并带来源序号。

        序号不能丢：识别文字的可信度判断（这是 OCR 读出来的，不是投标人写下的承诺）要靠
        "哪张图/哪一页"来落地；从前所有图共用一个「[内嵌图片·识别]」，模型分不清。
        反向变异：把注记改回裸标记（如 `[内嵌图片·识别]`），本用例失败——那正是线上被当成
        「投标文件里的多余标记」报出来的形态。
        """
        import agent.agents.bidding_agent.nodes.common as common_mod

        doc = _recognised_doc()
        monkeypatch.setattr(common_mod, "read_and_parse", lambda key: doc)
        chapters, _scanned = await parse_bid_docs(["uploads/u/x/商务文件.docx"])
        text = html_to_review_text(chapters["sec-1"])
        assert "统一社会信用代码 91310000MA1FL0XXXX" in text     # 识别内容确实进了送审材料
        assert text.count(SYSTEM_NOTE_PREFIX) == 2             # 两张图各自一条，自述"这是系统加的"
        assert "第1张" in text and "第2张" in text              # 归属没丢：哪张图认出来的
        assert "[内嵌图片" not in text and "扫描件识别]" not in text

    def test_outline_and_structure_lose_their_internal_clause_ids(self, submit_gateway):
        """提纲与构成清单挂着的 clause_ids（sec-19-c129…）一个都不许进送审材料。

        审查的产出里没有任何承载 id 的字段，给了只会被抄进报告、或被当成"文件里的多余编号"。
        反向变异：去掉 review.py 里的 strip_clause_ids，本用例失败。
        """
        gw = submit_gateway({"submit_risk_report": _RISK_ARGS})
        _system, user = _run_review(gw, {
            "read": {"risk_summary": [],
                     "required_structure": [{"id": "s1", "title": "投标报价一览表",
                                             "kind": "form", "required": True,
                                             "clause_ids": ["sec-19-c129"]}]},
            "outline": {"chapters": [{"id": "b4", "no": "第四章", "title": "企业资质",
                                      "group": "business", "clause_ids": ["sec-7-c31"]}]},
            "chapters": {"b4": "<p>已通过 ISO9001</p>"},
        })
        assert "投标报价一览表" in user and "企业资质" in user   # 该给的内容一个不少
        assert "clause_ids" not in user
        assert "sec-19-c129" not in user and "sec-7-c31" not in user

    def test_system_prompt_explains_the_notes_and_chapter_keys(self, submit_gateway):
        """提示词侧防线：讲清「【系统注记」与章节键是什么、识别文字为什么可信度更低、
        以及**不得据此判定任何风险**。反向变异：从 REVIEW_SYSTEM_PROMPT 摘掉这一段，本用例失败。"""
        gw = submit_gateway({"submit_risk_report": _RISK_ARGS})
        system, _user = _run_review(gw, {"read": {"risk_summary": []}, "outline": {}, "chapters": {}})
        assert SYSTEM_NOTE_PREFIX in system
        assert "不是投标人写的内容" in system
        assert "可信度低于可复制正文" in system                 # 识别误差要影响可信度判断
        assert "target_id" in system and "不得据此判定任何风险" in system


def _finding(title: str, advice: str = "补齐相应材料", **kw) -> dict:
    return {"level": "中风险", "tone": "warning", "title": title, "advice": advice,
            "chapter_title": "企业资质", "tender_ref": "对应：资格要求",
            "target_tab": "business", "target_id": "b4", "anchor_text": "", **kw}


class TestReportFilter:
    """数据侧的第二道防线：提示词说了不算，产出里谈论系统注记/内部编号的发现确定性丢弃。
    仅靠提示词纪律在弱模型上失效过（本仓 2026-08-01、08-08 两次先例）。"""

    def test_the_production_finding_is_dropped(self):
        """线上那条原文进来 → 丢弃，且 mid 计数跟着重算（不能留个"1 条中风险"的空壳）。

        反向变异：拿掉 _derive_counts 里的 mentions_system_note 判断，本用例失败。
        """
        report = RiskReport(score=72, items=[
            _finding("投标文件多处出现章节编号(如 sec-xxx)和内嵌图片标记，未作清理，影响文件整洁性和专业性"),
            _finding("缺少法定代表人授权委托书", advice="按招标文件格式补齐并盖章"),
        ], passed_items=[])
        assert [i.title for i in report.items] == ["缺少法定代表人授权委托书"]
        assert report.mid == 1

    @pytest.mark.parametrize("title,advice", [
        ("正文中残留【系统注记·图片识别 第3张】等标记", "清理后重新排版"),      # 抄了注记原文
        ("投标文件多处出现 sec-N 形式的章节编号", "清理后重新排版"),            # 编号占位符的另一种写法
        ("建议清理文件", "删除 sec-xxx 一类的编号与图片标记后再递交"),          # 只在建议里提
    ])
    def test_every_shape_of_the_same_complaint_is_dropped(self, title, advice):
        report = RiskReport(score=80, items=[_finding(title, advice)], passed_items=[])
        assert report.items == [] and report.mid == 0

    def test_a_complaint_that_quotes_a_real_clause_id_is_kept(self):
        """**硬约束**：宁可漏过一条格式抱怨，也不许删掉一条真发现。

        「章节 sec-12 的编号未清理」这种、抱怨里带**真** id 的写法，本过滤故意放过——
        分不清它和「按 sec-12 条款补齐材料」这条真发现，而删错的代价高一个数量级。
        真 id 由 clean_internal_ids 抹掉，卡片照常交付（这也是它当初存在的理由）。
        """
        report = RiskReport(score=80, items=[_finding("章节 sec-12 的编号未清理", "删除文中的编号")],
                            passed_items=[])
        assert len(report.items) == 1
        assert "sec-12" not in report.items[0].title      # 编号还是被清掉了，只是卡片留着

    @pytest.mark.parametrize("title,advice", [
        ("无法核验（扫描件）：法定代表人身份证明", "该材料可能已在扫描页中，请人工核对"),
        ("技术方案未说明 IPsec-3DES 加密强度", "补写加密算法与密钥长度"),   # 「sec-」不是我们的编号
        ("响应时间承诺写成 30-sec-timeout，与招标要求单位不符", "改用招标文件的单位重新承诺"),
        ("服务承诺未明确分级 SLA", "按招标要求写明各级响应时限"),
    ])
    def test_real_findings_are_never_touched(self, title, advice):
        """误伤检验：扫描件"无法核验"类结论、以及正文里天然含「sec-」的技术术语都必须留住。"""
        report = RiskReport(score=80, items=[_finding(title, advice)], passed_items=[])
        assert [i.title for i in report.items] == [title]

    # 模型把**真实**条款 id 写进这些字段是提示词点名要求的常态行为（2026-08-08 全量实测
    # 审查报告里 115 处），clean_internal_ids 当初正是为清洗它们而加。判据若把真 id 当成
    # "抱怨证据"，删掉的就是废标级的真发现——比多报一条格式抱怨严重一个数量级。
    @pytest.mark.parametrize("field,value", [
        ("tender_ref", "对应：评审办法（sec-2-c8）（★不可偏离）"),   # 线上实测 115 处的那个形态
        ("advice", "按 sec-8-c95 条款补齐授权委托书并盖章"),
        ("title", "sec-8-c95 要求的授权委托书缺失"),
        ("chapter_title", "sec-1"),                                # 模型把章节键填进章节标题
    ])
    def test_a_real_finding_carrying_a_real_clause_id_is_never_dropped(self, field, value):
        """反向变异：把判据放宽回裸 `sec-`（或把 tender_ref/chapter_title 也拿去当证据），本用例变红。"""
        kw = {field: value}
        report = RiskReport(score=40, items=[
            _finding(kw.pop("title", "未提供投标保证金缴纳凭证（废标项）"),
                     kw.pop("advice", "按招标文件要求缴纳并附银行回单"),
                     level="高风险", tone="destructive", **kw)], passed_items=[])
        assert len(report.items) == 1, "真发现被误杀"
        assert report.high == 1

    def test_the_worst_case_from_review_keeps_its_high_risk_count(self):
        """复审给的最坏形态：一条废标级高风险发现曾被静默删掉，报告变成「0 条高风险、体检分 40」——
        用户会带着一份要废标的标书去投。"""
        report = RiskReport(score=40, items=[
            _finding("未提供投标保证金缴纳凭证（废标项）", "按招标文件要求缴纳并附银行回单",
                     level="高风险", tone="destructive",
                     tender_ref="对应：投标人须知（sec-2-c8）（★不可偏离）")], passed_items=[])
        assert report.high == 1 and len(report.items) == 1

    def test_passed_items_go_through_the_same_filter(self):
        """通过项也可能在说我们自己的注记；真条款 id 同样只清洗、不丢弃。"""
        report = RiskReport(score=90, items=[], passed_items=[
            "文档整洁，无 sec-xxx 一类的残留编号",              # 说的是我们的占位符 → 丢
            "响应函已提供，含90天有效期承诺（sec-8-c10）",      # 真发现带真 id → 留，只清洗
        ], )
        assert len(report.passed_items) == 1 and report.passed == 1
        assert "sec-8-c10" not in report.passed_items[0] and "响应函已提供" in report.passed_items[0]


class TestDedupeKey:
    """审查载荷里已经没有任何内部条款 id（提纲/构成清单/读标三处都剥过），发现本身也就不带
    条款 id → 去重键只能靠位置分辨，否则同类问题会塌缩成一条、漏报其余位置。"""

    def test_same_wording_on_different_chapters_stays_two_findings(self):
        a = _finding("★条款未登入偏离表", "在偏离表中逐条登记", target_id="t1", anchor_text="技术偏离表")
        b = _finding("★条款未登入偏离表", "在偏离表中逐条登记", target_id="b2", anchor_text="商务偏离表")
        report = RiskReport(score=60, items=[a, b], passed_items=[])
        assert len(report.items) == 2 and report.mid == 2

    def test_same_wording_same_chapter_different_spot_stays_two_findings(self):
        a = _finding("★条款未登入偏离表", "逐条登记", target_id="t1", anchor_text="第 3 行 响应时间")
        b = _finding("★条款未登入偏离表", "逐条登记", target_id="t1", anchor_text="第 9 行 可用率")
        report = RiskReport(score=60, items=[a, b], passed_items=[])
        assert len(report.items) == 2

    def test_identical_cards_still_collapse(self):
        """真正的重复（三张一模一样的卡片）照旧塌缩成一条——去重本来要治的就是它。"""
        card = _finding("缺少法定代表人授权委托书", "补齐并盖章", target_id="b4", anchor_text="授权委托书")
        report = RiskReport(score=60, items=[card, dict(card), dict(card)], passed_items=[])
        assert len(report.items) == 1 and report.mid == 1
