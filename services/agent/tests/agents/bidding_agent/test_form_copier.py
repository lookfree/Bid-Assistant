"""表单章复印机 T3（spec 2026-08-14）：招标 docx 的表单节点区间深拷贝进导出文档。
版式（表格合并格/下划线/居中）靠**复制**保真，不是重建——这正是整条复印机路线的意义。"""
import io

import pytest

from agent.agents.bidding_agent.nodes.form_locate import FormSpan


def _tender_docx() -> bytes:
    """真实形态的招标表单页：居中抬头 + 致行 + 带合并格的表格 + 下划线空位落款。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    d.add_paragraph("前置说明段。")                                # body#0
    head = d.add_paragraph("供应商资格信用承诺函")                  # body#1 居中抬头
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("致（采购人）：云上（江西）安全技术有限公司")     # body#2
    t = d.add_table(rows=2, cols=3)                               # body#3
    t.cell(0, 0).merge(t.cell(0, 1))                              # 合并格
    t.cell(0, 0).text = "合计（大写）"
    t.cell(0, 2).text = "税率"
    t.cell(1, 0).text = "1"
    p = d.add_paragraph()                                          # body#4 下划线空位
    p.add_run("单位名称：")
    blank = p.add_run("____________")
    blank.underline = True
    d.add_paragraph("后续无关段。")                                # body#5
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class TestExtract:
    def test_span_nodes_are_copied_with_layout_intact(self):
        """区间抽取的节点保留合并格/居中/下划线——逐属性断言，防「文字在版式丢」。"""
        from agent.agents.bidding_agent.render.form_copier import extract_form_nodes

        nodes = extract_form_nodes(_tender_docx(), FormSpan(1, 4, -1))
        assert len(nodes) == 4
        xml = "".join(__import__("lxml").etree.tostring(n, encoding="unicode") for n in nodes)
        assert "供应商资格信用承诺函" in xml and "云上（江西）安全技术有限公司" in xml
        assert 'w:val="center"' in xml                     # 抬头居中
        assert "gridSpan" in xml or "hMerge" in xml or "vMerge" in xml or "restart" in xml \
            or "w:merge" in xml or "continue" in xml       # 合并格痕迹（python-docx 用 gridSpan/vMerge）
        assert "<w:u " in xml or 'w:u w:val' in xml.replace("'", '"')   # 下划线空位
        assert "后续无关段" not in xml and "前置说明段" not in xml

    def test_out_of_range_span_is_rejected(self):
        from agent.agents.bidding_agent.render.form_copier import (
            CopierUnsupported, extract_form_nodes)
        with pytest.raises(CopierUnsupported):
            extract_form_nodes(_tender_docx(), FormSpan(90, 99, -1))

    def test_numbered_list_in_span_falls_back(self):
        """区间里有编号列表（numPr 引用 numbering.xml）→ 拒收：光搬节点是悬空引用，
        Word 里编号会错乱。宁可回退 HTML 路线。"""
        from docx import Document
        from agent.agents.bidding_agent.render.form_copier import (
            CopierUnsupported, extract_form_nodes)

        from docx.oxml import parse_xml

        d = Document()
        d.add_paragraph("表单头")
        p = d.add_paragraph("第一条")          # 真实 Word 编号列表在 pPr 里直挂 numPr
        p._p.get_or_add_pPr().append(parse_xml(
            '<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'))
        buf = io.BytesIO()
        d.save(buf)
        with pytest.raises(CopierUnsupported):
            extract_form_nodes(buf.getvalue(), FormSpan(0, 1, -1))

    def test_image_in_span_falls_back(self):
        """区间里有图片引用（blip 指媒体部件）→ 拒收，同理悬空。"""
        from docx import Document
        from PIL import Image
        from agent.agents.bidding_agent.render.form_copier import (
            CopierUnsupported, extract_form_nodes)

        d = Document()
        d.add_paragraph("表单头")
        img = io.BytesIO()
        Image.new("RGB", (8, 8), "white").save(img, format="PNG")
        img.seek(0)
        d.add_picture(img)
        buf = io.BytesIO()
        d.save(buf)
        with pytest.raises(CopierUnsupported):
            extract_form_nodes(buf.getvalue(), FormSpan(0, 1, -1))


class TestGraft:
    def test_grafted_nodes_land_before_sectpr_and_keep_text(self):
        """嫁接进输出文档：文本与结构原样、sectPr 仍在文档末尾（插错位置 Word 打不开）。"""
        from docx import Document
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, graft_nodes)

        nodes = extract_form_nodes(_tender_docx(), FormSpan(1, 4, -1))
        out = Document()
        out.add_heading("供应商资格信用承诺函", level=1)
        graft_nodes(out, nodes)
        texts = [p.text for p in out.paragraphs]
        assert "致（采购人）：云上（江西）安全技术有限公司" in texts
        assert len(out.tables) == 1 and out.tables[0].cell(0, 0).text == "合计（大写）"
        body_tags = [el.tag.rsplit("}", 1)[-1] for el in out.element.body.iterchildren()]
        assert body_tags[-1] == "sectPr"


def _fill_docx() -> bytes:
    """填空形态大全（云上承诺函/情况一览表实测形状）：
    分离 run 空位、同 run 标签+空位、带括注标签、一行双空位、表格标签格。"""
    from docx import Document

    d = Document()
    p1 = d.add_paragraph()                                   # body#0 标签与空位分 run
    p1.add_run("单位名称：")
    blank = p1.add_run("________")
    blank.underline = True
    d.add_paragraph("统一社会信用代码（身份证号码）：______")     # body#1 同 run + 括注标签
    p3 = d.add_paragraph()                                   # body#2 一行双空位（后者无值）
    p3.add_run("联系电话：")
    p3.add_run("____")
    p3.add_run("　传真：")
    p3.add_run("____")
    t = d.add_table(rows=2, cols=2)                          # body#3 表格标签格
    t.cell(0, 0).text = "开户银行"
    t.cell(1, 0).text = "神秘字段"
    d.add_paragraph("项目名称：____")                          # body#4 项目信息白名单
    d.add_paragraph("我单位郑重承诺以上内容真实有效。")           # body#5 固定文字
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


_FIELDS = [("单位名称", "上海安几科技有限公司"),
           ("统一社会信用代码", "91310104MA1FRF3K3N"),
           ("联系电话", "021-52808586"),
           ("开户银行", "招商银行股份有限公司上海徐家汇支行")]


class TestFillBlanks:
    def _filled(self):
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks)
        nodes = extract_form_nodes(_fill_docx(), FormSpan(0, 5, -1))
        n = fill_blanks(nodes, _FIELDS, {"name": "云上零信任项目"})
        xml = "".join(__import__("lxml").etree.tostring(x, encoding="unicode") for x in nodes)
        return n, xml

    def test_known_labels_are_filled_and_formats_kept(self):
        """资料库有值的空位填上；分 run 空位保留原 run 格式（下划线还在=字写在横线上）。"""
        n, xml = self._filled()
        assert "上海安几科技有限公司" in xml
        assert "91310104MA1FRF3K3N" in xml                  # 同 run、带括注标签也认
        assert "021-52808586" in xml
        assert "招商银行股份有限公司上海徐家汇支行" in xml     # 表格标签格右侧
        assert "云上零信任项目" in xml                       # 项目信息白名单（meta.name）
        assert 'w:val="single"' in xml or "<w:u " in xml     # 下划线格式没被抹掉
        assert n == 5

    def test_unknown_blanks_stay_blank_and_fixed_text_untouched(self):
        """没值的空位原样留白（传真）；未知标签的表格格不填；固定文字一个字不动——绝不虚构。"""
        n, xml = self._filled()
        assert "传真：" in xml and xml.count("____") >= 1     # 传真的空位还在
        assert "神秘字段" in xml
        assert "我单位郑重承诺以上内容真实有效。" in xml
        import re as _re
        row = _re.search(r"神秘字段.*?</w:tr>", xml, _re.S).group(0)
        assert _re.sub(r"<[^>]+>", "", row).replace("神秘字段", "").strip() == ""

    def test_empty_fields_fill_nothing(self):
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks)
        nodes = extract_form_nodes(_fill_docx(), FormSpan(0, 5, -1))
        assert fill_blanks(nodes, [], {}) == 0


def _copier_tender() -> bytes:
    """带边界的招标 docx：承诺函（裸抬头边界）→ 内容 → 下一份表单抬头收段。"""
    from docx import Document

    d = Document()
    d.add_paragraph("供应商资格信用承诺函")
    d.add_paragraph("致（采购人）：云上（江西）安全技术有限公司")
    d.add_paragraph("单位名称：________")
    d.add_paragraph("我单位自愿参加本次采购询价活动并郑重承诺守信。")
    d.add_paragraph("供应商情况一览表")
    d.add_paragraph("供应商名称：________")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class TestExportWiring:
    """T5 导出接线（评审 2026-08-14 整改后）：pristine 表单章走复印机；手改章/偏离表/方案章/
    非 docx 主文件/自定义导出格式/基线缺失全部让路；招标 key 取 state.files 主文件位。"""

    _CHAPTERS = {"b1": "<p>承诺函生成稿</p>", "b2": "<p>偏离表稿</p>", "t5": "<p>方案稿</p>"}
    _OUTLINE = {"chapters": [
        {"id": "b1", "title": "供应商资格信用承诺函", "group": "business"},
        {"id": "b2", "title": "技术偏离表", "group": "tech"},
        {"id": "t5", "title": "技术方案", "group": "tech"},
    ]}

    class _Ctx:
        thread_id = "proj-x"
        run_id = None          # 事件助手对 None 直接跳过，测试无需 recorder
        recorder = None
        agent_type = "bidding_agent"

    def _state(self, *, main="uploads/u/招标.docx", fmt=None):
        return {"chapters": dict(self._CHAPTERS),
                "files": [{"key": main, "name": main.rsplit("/", 1)[-1]},
                          {"key": "uploads/u/答疑.docx", "name": "答疑.docx"}],
                "read": {"project_meta": {"name": "云上零信任项目"}},
                "run_input": {"library_refs": {"company": [
                    {"body": "单位名称：上海安几科技有限公司"}]},
                    **({"format": fmt} if fmt else {})}}

    def _run(self, monkeypatch, *, original=None, state=None):
        import asyncio
        import agent.agents.bidding_agent.nodes.export as export_mod
        tender = _copier_tender()
        monkeypatch.setattr(export_mod, "_copier_baseline",
                            lambda tid: original if original is not None else dict(self._CHAPTERS))
        monkeypatch.setattr(export_mod.storage_read, "read_bytes", lambda k: tender)
        return asyncio.run(export_mod._copier_nodes(
            self._Ctx(), state or self._state(), self._OUTLINE))

    def test_pristine_form_chapter_is_copied_and_filled(self, monkeypatch):
        out = self._run(monkeypatch)
        assert set(out) == {"b1"}, "只有未手改的表单章走复印机（偏离表/方案章让路）"
        xml = "".join(__import__("lxml").etree.tostring(n, encoding="unicode")
                      for n in out["b1"]["nodes"])
        assert "我单位自愿参加本次采购询价活动并郑重承诺守信。" in xml
        assert "上海安几科技有限公司" in xml            # 企业信息已由代码填进空位
        assert "供应商情况一览表" not in xml            # 下一份表单没被裹进来

    def test_edited_chapter_lets_the_html_route_win(self, monkeypatch):
        """用户手改过（当前 html ≠ 原始产物）→ 复印机让路，绝不静默覆盖手改。"""
        out = self._run(monkeypatch, original={**self._CHAPTERS, "b1": "<p>模型原稿</p>"})
        assert out == {}

    def test_non_docx_main_tender_disables_the_copier(self, monkeypatch):
        """主文件位（files[0]）非 docx 即让路——**不许**退而取任意第一个 docx，
        多文件项目里那可能是答疑册（评审 F13）。"""
        out = self._run(monkeypatch, state=self._state(main="uploads/u/招标.pdf"))
        assert out == {}

    def test_custom_export_format_keeps_the_copier_with_indent_immunity(self, monkeypatch):
        """配置了导出格式（spec330）**不再让路**（2026-08-14 生产实证：这家客户每次导出都带
        格式配置，让路等于复印机永久关闭）。嫁接段落打缩进免疫：无显式缩进的补 firstLine=0，
        防被改过的 Normal 顶成首行缩进——标签列全体右移。"""
        from docx import Document
        from agent.agents.bidding_agent.render.form_copier import graft_nodes

        out = self._run(monkeypatch, state=self._state(fmt={"font": "仿宋"}))
        assert set(out) == {"b1"}
        doc = Document()
        graft_nodes(doc, out["b1"]["nodes"])   # 免疫发生在嫁接时（进入可能被改样式的文档）
        xml = doc.element.body.xml
        assert 'w:firstLine="0"' in xml

    def test_baseline_failure_falls_back_globally(self, monkeypatch):
        import asyncio
        import agent.agents.bidding_agent.nodes.export as export_mod

        def boom(tid):
            raise RuntimeError("pg down")
        monkeypatch.setattr(export_mod, "_copier_baseline", boom)
        monkeypatch.setattr(export_mod.storage_read, "read_bytes",
                            lambda k: (_ for _ in ()).throw(AssertionError("零候选前不许下载")))
        out = asyncio.run(export_mod._copier_nodes(self._Ctx(), self._state(), self._OUTLINE))
        assert out == {}


class TestRenderIntegration:
    def test_render_docx_grafts_copier_chapter_instead_of_html(self):
        """copier_nodes 命中的章：招标原样 XML 进文档、该章 HTML 一个字不渲染；
        章标题照常；未命中章行为与从前逐字节一致（None 传参兼容由既有测试覆盖）。"""
        from docx import Document
        from agent.agents.bidding_agent.render.docx import render_docx
        from agent.agents.bidding_agent.render.form_copier import extract_form_nodes

        nodes = extract_form_nodes(_copier_tender(), FormSpan(0, 3, -1))
        outline = {"chapters": [{"id": "b1", "no": "第一章",
                                 "title": "供应商资格信用承诺函", "group": "business"}]}
        data = render_docx(outline, {"b1": "<p>HTML旧稿不该出现</p>"},
                           copier_nodes={"b1": nodes})
        doc = Document(io.BytesIO(data))
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "我单位自愿参加本次采购询价活动并郑重承诺守信。" in texts
        assert "单位名称：" in texts
        assert "HTML旧稿不该出现" not in texts
        assert "供应商资格信用承诺函" in texts


def _parent_child_tender() -> bytes:
    """父子表单（云上式）：3.一览表 含 3-1.明细表，4. 收段。"""
    from docx import Document

    d = Document()
    d.add_paragraph("3.报价一览表")
    d.add_paragraph("序号\t项目名称\t数量\t单价\t总价")
    d.add_paragraph("合计（大写）：报价一览合计栏")
    d.add_paragraph("3-1.报价明细表")
    d.add_paragraph("序号\t产品名称\t品牌\t明细专属列")
    d.add_paragraph("4.资格文件")
    d.add_paragraph("按要求提供。")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class TestReviewFixes0814:
    def test_edited_child_is_still_cut_out_of_the_copied_parent(self, monkeypatch):
        """评审 F2：去重必须先于 pristine 过滤——明细表被手改（不复印）时，
        被复印的一览表父章里**仍不得**裹着招标的空白明细表，否则成书里同一张表两份打架。"""
        import asyncio
        import agent.agents.bidding_agent.nodes.export as export_mod

        chapters = {"b1": "<p>一览表原稿</p>", "b2": "<p>明细表被手改</p>"}
        outline = {"chapters": [{"id": "b1", "title": "报价一览表", "group": "business"},
                                {"id": "b2", "title": "报价明细表", "group": "business"}]}

        class _Ctx:
            thread_id = "proj-x"
            run_id = None
            recorder = None
            agent_type = "bidding_agent"

        state = {"chapters": chapters,
                 "files": [{"key": "uploads/u/招标.docx", "name": "招标.docx"}],
                 "read": {}, "run_input": {}}
        monkeypatch.setattr(export_mod, "_copier_baseline",
                            lambda tid: {"b1": "<p>一览表原稿</p>", "b2": "<p>明细表原稿</p>"})
        monkeypatch.setattr(export_mod.storage_read, "read_bytes",
                            lambda k: _parent_child_tender())
        out = asyncio.run(export_mod._copier_nodes(_Ctx(), state, outline))
        assert set(out) == {"b1"}
        xml = "".join(__import__("lxml").etree.tostring(n, encoding="unicode")
                      for n in out["b1"]["nodes"])
        assert "报价一览合计栏" in xml
        assert "明细专属列" not in xml, "手改的子表单仍留在被复印的父表里（去重晚于 pristine 过滤）"
        assert "报价明细表" not in xml

    def test_hyperlink_is_refused(self):
        """评审 F4：超链接引用目标文档没有的关系，搬过去轻则修复弹窗重则版面断裂——
        诚实拒收走 HTML 退路。（顶层段落的段内 sectPr 2026-08-14 起改为抽取时剥离，
        见 TestInlineSectPrStripped；藏在表格里的仍拒。）"""
        from docx import Document
        from docx.oxml import parse_xml
        from agent.agents.bidding_agent.render.form_copier import (
            CopierUnsupported, extract_form_nodes)

        W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        d = Document()
        p = d.add_paragraph("表单头")
        p._p.append(parse_xml(
            f'<w:hyperlink {W} xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/'
            f'relationships" r:id="rId9"><w:r><w:t>官网</w:t></w:r></w:hyperlink>'))
        buf = io.BytesIO()
        d.save(buf)
        with pytest.raises(CopierUnsupported, match="超链接"):
            extract_form_nodes(buf.getvalue(), FormSpan(0, 0, -1))

    def test_colon_slot_placeholder_is_filled_but_label_bracket_is_not(self):
        """评审 F5 收窄：「致：【XX公司[采购人名称]】：」冒号后的括注是值槽 → 替换；
        「致（采购人）：」的括注是标签限定语、「____（供应商全称）」是空位说明 → 一个字不动。"""
        from docx import Document
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks)

        d = Document()
        d.add_paragraph("致：【XX公司[采购人名称]】：")
        d.add_paragraph("致（采购人）：____")
        d.add_paragraph("____（供应商全称）法定代表人授权如下")
        buf = io.BytesIO()
        d.save(buf)
        nodes = extract_form_nodes(buf.getvalue(), FormSpan(0, 2, -1))
        fill_blanks(nodes, [("单位名称", "上海安几科技有限公司")], {"buyer": "云上（江西）安全技术有限公司"})
        xml = "".join(__import__("lxml").etree.tostring(n, encoding="unicode") for n in nodes)
        assert "致：云上（江西）安全技术有限公司：" in xml       # 冒号槽已替换
        assert "致（采购人）：" in xml                          # 标签限定语原样
        assert "（供应商全称）" in xml                          # 空位说明原样
        assert "上海安几科技有限公司（供应商全称）" in xml       # 空位本身由下划线填充

    def test_two_blanks_in_one_run_fill_independently(self):
        """评审 F14：「电话：____　传真：____」打在同一个 run 里，两个空位各认各的。"""
        from docx import Document
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks)

        d = Document()
        d.add_paragraph("电话：____　传真：____")
        buf = io.BytesIO()
        d.save(buf)
        nodes = extract_form_nodes(buf.getvalue(), FormSpan(0, 0, -1))
        n = fill_blanks(nodes, [("电话", "021-52808586"), ("传真", "021-99999999")], {})
        xml = "".join(__import__("lxml").etree.tostring(x, encoding="unicode") for x in nodes)
        assert n == 2 and "021-52808586" in xml and "021-99999999" in xml

    def test_placed_cert_images_survive_as_the_copied_chapter_tail(self, monkeypatch):
        """2026-08-14 授权书实测：含 data-file-id 图的表单章照常复印招标版式（粘贴框/签章行
        原样），已就位的证照块抽出来挂章尾——版式与证照两头都保住，一头都不丢。"""
        import asyncio
        import agent.agents.bidding_agent.nodes.export as export_mod

        chapters = {"b1": ('<p>承诺函稿</p><p>【营业执照】见下图：</p>'
                           '<p><img data-file-id="x" data-object-key="k"></p>')}
        outline = {"chapters": [{"id": "b1", "title": "供应商资格信用承诺函", "group": "business"}]}

        class _Ctx:
            thread_id = "proj-x"
            run_id = None
            recorder = None
            agent_type = "bidding_agent"

        state = {"chapters": chapters,
                 "files": [{"key": "uploads/u/招标.docx", "name": "招标.docx"}],
                 "read": {}, "run_input": {}}
        monkeypatch.setattr(export_mod, "_copier_baseline", lambda tid: dict(chapters))
        monkeypatch.setattr(export_mod.storage_read, "read_bytes", lambda k: _copier_tender())
        out = asyncio.run(export_mod._copier_nodes(_Ctx(), state, outline))
        assert set(out) == {"b1"}
        assert 'data-file-id="x"' in out["b1"]["tail"]
        assert "【营业执照】见下图：" in out["b1"]["tail"]


class TestHtmlFill:
    """HTML 版同值填空（2026-08-14 用户口径：审查材料必须与最终交付同值）。
    与 XML 版共用 build_lut 与规则；标签零改动、匹配不上留白。"""

    _FIELDS = [("单位名称", "上海安几科技有限公司"), ("开户银行", "招商银行上海徐家汇支行"),
               ("全权代表姓名", "胡月")]

    def test_paragraph_blank_and_td_pair_and_slot(self):
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        html = ("<p>单位名称：________</p>"
                "<table><tr><td>开户银行</td><td></td><td>银行账号</td><td></td></tr></table>"
                "<p>致：【XX公司[采购人名称]】：</p>")
        out, n = fill_blanks_html(html, self._FIELDS, {"buyer": "云上（江西）安全技术有限公司"})
        assert "单位名称：上海安几科技有限公司" in out
        assert "<td>开户银行</td><td>招商银行上海徐家汇支行</td>" in out
        assert "<td>银行账号</td><td></td>" in out          # 没值的格留空
        assert "致：云上（江西）安全技术有限公司：" in out
        assert n == 3

    def test_trailing_bracket_label_across_tokens(self):
        """授权书形态：「____（全权代表姓名）」空位与后括注被行内标签分开也认。"""
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        html = "<p>授权 <strong>____</strong>（全权代表姓名）为全权代表。</p>"
        out, n = fill_blanks_html(html, self._FIELDS, {})
        assert n == 1 and "胡月" in out
        assert "（全权代表姓名）" in out                     # 括注说明原样保留

    def test_fixed_text_and_tags_untouched(self):
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        html = '<p style="text-align:right">供应商签章：</p><p>我单位郑重承诺。</p>'
        out, n = fill_blanks_html(html, self._FIELDS, {})
        assert out == html and n == 0


class TestHtmlFillHardening:
    """评审 2026-08-14 二轮(7931d1b)硬化钉:转义/裸</值不回扫/宽容格。"""

    def test_values_are_html_escaped(self):
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        out, n = fill_blanks_html("<p>单位名称：____</p>",
                                  [("单位名称", 'AB<证券>公司 & "Co"')], {})
        assert n == 1
        assert "AB&lt;证券&gt;公司 &amp;" in out
        assert "<证券>" not in out

    def test_bare_less_than_survives(self):
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        out, n = fill_blanks_html("<p>a</p>x < y", [("单位名称", "上海安几")], {})
        assert out == "<p>a</p>x < y" and n == 0

    def test_inserted_value_is_not_rescanned(self):
        """值里自带的【】/____ 不许被后续遍改写,填空数不虚增。"""
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        out, n = fill_blanks_html(
            "<table><tr><td>开户银行</td><td></td></tr></table>",
            [("开户银行", "招商银行：【网点[项目名称]】____")], {"name": "XX项目"})
        assert n == 1
        assert "招商银行：【网点[项目名称]】____" in out.replace("&#x27;", "'")

    def test_strong_label_nbsp_cell_and_th_are_filled(self):
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        html = ("<table><tr><th><strong>开户银行</strong></th><td>&nbsp;</td></tr>"
                "<tr><td>银行账号</td><td> </td></tr></table>")
        out, n = fill_blanks_html(html, [("开户银行", "招行徐家汇"), ("银行账号", "121932027710506")], {})
        assert n == 2 and "招行徐家汇" in out and "121932027710506" in out


def _colon_docx() -> bytes:
    """行尾冒号落款（2026-08-14 云上导出实测形态）：招标落款行冒号后**没有下划线**，
    此前 XML 填空无落点（filled=0）——审查页模型填了值，导出却是空标签。"""
    from docx import Document

    d = Document()
    d.add_paragraph("供应商名称：")                       # body#0 别名 → 单位名称
    d.add_paragraph("地址：")                             # body#1 别名 → 注册地址
    d.add_paragraph("联系地址和电话：")                    # body#2 组合值
    d.add_paragraph("电子邮箱：")                          # body#3 无档案 → 留白
    d.add_paragraph("我方承诺如下内容：")                   # body#4 正文引导句 → 绝不填
    p = d.add_paragraph()                                  # body#5 带空位的行归空位规则管
    p.add_run("供应商盖章：")
    p.add_run("________")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


_COLON_FIELDS = [("单位名称", "上海安几科技有限公司"),
                 ("注册地址", "上海市普陀区祁连山路111弄6号"),
                 ("联系电话", "021-52808586")]


class TestLineEndColonFill:
    """2026-08-14 云上导出实测：响应函/承诺函/报价明细表落款 filled=0 的根修。"""

    def _filled(self):
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks)
        nodes = extract_form_nodes(_colon_docx(), FormSpan(0, 5, -1))
        n = fill_blanks(nodes, _COLON_FIELDS, {})
        xml = "".join(__import__("lxml").etree.tostring(x, encoding="unicode") for x in nodes)
        return n, xml

    def test_label_colon_line_gets_value_appended(self):
        n, xml = self._filled()
        assert "供应商名称：" in xml and "上海安几科技有限公司" in xml
        assert "地址：" in xml and "上海市普陀区祁连山路111弄6号" in xml

    def test_combined_address_phone_label(self):
        """「联系地址和电话：」一行要装两个字段，组合成一个值。"""
        _, xml = self._filled()
        assert "上海市普陀区祁连山路111弄6号" in xml and "021-52808586" in xml

    def test_unknown_and_prose_lines_stay_blank(self):
        """无档案的标签留白；正文引导句（「我方承诺如下内容：」）一个字不加。"""
        _, xml = self._filled()
        assert "电子邮箱：</w:t>" in xml or "电子邮箱：" in xml
        import re as _re
        mail = _re.search(r"电子邮箱：([^<]*)", xml)
        assert mail and not mail.group(1).strip()
        prose = _re.search(r"我方承诺如下内容：([^<]*)", xml)
        assert prose is None or not prose.group(1).strip()

    def test_blank_carrying_line_not_double_filled(self):
        """带下划线空位的行归空位规则管——盖章无档案，下划线原样留白。"""
        n, xml = self._filled()
        assert "供应商盖章：" in xml and "________" in xml
        assert n == 3            # 供应商名称 + 地址 + 联系地址和电话，一处不多

    def test_table_label_cell_not_line_end_filled(self):
        """表格标签格「单位名称：」的值走**右侧空格**，标签格自身绝不追加——否则双份。"""
        from docx import Document
        from agent.agents.bidding_agent.render.form_copier import extract_form_nodes, fill_blanks
        d = Document()
        t = d.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "单位名称："
        buf = io.BytesIO()
        d.save(buf)
        nodes = extract_form_nodes(buf.getvalue(), FormSpan(0, 0, -1))
        n = fill_blanks(nodes, _COLON_FIELDS, {})
        xml = "".join(__import__("lxml").etree.tostring(x, encoding="unicode") for x in nodes)
        assert n == 1 and xml.count("上海安几科技有限公司") == 1

    def test_html_line_end_parity(self):
        """HTML 引擎同形态同值（同值三视图：审查/编辑器与导出一个规则）。"""
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        html = "<p>供应商名称：</p><p>电子邮箱：</p><p>我方承诺如下内容：</p>"
        out, n = fill_blanks_html(html, _COLON_FIELDS, {})
        assert n == 1
        assert "<p>供应商名称：上海安几科技有限公司</p>" in out
        assert "<p>电子邮箱：</p>" in out
        assert "<p>我方承诺如下内容：</p>" in out

    def test_html_line_end_value_is_escaped(self):
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        out, n = fill_blanks_html("<p>单位名称：</p>", [("单位名称", "A<B>&C")], {})
        assert n == 1 and "A&lt;B&gt;&amp;C" in out and "<B>" not in out


class TestFallbackImageExempt:
    """2026-08-14 云上 b2 授权书白拒根修：招标身份证粘贴框的 imagedata 全在 mc:Fallback
    （Word 只读 mc:Choice 的兼容降级层）且不引用任何图片文件——搬运无害，不该拒。"""

    _NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
           'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
           'xmlns:v="urn:schemas-microsoft-com:vml" '
           'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
           'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"')

    def _p(self, inner: str):
        from lxml import etree
        return etree.fromstring(f"<w:p {self._NS}><w:r>{inner}</w:r></w:p>")

    def test_fallback_imagedata_without_rel_passes(self):
        from agent.agents.bidding_agent.render.form_copier import _check_portable
        el = self._p("<mc:AlternateContent><mc:Choice Requires=\"wps\"><w:t>框</w:t></mc:Choice>"
                     "<mc:Fallback><w:pict><v:shape><v:imagedata/></v:shape></w:pict>"
                     "</mc:Fallback></mc:AlternateContent>")
        _check_portable(el)      # 不抛 = 通过

    def test_fallback_imagedata_with_rel_still_rejected(self):
        """Fallback 里带 r:id 关系引用的，搬过去是悬空 rId——照旧诚实拒收。"""
        from agent.agents.bidding_agent.render.form_copier import CopierUnsupported, _check_portable
        el = self._p("<mc:AlternateContent><mc:Fallback><w:pict><v:shape>"
                     "<v:imagedata r:id=\"rId9\"/></v:shape></w:pict></mc:Fallback>"
                     "</mc:AlternateContent>")
        with pytest.raises(CopierUnsupported):
            _check_portable(el)

    def test_real_image_outside_fallback_still_rejected(self):
        from agent.agents.bidding_agent.render.form_copier import CopierUnsupported, _check_portable
        el = self._p("<w:drawing><a:blip r:embed=\"rId5\"/></w:drawing>")
        with pytest.raises(CopierUnsupported):
            _check_portable(el)


class TestInlineSectPrStripped:
    """2026-08-14 云上 b2 授权书第二道白拒：表单末行挂着段内分节符（几何与文档级相同，
    只带招标页脚引用）。原样搬=悬空引用+招标页面设置改写全书；整章拒=版式退回 HTML 重建。
    根修：抽取时剥掉顶层段落的 sectPr——内容并入输出文档当前节；藏在表格里的照拒。"""

    _NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
           'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"')

    def _sect_p(self):
        from lxml import etree
        return etree.fromstring(
            f"<w:p {self._NS}><w:pPr><w:sectPr>"
            "<w:footerReference r:id=\"rId4\" w:type=\"default\"/>"
            "<w:pgSz w:w=\"11906\" w:h=\"16838\"/></w:sectPr></w:pPr>"
            "<w:r><w:t>说明：法定代表人参加采购，不用提供授权书</w:t></w:r></w:p>")

    def _doc_sect(self):
        """文档级 sectPr（body 末尾）——剥离只在段内几何与它一致时发生（三轮 F6）。"""
        from lxml import etree
        return etree.fromstring(
            f'<w:sectPr {self._NS}><w:pgSz w:w="11906" w:h="16838"/></w:sectPr>')

    def test_top_level_sectpr_is_stripped_not_rejected(self):
        from agent.agents.bidding_agent.render.form_copier import extract_span
        nodes = extract_span([self._sect_p(), self._doc_sect()], FormSpan(0, 0, -1))
        xml = __import__("lxml").etree.tostring(nodes[0], encoding="unicode")
        assert "说明：法定代表人参加采购" in xml
        assert "sectPr" not in xml and "footerReference" not in xml

    def test_original_node_is_untouched(self):
        """剥离只发生在深拷贝上——招标文档的节点一个字不动。"""
        from agent.agents.bidding_agent.render.form_copier import extract_span
        src = self._sect_p()
        extract_span([src, self._doc_sect()], FormSpan(0, 0, -1))
        xml = __import__("lxml").etree.tostring(src, encoding="unicode")
        assert "sectPr" in xml

    def test_sectpr_hidden_in_table_still_rejected(self):
        from lxml import etree
        from agent.agents.bidding_agent.render.form_copier import CopierUnsupported, extract_span
        tbl = etree.fromstring(
            f"<w:tbl {self._NS}><w:tr><w:tc><w:p><w:pPr><w:sectPr/></w:pPr>"
            "</w:p></w:tc></w:tr></w:tbl>")
        with pytest.raises(CopierUnsupported):
            extract_span([tbl], FormSpan(0, 0, -1))


class TestUnderlinedSpaceBlank:
    """2026-08-14 云上授权书实证（V1 注释预留的形态）：空位是**带下划线格式的纯空格 run**
    （「（供应商全称）法定代表人 ____ 授权 ____（全权代表姓名）为全权代表」）。
    下划线是「在此线上填写」的显式标记——比放宽到任意长空格安全得多；
    值写回原 run（rPr 原样），字落在横线上。"""

    def _auth_docx(self) -> bytes:
        """按云上招标授权正文的真实 run 结构造：缩进空格 run（无下划线）＋
        下划线空格 run×3，标签分别是后括注/前文/后括注。"""
        from docx import Document

        d = Document()
        p = d.add_paragraph()
        p.add_run("    ")                                    # 纯缩进,不是空位
        p.add_run("                 ").underline = True       # 空位1 ← 后括注(供应商全称)
        p.add_run("（供应商全称）法定代表人")
        p.add_run("           ").underline = True             # 空位2 ← 前文 法定代表人
        p.add_run("授权")
        p.add_run("         ").underline = True               # 空位3 ← 后括注(全权代表姓名)
        p.add_run("（全权代表姓名）为全权代表，参加贵处组织的询比活动。")
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    def _filled_xml(self):
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks)
        nodes = extract_form_nodes(self._auth_docx(), FormSpan(0, 0, -1))
        n = fill_blanks(nodes, [("单位名称", "上海安几科技有限公司"),
                                ("法定代表人", "于新宇"),
                                ("全权代表姓名", "胡月")], {})
        return n, "".join(__import__("lxml").etree.tostring(x, encoding="unicode") for x in nodes)

    def test_three_slots_filled_via_bracket_and_leading_labels(self):
        n, xml = self._filled_xml()
        assert n == 3
        assert "上海安几科技有限公司" in xml and "于新宇" in xml and "胡月" in xml

    def test_value_keeps_underline_and_indent_run_untouched(self):
        """值写在原下划线 run 里（字在横线上）；无下划线的缩进空格 run 一个字不动。"""
        n, xml = self._filled_xml()
        assert '<w:t xml:space="preserve">    </w:t>' in xml   # 缩进 run 原样
        assert xml.count('w:val="single"') >= 3 or xml.count("<w:u ") >= 3

    def test_unknown_label_slot_stays_blank(self):
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks)
        from docx import Document
        d = Document()
        p = d.add_paragraph()
        p.add_run("神秘字段")
        p.add_run("        ").underline = True
        buf = io.BytesIO()
        d.save(buf)
        nodes = extract_form_nodes(buf.getvalue(), FormSpan(0, 0, -1))
        n = fill_blanks(nodes, [("单位名称", "上海安几科技有限公司")], {})
        xml = "".join(__import__("lxml").etree.tostring(x, encoding="unicode") for x in nodes)
        assert n == 0 and "上海安几科技有限公司" not in xml


class TestReviewFixesRound3:
    """评审 2026-08-14 三轮(bbc8be3)：行尾冒号/下划线空格/Fallback 豁免各自的越权面收回。"""

    _F = [("单位名称", "上海安几科技有限公司"), ("法定代表人", "于新宇"),
          ("全权代表姓名", "胡月")]
    _NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
           'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
           'xmlns:v="urn:schemas-microsoft-com:vml" '
           'xmlns:o="urn:schemas-microsoft-com:office:office"')

    def test_f1_table_cell_paragraph_not_double_filled(self):
        """F1：<td><p>标签：</p></td> 的值走邻格（_TD_PAIR），段落遍不得再补一份。"""
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        out, n = fill_blanks_html(
            '<table><tr><td><p>单位名称：</p></td><td><p></p></td></tr></table>', self._F, {})
        assert n == 1 and out.count("上海安几科技有限公司") == 1

    def test_f2_html_space_blank_with_bracket_label_fills(self):
        """F2：HTML 侧同样认「长空格串＋括注标签」空位（模板退路的授权书形态），
        审查/导出同值。"""
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        html = ("<p>                 （供应商全称）法定代表人           授权"
                "         （全权代表姓名）为全权代表，参加询比活动。</p>")
        out, n = fill_blanks_html(html, self._F, {})
        assert n == 3
        assert "上海安几科技有限公司" in out and "于新宇" in out and "胡月" in out

    def test_f3_prefixed_label_is_not_filled(self):
        """F3：「分供应商名称：」「外协供应商名称：」不是我方名称槽——行尾/尾括注一律
        精确查表，绝不子串误配。"""
        from agent.agents.bidding_agent.render.form_copier import (
            extract_form_nodes, fill_blanks, fill_blanks_html)
        from docx import Document
        out, n = fill_blanks_html('<p>分供应商名称：</p>', self._F, {})
        assert n == 0 and "上海安几" not in out
        d = Document()
        d.add_paragraph("分供应商名称：")
        buf = io.BytesIO()
        d.save(buf)
        nodes = extract_form_nodes(buf.getvalue(), FormSpan(0, 0, -1))
        assert fill_blanks(nodes, self._F, {}) == 0

    def test_f4_id_reference_inside_fallback_still_rejected(self):
        """F4：Fallback 豁免只豁**无引用的图片壳**；numPr 这类按 ID 引用 numbering.xml
        的节点在降级层里照样悬空（WPS/LibreOffice 会读降级层）——照拒。"""
        from lxml import etree
        from agent.agents.bidding_agent.render.form_copier import CopierUnsupported, _check_portable
        el = etree.fromstring(
            f'<w:p {self._NS}><w:r><mc:AlternateContent><mc:Fallback>'
            '<w:numPr><w:numId w:val="7"/></w:numPr>'
            '</mc:Fallback></mc:AlternateContent></w:r></w:p>')
        with pytest.raises(CopierUnsupported):
            _check_portable(el)

    def test_f5_vml_o_relid_counts_as_relationship_ref(self):
        """F5：VML 的 o:relid 也是关系引用（老 Word 常只写它不写 r:id）——搬过去同样悬空。"""
        from lxml import etree
        from agent.agents.bidding_agent.render.form_copier import CopierUnsupported, _check_portable
        el = etree.fromstring(
            f'<w:p {self._NS}><w:r><mc:AlternateContent><mc:Fallback><w:pict>'
            '<v:shape><v:imagedata o:relid="rId9"/></v:shape>'
            '</w:pict></mc:Fallback></mc:AlternateContent></w:r></w:p>')
        with pytest.raises(CopierUnsupported):
            _check_portable(el)

    def test_f6_landscape_sectpr_is_kept_and_rejected(self):
        """F6：段内 sectPr 只有页面几何与文档级**相同**才剥；横版表单剥了会把宽表塞进
        竖版页——留着走黑名单拒收，HTML 退路重排适配页面，不比从前差。"""
        from lxml import etree
        from agent.agents.bidding_agent.render.form_copier import CopierUnsupported, extract_span
        W = self._NS
        land = etree.fromstring(
            f'<w:p {W}><w:pPr><w:sectPr><w:pgSz w:w="16838" w:h="11906" w:orient="landscape"/>'
            '</w:sectPr></w:pPr><w:r><w:t>横版报价表尾</w:t></w:r></w:p>')
        ref = etree.fromstring(f'<w:sectPr {W}><w:pgSz w:w="11906" w:h="16838"/></w:sectPr>')
        with pytest.raises(CopierUnsupported):
            extract_span([land, ref], FormSpan(0, 0, -1))

    def test_f8_nbsp_padded_colon_line_fills_and_keeps_whitespace(self):
        """F8/F11：冒号与 </p> 之间的 &nbsp;/空白不挡填空，且模板字符（含空白）零丢失。"""
        from agent.agents.bidding_agent.render.form_copier import fill_blanks_html
        out, n = fill_blanks_html('<p>供应商名称：&nbsp; </p>', self._F, {})
        assert n == 1
        assert "供应商名称：上海安几科技有限公司&nbsp; </p>" in out
