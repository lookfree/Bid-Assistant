import asyncio
from agent.runtime.registry import RunContext
from agent.agents.bidding_agent.nodes import common as common_mod
from agent.agents.bidding_agent.nodes import export as export_mod
from agent.agents.bidding_agent.nodes.export import make_export_node


def test_export_node_writes_docx_key(monkeypatch):
    saved = {}

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            saved["key"], saved["len"], saved["ct"] = key, len(data), content_type

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)  # PDF 转换与本测试无关
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-7"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<h3>1.1</h3><p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert out["artifacts"]["docx"] == "artifacts/proj-7/bid.docx"
    assert saved["len"] > 0 and "wordprocessingml" in saved["ct"]


def test_export_node_pdf_conversion_failure_keeps_docx_only(monkeypatch):
    """spec323：docx_to_pdf 返回 None（soffice 缺失/失败）→ pdf/pdf_pages 显式置 None（清 merge 残留,
    评审 F1:否则重导出失败时上一版 PDF/页数会混进新结果），docx 仍产出、不上传 pdf。"""
    saved = {}

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            saved[key] = len(data)

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-9"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert isinstance(out["artifacts"].pop("exported_at"), str)  # 终审 C1：本册渲染时刻标记，逐次不同不断言定值
    assert out["artifacts"] == {"docx": "artifacts/proj-9/bid.docx", "pdf": None, "pdf_pages": None}
    assert "artifacts/proj-9/bid.pdf" not in saved


def test_export_node_pdf_conversion_success_adds_pdf_key(monkeypatch):
    """docx_to_pdf 返回字节 → 上传 bid.pdf，artifacts 携带 pdf key。"""
    saved = {}

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            saved[key] = (len(data), content_type)

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: b"%PDF-1.4 fake")
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-10"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    # 假 PDF 字节解析不出页数 → pdf_pages 显式 None（同样要压过 merge 里的旧值）
    assert isinstance(out["artifacts"].pop("exported_at"), str)  # 终审 C1：本册渲染时刻标记
    assert out["artifacts"] == {"docx": "artifacts/proj-10/bid.docx", "pdf": "artifacts/proj-10/bid.pdf",
                                "pdf_pages": None}
    assert saved["artifacts/proj-10/bid.pdf"] == (13, "application/pdf")


def test_export_node_rerenders_pptx_when_deck_present(monkeypatch):
    """spec315a 契约 5：state 有 deck（含编辑回灌的）→ export 同时重渲 .pptx，docx+pptx 并出。"""
    saved = {}

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            saved[key] = len(data)

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)  # PDF 转换与本测试无关
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-8"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "deck": {"title": "述标", "template": "tech",
                 "slides": [{"id": "s0", "title": "封面", "kind": "cover"}]},
    }))
    assert isinstance(out["artifacts"].pop("exported_at"), str)  # 终审 C1：本册渲染时刻标记
    assert out["artifacts"] == {"docx": "artifacts/proj-8/bid.docx", "pdf": None, "pdf_pages": None,
                                "pptx": "artifacts/proj-8/present.pptx"}
    assert saved["artifacts/proj-8/bid.docx"] > 0 and saved["artifacts/proj-8/present.pptx"] > 0


def test_export_node_adds_package_cover_line_when_run_input_package_present(monkeypatch):
    """spec324：state.run_input.package 存在 → 渲染出的 docx 封面含「包件：《name》」一行。"""
    from docx import Document
    import io as io_mod

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    captured = {}
    real_render_docx = export_mod.render_docx

    def _capturing_render_docx(*args, **kwargs):
        data = real_render_docx(*args, **kwargs)
        captured["data"] = data
        return data
    monkeypatch.setattr(export_mod, "render_docx", _capturing_render_docx)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-11"))
    asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
        "run_input": {"package": {"id": "p1", "name": "实网攻防"}},
    }))
    doc = Document(io_mod.BytesIO(captured["data"]))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "包件：《实网攻防》" in texts


def test_export_node_without_credentials_render_call_identical(monkeypatch):
    """spec325：run_input 无 credentials 键 → render_docx 调用不带 credentials（或为 None），
    产出字节与今天一致（回归：不因新增功能改变现有导出结果）。"""
    from docx import Document
    import io as io_mod

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    captured = {}
    real_render_docx = export_mod.render_docx

    def _capturing_render_docx(*args, **kwargs):
        data = real_render_docx(*args, **kwargs)
        captured["data"] = data
        captured["kwargs"] = kwargs
        return data
    monkeypatch.setattr(export_mod, "render_docx", _capturing_render_docx)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-12"))
    asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert captured["kwargs"].get("credentials") is None
    doc = Document(io_mod.BytesIO(captured["data"]))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "资格证明文件" not in texts


def test_export_node_prefetches_credential_images(monkeypatch):
    """spec325：run_input.credentials 非空 → 节点按 key 预取字节（storage_read.read_bytes），
    渲染出的 docx 含附录标题与图片 media。"""
    from agent.parsing import storage_read as storage_read_mod

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    tiny_png = (b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06"
                b"\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01"
                b"\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82")
    fetched_keys = []

    def _fake_read_bytes(key):
        fetched_keys.append(key)
        return tiny_png

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    monkeypatch.setattr(storage_read_mod, "read_bytes", _fake_read_bytes)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-13"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
        "run_input": {"credentials": [
            {"title": "营业执照", "images": ["library/user1/license.png"]},
        ]},
    }))
    assert out["artifacts"]["docx"] == "artifacts/proj-13/bid.docx"
    assert fetched_keys == ["library/user1/license.png"]


def test_export_node_credential_fetch_failure_no_crash(monkeypatch):
    """spec325：图片 key 取图抛错（MinIO 404/网络）→ 节点不崩，占位段落进入 docx。"""
    from docx import Document
    import io as io_mod
    from agent.parsing import storage_read as storage_read_mod

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    def _raising_read_bytes(key):
        raise RuntimeError("object not found")

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    monkeypatch.setattr(storage_read_mod, "read_bytes", _raising_read_bytes)
    captured = {}
    real_render_docx = export_mod.render_docx

    def _capturing_render_docx(*args, **kwargs):
        data = real_render_docx(*args, **kwargs)
        captured["data"] = data
        return data
    monkeypatch.setattr(export_mod, "render_docx", _capturing_render_docx)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-14"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
        "run_input": {"credentials": [
            {"title": "营业执照", "images": ["library/user1/missing.png"]},
        ]},
    }))
    assert out["artifacts"]["docx"] == "artifacts/proj-14/bid.docx"
    doc = Document(io_mod.BytesIO(captured["data"]))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "（图片加载失败：missing.png）" in texts


def test_export_node_rerender_fetches_master_from_deck_enterprise_template_id(monkeypatch):
    """企业母版：deck.enterprise_template_id 给出 → export 重渲时按它重新预取母版字节，
    传给 render_pptx 的 master_bytes，保持编辑后重导出仍套用同一份企业母版。"""
    from agent.parsing import storage_read as storage_read_mod
    saved = {}

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            saved[key] = len(data)

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    fetched = []
    monkeypatch.setattr(storage_read_mod, "read_bytes",
                        lambda key: fetched.append(key) or b"fake-master-bytes")
    captured = {}

    def _fake_render_pptx(deck, *, template=None, master_bytes=None):
        captured["master_bytes"] = master_bytes
        return b"PK\x03\x04fake"
    monkeypatch.setattr(export_mod, "render_pptx", _fake_render_pptx)

    key = "library/u1/master.pptx"
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-15"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "deck": {"title": "述标", "template": "tech", "enterprise_template_id": key,
                 "slides": [{"id": "s0", "title": "封面", "kind": "cover"}]},
    }))
    assert fetched == [key]
    assert captured["master_bytes"] == b"fake-master-bytes"
    assert out["artifacts"]["pptx"] == "artifacts/proj-15/present.pptx"


class _FakeRecorder:
    """spec326 敏感词扫描测试用 fake recorder：raise_on_log 模拟落库失败。"""

    def __init__(self, raise_on_log: bool = False):
        self.calls: list[dict] = []
        self.raise_on_log = raise_on_log

    def log_event(self, run_id, agent_type, event_type, **kwargs):
        if self.raise_on_log:
            raise RuntimeError("db down")
        self.calls.append({"run_id": run_id, "agent_type": agent_type,
                            "event_type": event_type, **kwargs})


def test_export_node_scans_and_logs_sensitive_word_hit(monkeypatch):
    """spec326：chapters 命中词库违禁词 → recorder.log_event 被调用，event_type=='content_flag'。"""
    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass
    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    recorder = _FakeRecorder()
    node = make_export_node(RunContext(run_id="r1", agent_type="bidding_agent",
                                        thread_id="proj-20", recorder=recorder))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]},
        "chapters": {"t1": "<p>这是赌博网站的广告内容</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert out["artifacts"]["docx"] == "artifacts/proj-20/bid.docx"
    assert len(recorder.calls) == 1
    call = recorder.calls[0]
    assert call["event_type"] == "content_flag"
    assert "赌博" in call["data"]["words"]


def test_export_node_clean_chapters_no_flag_event(monkeypatch):
    """spec326：干净 chapters（无命中）→ recorder.log_event 不被调用。"""
    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass
    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    recorder = _FakeRecorder()
    node = make_export_node(RunContext(run_id="r2", agent_type="bidding_agent",
                                        thread_id="proj-21", recorder=recorder))
    asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]},
        "chapters": {"t1": "<p>技术方案与商务报价正常内容</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert recorder.calls == []


def test_export_node_scan_failure_does_not_block_export(monkeypatch):
    """spec326 生产铁律：recorder.log_event 抛异常 → 扫描只 warning，export 仍成功返回 artifacts。"""
    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass
    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    recorder = _FakeRecorder(raise_on_log=True)
    node = make_export_node(RunContext(run_id="r3", agent_type="bidding_agent",
                                        thread_id="proj-22", recorder=recorder))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]},
        "chapters": {"t1": "<p>这是赌博网站的广告内容</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert out["artifacts"]["docx"] == "artifacts/proj-22/bid.docx"


def test_artifacts_reducer_keeps_pptx_and_docx():
    """spec201 state.artifacts 合并 reducer：present(pptx) 与 export(docx) 并存不互相覆盖。"""
    from agent.agents.bidding_agent.state import _merge_dict
    merged = _merge_dict({"pptx": "artifacts/p/present.pptx"}, {"docx": "artifacts/p/bid.docx"})
    assert merged == {"pptx": "artifacts/p/present.pptx", "docx": "artifacts/p/bid.docx"}


def test_scope_marker_survives_merge_without_being_overwritten_by_other_scope():
    """终审 C1：不同册各自导出的 artifacts 经跨 run 合并 reducer（present/export 共用的同一个
    _merge_dict，thread_id 稳定不变即靠它续状态）叠加后，各自的 exported_at{_sfx} 互不覆盖——
    这是 App 侧 export-preview 判断"某册最近一次真渲染时刻"的前提：docx/docx_tech 等键值一旦
    产出就不再变化（确定性 MinIO key、原地覆盖），单看键是否存在分不出"这行是不是真重渲了那册"，
    只有 exported_at 每次真渲染才刷新，不改动的册原样带旧值——先导全量、后单独导技术册，
    全量那次的 exported_at 必须原样保留（不能被技术册那次运行带偏，那正是 bug 成因）。"""
    from agent.agents.bidding_agent.state import _merge_dict
    full_run = {"docx": "artifacts/p/bid.docx", "pdf": None, "pdf_pages": None,
                "exported_at": "2026-08-01T00:00:00.000+00:00"}
    tech_run = {"docx_tech": "artifacts/p/bid_tech.docx", "pdf_tech": None, "pdf_pages_tech": None,
                "exported_at_tech": "2026-08-09T00:00:00.000+00:00"}
    merged = _merge_dict(full_run, tech_run)
    assert merged["exported_at"] == "2026-08-01T00:00:00.000+00:00"  # 全量渲染时刻未被技术册运行动过
    assert merged["docx"] == "artifacts/p/bid.docx"  # 全量 docx 键值同样原样保留（正是 bug 的成因）
    assert merged["exported_at_tech"] == "2026-08-09T00:00:00.000+00:00"


def test_scope_tech_filters_chapters_and_writes_suffixed_keys(monkeypatch):
    """技术册：只渲 group=tech 章；产物键 docx_tech/pdf_tech/pdf_pages_tech，
    render_docx 收到 scope='tech'。"""
    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: b"%PDF-1.4 fake")
    captured = {}
    real_render_docx = export_mod.render_docx

    def _capturing_render_docx(*args, **kwargs):
        data = real_render_docx(*args, **kwargs)
        captured["outline"] = args[0]
        captured["kwargs"] = kwargs
        return data
    monkeypatch.setattr(export_mod, "render_docx", _capturing_render_docx)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-30"))
    out = asyncio.run(node({
        "outline": {"chapters": [
            {"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"},
            {"id": "b1", "no": "第二章", "title": "商务报价", "group": "business"},
        ]},
        "chapters": {"t1": "<p>技术正文</p>", "b1": "<p>商务正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
        "run_input": {"export_scope": "tech"},
    }))
    assert [c["id"] for c in captured["outline"]["chapters"]] == ["t1"]
    assert captured["kwargs"]["scope"] == "tech"
    assert set(out["artifacts"].keys()) == {"docx_tech", "pdf_tech", "pdf_pages_tech", "exported_at_tech"}
    assert out["artifacts"]["docx_tech"] == "artifacts/proj-30/bid_tech.docx"
    assert out["artifacts"]["pdf_tech"] == "artifacts/proj-30/bid_tech.pdf"
    assert isinstance(out["artifacts"]["exported_at_tech"], str)


def test_scope_business_takes_untagged_chapters(monkeypatch):
    """未标组章节归商务册（与预算口径一致）。run_input export_scope='business'，
    outline 含 group='tech' 与无 group 字段各一章 → 渲染只收无组那章；键带 _biz。"""
    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    captured = {}
    real_render_docx = export_mod.render_docx

    def _capturing_render_docx(*args, **kwargs):
        data = real_render_docx(*args, **kwargs)
        captured["outline"] = args[0]
        captured["kwargs"] = kwargs
        return data
    monkeypatch.setattr(export_mod, "render_docx", _capturing_render_docx)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-31"))
    out = asyncio.run(node({
        "outline": {"chapters": [
            {"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"},
            {"id": "u1", "no": "第二章", "title": "商务条款"},  # 未标组
        ]},
        "chapters": {"t1": "<p>技术正文</p>", "u1": "<p>商务正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
        "run_input": {"export_scope": "business"},
    }))
    assert [c["id"] for c in captured["outline"]["chapters"]] == ["u1"]
    assert captured["kwargs"]["scope"] == "business"
    assert isinstance(out["artifacts"].pop("exported_at_biz"), str)  # 终审 C1：本册渲染时刻标记
    assert out["artifacts"] == {
        "docx_biz": "artifacts/proj-31/bid_biz.docx",
        "pdf_biz": None,
        "pdf_pages_biz": None,
    }


def test_scope_default_full_unchanged(monkeypatch):
    """缺省无 export_scope：行为与今天逐字节一致——键集合 {"docx","pdf","pdf_pages"}，
    render_docx 收到 scope='full'、章节未过滤。既有全量测试同时守护此项。"""
    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    captured = {}
    real_render_docx = export_mod.render_docx

    def _capturing_render_docx(*args, **kwargs):
        data = real_render_docx(*args, **kwargs)
        captured["outline"] = args[0]
        captured["kwargs"] = kwargs
        return data
    monkeypatch.setattr(export_mod, "render_docx", _capturing_render_docx)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-32"))
    out = asyncio.run(node({
        "outline": {"chapters": [
            {"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"},
            {"id": "b1", "no": "第二章", "title": "商务报价", "group": "business"},
        ]},
        "chapters": {"t1": "<p>技术正文</p>", "b1": "<p>商务正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert [c["id"] for c in captured["outline"]["chapters"]] == ["t1", "b1"]
    assert captured["kwargs"]["scope"] == "full"
    assert isinstance(out["artifacts"].pop("exported_at"), str)  # 终审 C1：本册渲染时刻标记
    assert out["artifacts"] == {
        "docx": "artifacts/proj-32/bid.docx",
        "pdf": None,
        "pdf_pages": None,
    }


def test_scope_unknown_value_normalizes_to_full(monkeypatch):
    """终审 M2：export_scope 给了未知字面量（不是 tech/business）→ 按 full 处理：章节不过滤，
    render_docx 收到的 scope 归一成 'full'（不是原始未知值），产物键不带后缀。此前 sfx 归一但
    scope 参未归一，会让全量章节配上空组尾巴（render_docx 的 tag 逻辑判 scope=="full" 才加尾巴）。"""
    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    captured = {}
    real_render_docx = export_mod.render_docx

    def _capturing_render_docx(*args, **kwargs):
        data = real_render_docx(*args, **kwargs)
        captured["outline"] = args[0]
        captured["kwargs"] = kwargs
        return data
    monkeypatch.setattr(export_mod, "render_docx", _capturing_render_docx)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-34"))
    out = asyncio.run(node({
        "outline": {"chapters": [
            {"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"},
            {"id": "b1", "no": "第二章", "title": "商务报价", "group": "business"},
        ]},
        "chapters": {"t1": "<p>技术正文</p>", "b1": "<p>商务正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
        "run_input": {"export_scope": "bogus-scope"},
    }))
    assert [c["id"] for c in captured["outline"]["chapters"]] == ["t1", "b1"]  # 未过滤，与 full 一致
    assert captured["kwargs"]["scope"] == "full"  # 不是原始的 "bogus-scope"
    assert set(out["artifacts"].keys()) == {"docx", "pdf", "pdf_pages", "exported_at"}  # 无分册后缀


def test_scope_with_no_matching_chapters_raises(monkeypatch):
    """全部章节同组时另一册为空 → RuntimeError（防御；前端本已置灰不该到这）。"""
    import pytest

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: None)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-33"))
    with pytest.raises(RuntimeError, match="该册没有章节"):
        asyncio.run(node({
            "outline": {"chapters": [
                {"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"},
            ]},
            "chapters": {"t1": "<p>技术正文</p>"},
            "read": {"project_meta": {"name": "投标文件"}},
            "run_input": {"export_scope": "business"},
        }))


def test_export_node_reports_real_pdf_pages(monkeypatch):
    """真实页数回报（篇幅控制地面真值）：可解析的 PDF → artifacts 带 pdf_pages；
    解析不了（上一用例的假字节）则静默缺省,绝不影响导出。"""
    import io
    from pypdf import PdfWriter
    w = PdfWriter()
    for _ in range(3):
        w.add_blank_page(width=595, height=842)
    buf = io.BytesIO()
    w.write(buf)
    real_pdf = buf.getvalue()

    class _Storage:
        async def put_bytes(self, key, data, content_type=None):
            pass

    monkeypatch.setattr(common_mod, "storage", _Storage())
    monkeypatch.setattr(export_mod, "docx_to_pdf", lambda data: real_pdf)
    node = make_export_node(RunContext(run_id="r", agent_type="bidding_agent", thread_id="proj-11"))
    out = asyncio.run(node({
        "outline": {"chapters": [{"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"}]},
        "chapters": {"t1": "<p>正文</p>"},
        "read": {"project_meta": {"name": "投标文件"}},
    }))
    assert out["artifacts"]["pdf_pages"] == 3
