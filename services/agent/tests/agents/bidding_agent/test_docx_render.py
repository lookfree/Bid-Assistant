import io
import zipfile
from docx import Document
from agent.agents.bidding_agent.render.docx import render_docx

# 1x1 透明 PNG，最小合法图片字节（spec325 测试专用）
_TINY_PNG = (b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06"
             b"\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01"
             b"\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82")


def _doc_text(data: bytes) -> str:
    """docx 字节 → 全部段落文本拼接（该文件既有手法，抽成局部帮助函数供 scope 测试复用）。"""
    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


def test_render_docx_assembles_chapters():
    outline = {"chapters": [
        {"id": "t1", "no": "第一章", "title": "项目理解", "group": "tech"},
        {"id": "b3", "no": "第三章", "title": "商务报价", "group": "business"},
        {"id": "t5", "no": "第五章", "title": "应急预案", "group": "tech"},  # 无正文 → 占位
    ]}
    chapters = {"t1": "<h3>1.1 需求理解</h3><p>政务云运维…</p><ul><li>7×24</li></ul>",
                "b3": "<h3>3.1 报价</h3><p>1560 万元</p><table><tr><th>项</th><th>金额</th></tr><tr><td>运维</td><td>1560</td></tr></table>"}
    data = render_docx(outline, chapters, meta={"name": "某市政务云运维 投标文件", "buyer": "某市大数据局"})
    assert data[:2] == b"PK"
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "某市政务云运维 投标文件" in texts
    assert "（本章正文待生成）" in texts          # t5 无正文 → 占位
    assert "7×24" in texts                         # 列表项进入 docx
    assert doc.tables and doc.tables[0].rows[1].cells[0].text == "运维"   # 表格映射


def test_render_docx_handles_ragged_table_and_container():
    """模型产出不规整时不崩：表格行列参差取最大列数；div 包裹递归展开不压扁。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    chapters = {"t1": "<div><h3>标题</h3><p>正文</p>"
                      "<table><tr><td>a</td></tr><tr><td>b</td><td>c</td></tr></table></div>"}
    data = render_docx(outline, chapters)
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert "标题" in texts and "正文" in texts        # div 内结构保留（各自成段）
    assert doc.tables[0].rows[1].cells[1].text == "c"  # 参差行不越界


def test_render_docx_has_real_toc_and_page_number_fields():
    """spec323：目录不是静态文本而是真域；页脚是居中 PAGE 域页码。
    2026-08-11：域**必须配 settings.xml 的 updateFields**，否则 Word 打开时目录是空的——
    用户看到的就是「导出的文档没有目录」（实测反馈）；原先靠正文写一句「按 F9 更新」提示，
    等于把实现细节转嫁给用户，还把一行废话印进要交给评委的标书。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    data = render_docx(outline, {"t1": "<p>正文</p>"}, meta={"name": "某项目 投标文件"})
    zf = zipfile.ZipFile(io.BytesIO(data))
    document_xml = zf.read("word/document.xml").decode("utf-8")
    footer_xml = "".join(
        zf.read(n).decode("utf-8") for n in zf.namelist() if n.startswith("word/footer")
    )
    header_xml = "".join(
        zf.read(n).decode("utf-8") for n in zf.namelist() if n.startswith("word/header")
    )
    settings_xml = zf.read("word/settings.xml").decode("utf-8")
    assert 'TOC \\o "1-4" \\h \\z \\u' in document_xml     # 真 TOC 域 instrText（1-4:章/节/小节/细项）
    assert "updateFields" in settings_xml                  # 打开即自动更新域（否则目录是空的）
    assert "F9" not in document_xml                        # 不把「请按 F9」这类提示印进标书正文
    assert "PAGE" in footer_xml                            # 页脚 PAGE 域
    assert "某项目 投标文件" in header_xml                  # 页眉=项目名


def test_render_docx_honors_text_align_center_and_right():
    """行内 text-align 必须落到 docx 段落对齐——表单抬头（「响   应   函」）与落款/日期
    靠它才能和招标模板一样居中/靠右。此前渲染层无视对齐，编辑器里居中的文字导出也
    全部变左对齐（2026-08-13 用户实测「响应函少了居中的标题」）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "响应函", "group": "tech"}]}
    html = ('<h3 style="text-align:center">响   应   函</h3>'
            '<p>致：采购人</p>'
            '<p style="text-align: right">法定代表人签字或签章：</p>')
    doc = Document(io.BytesIO(render_docx(outline, {"t1": html})))
    by_text = {p.text: p for p in doc.paragraphs}
    assert by_text["响   应   函"].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert by_text["法定代表人签字或签章："].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert by_text["致：采购人"].alignment is None, "没写对齐的段落不该被动"


def test_render_docx_toc_cache_holds_static_entries():
    """目录域的缓存区必须有静态条目。空域赌「打开时自动更新」——Word 只是弹一次确认，
    WPS 和导出的 PDF 根本不理 updateFields，用户看到的就是空白目录页（2026-08-11 修过
    一次 updateFields，2026-08-12 用户实测「还是没有目录」）。
    条目必须位于 fldChar separate 与 end **之间**：放在域外，Word 更新域后正文会同时
    出现静态一份 + 重建一份，两份目录。"""
    outline = {"chapters": [
        {"id": "t1", "no": "第一章", "title": "响应函", "group": "tech"},
        {"id": "t2", "no": "第二章", "title": "技术方案", "group": "tech"}]}
    data = render_docx(outline, {
        "t1": "<p>正文</p>",
        "t2": "<h3>一、总体架构</h3><p>x</p><h6>（1）五级明细</h6>"})
    xml = zipfile.ZipFile(io.BytesIO(data)).read("word/document.xml").decode("utf-8")
    sep = xml.index('w:fldCharType="separate"')
    end = xml.index('w:fldCharType="end"', sep)
    cache = xml[sep:end]
    assert "第一章 响应函" in cache, "章级条目没进目录缓存——WPS/PDF 里目录仍是空白"
    assert "一、总体架构" in cache                    # 节级（Heading 2）也进
    assert "投标人承诺与签章" in cache                # 系统章同样进目录
    assert "（1）五级明细" not in cache               # 没有页码的深层明细堆一页纯属噪音
    # 目录一份 + 正文一份，恰好两处；缓存条目错插到域外时这里会数出错位
    assert xml.count("第一章 响应函") == 2
    assert "updateFields" in zipfile.ZipFile(io.BytesIO(data)).read("word/settings.xml").decode("utf-8")


def test_render_docx_without_package_byte_identical():
    """spec324：不传 package（未选包/单包）→ 输出与今天逐字节一致（无「包件：」行）。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    meta = {"name": "某项目 投标文件", "buyer": "某局"}
    without_kw = render_docx(outline, {"t1": "<p>正文</p>"}, meta=meta)
    without_default = render_docx(outline, {"t1": "<p>正文</p>"}, meta=meta, package=None)
    assert without_kw == without_default
    doc = Document(io.BytesIO(without_kw))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "包件：" not in texts


def test_render_docx_with_package_adds_cover_line():
    """spec324：package 存在 → 封面项目名下加「包件：《name》」一行。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    meta = {"name": "某项目 投标文件", "buyer": "某局"}
    data = render_docx(outline, {"t1": "<p>正文</p>"}, meta=meta,
                        package={"id": "p1", "name": "实网攻防"})
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert "包件：《实网攻防》" in texts
    assert texts.index("包件：《实网攻防》") < texts.index("采购人：某局")  # 位于项目名之下、其它信息之上


def test_render_docx_placeholder_image_fetches_bytes_via_fetch_object():
    """2026-08-09 资质附录系统章节 Plan A①：附录已前置为普通系统章（sys-creds），章内占位图
    `<img data-object-key>` 无 src 无字节 → fetch_object 按 key 取回 PNG 字节 → docx 含图
    （word/media/ 计数与占位图数一致，沿用该文件既有断言手法）。"""
    outline = {"chapters": [
        {"id": "sys-creds", "no": "附录", "title": "资格证明文件", "group": "business", "system": True},
    ]}
    chapters = {"sys-creds": "<h3>营业执照</h3>"
                             '<p><img data-file-id="f1" data-object-key="library/u1/license.png" alt="营业执照" /></p>'}
    fetched_keys = []

    def _fake_fetch(key):
        fetched_keys.append(key)
        return _TINY_PNG

    data = render_docx(outline, chapters, fetch_object=_fake_fetch)
    assert fetched_keys == ["library/u1/license.png"]
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "资格证明文件" in texts and "营业执照" in texts
    zf = zipfile.ZipFile(io.BytesIO(data))
    media = [n for n in zf.namelist() if n.startswith("word/media/")]
    assert len(media) == 1


def test_render_docx_placeholder_image_fetch_none_falls_back_to_placeholder():
    """fetch_object 回 None（取图失败/对象不存在）→ 占位一行「（图片加载失败：alt）」，不崩，
    与既有 data: 坏图分支同语义。"""
    outline = {"chapters": [
        {"id": "sys-creds", "no": "附录", "title": "资格证明文件", "group": "business", "system": True},
    ]}
    chapters = {"sys-creds": '<p><img data-file-id="f1" data-object-key="library/u1/missing.png" alt="营业执照" /></p>'}
    data = render_docx(outline, chapters, fetch_object=lambda key: None)
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "（图片加载失败：营业执照）" in texts
    zf = zipfile.ZipFile(io.BytesIO(data))
    media = [n for n in zf.namelist() if n.startswith("word/media/")]
    assert len(media) == 0


def test_render_docx_placeholder_image_fetch_object_raises_falls_back_to_placeholder():
    """fetch_object 抛错（MinIO 404/网络抖动，storage_read.read_bytes 不吞异常）→ 同样落占位行，
    不让异常炸穿整本渲染。"""
    outline = {"chapters": [
        {"id": "sys-creds", "no": "附录", "title": "资格证明文件", "group": "business", "system": True},
    ]}
    chapters = {"sys-creds": '<p><img data-file-id="f1" data-object-key="library/u1/missing.png" alt="资质证书" /></p>'}

    def _raising_fetch(key):
        raise RuntimeError("object not found")

    data = render_docx(outline, chapters, fetch_object=_raising_fetch)
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "（图片加载失败：资质证书）" in texts


def test_render_docx_placeholder_image_corrupt_bytes_falls_back_to_placeholder():
    """fetch_object 取到坏字节（add_picture 解码失败）→ 占位行，不崩，不影响导出整体。"""
    outline = {"chapters": [
        {"id": "sys-creds", "no": "附录", "title": "资格证明文件", "group": "business", "system": True},
    ]}
    chapters = {"sys-creds": '<p><img data-file-id="f1" data-object-key="library/u1/bad.png" alt="资质证书" /></p>'}
    data = render_docx(outline, chapters, fetch_object=lambda key: b"not a real image")
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "（图片加载失败：资质证书）" in texts


def test_render_docx_applies_bid_convention_styles():
    """标书排版惯例：正文宋体小四(12pt)、标题黑体加粗黑色（覆盖 Word 默认标题蓝）。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    data = render_docx(outline, {"t1": "<p>正文</p>"})
    doc = Document(io.BytesIO(data))
    normal = doc.styles["Normal"]
    assert normal.font.name == "宋体"
    assert normal.font.size.pt == 12
    assert normal.element.rPr.rFonts.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
    ) == "宋体"
    h1 = doc.styles["Heading 1"]
    assert h1.font.color.rgb == (0, 0, 0)
    assert h1.font.name == "黑体"
    assert h1.font.bold is True
    assert h1.font.size.pt == 16


def test_render_docx_appends_ai_generated_notice():
    """spec326 算法备案：导出文件末尾自动写入 AI 生成提示（长版文案，逐字不可改）。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    data = render_docx(outline, {"t1": "<p>正文</p>"})
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert ("本内容由智启元投标助手生成合成类算法辅助生成，"
            "仅供投标文件编制参考，请结合招标文件原文和企业实际情况复核确认后使用。") in texts


def test_render_docx_embeds_inline_data_url_images():
    """编辑器插图（data URL 内嵌）要落进导出文件：裸 img 与 p 内嵌 img 都算；坏 base64 丢图保文。"""
    import base64
    src = "data:image/png;base64," + base64.b64encode(_TINY_PNG).decode()
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "方案", "group": "tech"}]}
    html = f'<p>拓扑如下：<img src="{src}"/></p><img src="{src}"/><img src="data:image/png;base64,@@bad@@"/>'
    data = render_docx(outline, {"t1": html})
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    assert len(media) >= 1  # 两张同内容图（python-docx 按内容去重）至少落一份；坏图被跳过未炸
    doc = Document(io.BytesIO(data))
    assert any("拓扑如下" in p.text for p in doc.paragraphs)  # 段落文字仍在


def test_render_docx_custom_format_applies(caplog):
    """spec330：传 fmt → A4 + 页边距 2.2/2.3、正文仿宋小四缩进2字符、标题字体字号加粗、1.5倍行距;
    改哪项覆盖哪项（body_font 覆盖为仿宋,其余走默认）。"""
    from docx.shared import Cm, Pt
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "方案", "group": "tech"}]}
    data = render_docx(outline, {"t1": "<p>正文</p>"}, fmt={"body_font": "仿宋"})
    doc = Document(io.BytesIO(data))
    sec = doc.sections[0]
    assert round(sec.page_width.cm, 1) == 21.0 and round(sec.page_height.cm, 1) == 29.7  # A4 纵向
    assert round(sec.top_margin.cm, 1) == 2.2 and round(sec.left_margin.cm, 1) == 2.3
    normal = doc.styles["Normal"]
    assert normal.font.name == "仿宋"                    # 覆盖项生效
    assert normal.font.size == Pt(12)                    # 默认小四
    assert normal.paragraph_format.first_line_indent == Pt(24)  # 首行缩进 2 字符 = 2×12pt
    assert normal.paragraph_format.line_spacing == 1.5
    h2 = doc.styles["Heading 2"]
    assert h2.font.name == "宋体" and h2.font.size == Pt(14) and h2.font.bold  # 标题默认宋体四号加粗
    assert h2.paragraph_format.first_line_indent == Pt(0)


def test_render_docx_without_fmt_keeps_legacy_styles():
    """不传 fmt：与现行样式一致（正文宋体12pt、标题黑体、默认页面）——旧路径零变化。"""
    from docx.shared import Pt
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "方案", "group": "tech"}]}
    doc = Document(io.BytesIO(render_docx(outline, {"t1": "<p>正文</p>"})))
    normal = doc.styles["Normal"]
    assert normal.font.name == "宋体" and normal.font.size == Pt(12)
    assert normal.paragraph_format.first_line_indent is None       # 旧路径不设缩进
    assert doc.styles["Heading 2"].font.name == "黑体"             # 旧路径标题黑体


def test_render_docx_ordered_list_and_merged_cells():
    """审查修正：<ol> 按编号列表输出;colspan/rowspan 按网格展开+合并,列不错位。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "方案", "group": "tech"}]}
    html = ('<ol><li>第一步</li><li>第二步</li></ol>'
            '<table><tr><td colspan="2">合并表头</td><td>C</td></tr>'
            '<tr><td>a</td><td>b</td><td>c</td></tr></table>')
    doc = Document(io.BytesIO(render_docx(outline, {"t1": html})))
    texts = [p.text for p in doc.paragraphs]
    assert "第一步" in texts and "第二步" in texts
    tbl = doc.tables[0]
    assert len(tbl.columns) == 3
    # 合并表头占前两格（同一合并单元格文本一致）,第三列 C 不左移
    assert tbl.cell(0, 0).text == "合并表头" and tbl.cell(0, 2).text == "C"
    assert tbl.cell(1, 0).text == "a" and tbl.cell(1, 2).text == "c"


def test_render_docx_fmt_indent_not_bleeding_into_cells():
    """spec330 审查修正：配置格式时正文缩进不溢入表格单元格。"""
    from docx.shared import Pt
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "方案", "group": "tech"}]}
    html = "<p>正文</p><table><tr><td>格</td></tr></table>"
    doc = Document(io.BytesIO(render_docx(outline, {"t1": html}, fmt={})))
    assert doc.styles["Normal"].paragraph_format.first_line_indent == Pt(24)  # 正文缩进 2 字符
    cell_para = doc.tables[0].cell(0, 0).paragraphs[0]
    assert cell_para.paragraph_format.first_line_indent == Pt(0)              # 单元格显式置零


def test_render_docx_follows_edited_outline_numbering():
    """230 生产实测：正文内嵌生成时的旧章标题（第一章…）与旧层级编号；用户把商务标重排为
    全文连续（第七章）后导出，标题/编号必须按提纲现值出——旧章标题剥掉、小节编号首段跟随。"""
    outline = {"chapters": [
        {"id": "b1", "no": "第七章", "title": "变更申请基本信息", "group": "business"},
    ]}
    chapters = {"b1": "<h1>第一章 变更申请基本信息</h1><h2>1.1 项目名称与申请单号</h2>"
                      "<h3>1.1.1 项目名称</h3><p>正文</p>"}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    texts = [p.text for p in doc.paragraphs]
    assert "第七章 变更申请基本信息" in texts   # 章标题唯一来源=提纲（组尾巴 2026-08-15 起不再显示）
    assert "第一章 变更申请基本信息" not in texts          # 内嵌旧章标题被剥
    assert "7.1 项目名称与申请单号" in texts               # 层级编号跟随当前章号
    assert "7.1.1 项目名称" in texts


def test_render_docx_three_level_heading_hierarchy_and_toc():
    """三级层级贯通（评审:章下标题全平级）:绝对映射 节<h3>→Word2、小节<h4>→Word3（评审二轮:
    相对归一会让同一 <h4> 在不同章落不同级,改绝对映射保证跨章一致、编辑器可预期）;
    旧文档全 <h3> 同样得到 章(1)→节(2);TOC 域扩到 1-4;Heading 4 样式已配置(防主题蓝)。"""
    outline = {"chapters": [
        {"id": "t1", "no": "第一章", "title": "技术方案", "group": "tech"},
        {"id": "t2", "no": "第二章", "title": "实施方案", "group": "tech"},
        {"id": "t3", "no": "第三章", "title": "培训方案", "group": "tech"},
    ]}
    chapters = {
        "t1": "<h3>1.1 需求理解</h3><p>正文</p><h4>1.1.1 现状分析</h4><p>细项</p><h3>1.2 总体设计</h3>",
        "t2": "<h3>2.1 计划</h3><p>旧式单级正文</p><h3>2.2 保障</h3>",
        "t3": "<h4>3.0.1 只有小节的章</h4><p>跨章一致:h4 恒为第三级,不随本章标签构成漂移</p>",
    }
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    styles = {p.text: p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")}
    assert styles["第一章 技术方案"] == "Heading 1"
    assert styles["1.1 需求理解"] == "Heading 2"      # 节
    assert styles["1.1.1 现状分析"] == "Heading 3"    # 小节（此前与节平级）
    assert styles["1.2 总体设计"] == "Heading 2"
    assert styles["2.1 计划"] == "Heading 2"          # 旧式全 h3 章:节级,平级问题修复
    assert styles["3.0.1 只有小节的章"] == "Heading 3"  # 绝对映射:h4 不因缺 h3 被抬级
    assert doc.styles["Heading 4"].font.size.pt == 12  # 第四级样式已配置
    assert 'TOC \\o "1-4"' in doc.element.xml


def test_render_docx_maps_five_heading_levels():
    """五级提纲（第一章 → 一、 → 1. → （1） → ①）逐级落到 Word 1-5 级。
    写手契约：二级节 h3 / 三级小节 h4 / 四级细分 h5 / 五级明细 h6；章标题由渲染层占 1 级。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "整体服务方案", "group": "tech"}]}
    html = ("<h3>一、项目理解</h3><p>x</p>"
            "<h4>1. 项目背景分析</h4><p>y</p>"
            "<h5>（1）人员配置</h5><p>z</p>"
            "<h6>① 值班安排</h6><p>w</p>")
    doc = Document(io.BytesIO(render_docx(outline, {"t1": html})))
    levels = {p.text: p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")}
    assert levels["第一章 整体服务方案"] == "Heading 1"
    assert levels["一、项目理解"] == "Heading 2"
    assert levels["1. 项目背景分析"] == "Heading 3"
    assert levels["（1）人员配置"] == "Heading 4"
    assert levels["① 值班安排"] == "Heading 5"


def test_render_docx_strips_leaked_outline_ids_and_breaks_pages():
    """生产实测（7月28日反馈）：导出标题成了「t3.1 升级改造部署实施方案」——提纲内部 id 被写手
    抄进了标题，目录里也全是 t2.3/t3.1。渲染层确定性摘掉打头的 t/b，编号数字保留；
    另：每章另起一页（首章不加，否则目录后多一页空白）。"""
    outline = {"chapters": [
        {"id": "t1", "no": "第一章", "title": "整体服务方案", "group": "tech"},
        {"id": "t2", "no": "第二章", "title": "实施方案", "group": "tech"},
    ]}
    html = {
        "t1": "<h3>t1.1 项目理解</h3><p>x</p><h4>t1.1.1 背景分析</h4><p>y</p>",
        "t2": "<h3>b2b 服务体系</h3><p>z</p>",  # 「b2b」不是 id，一个字母都不能动
    }
    doc = Document(io.BytesIO(render_docx(outline, html)))
    heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "1.1 项目理解" in heads and "1.1.1 背景分析" in heads
    assert not any(h.startswith("t1.") for h in heads)
    assert "b2b 服务体系" in heads
    # 分页：两章之间恰好一个分页符（首章前不加）
    breaks = sum(1 for p in doc.paragraphs for r in p.runs if "w:br" in r._element.xml and 'type="page"' in r._element.xml)
    assert breaks >= 1


def test_render_docx_survives_garbage_span():
    """2026-08-06 生产事故复现：模型写出 rowspan="wer"，导出整步以 ValueError 失败，
    用户连点九次每次 0.2 秒就崩。非数字跨度按不合并处理，整本标书照样导得出来。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    chapters = {"t1": '<table><tr>'
                      '<td colspan="1" rowspan="wer"><p>附：</p></td>'
                      '<td colspan="abc" rowspan="1"><p>说明</p></td>'
                      '</tr><tr><td>甲</td><td>乙</td></tr></table>'}
    data = render_docx(outline, chapters)          # 不抛 = 事故不再复现
    doc = Document(io.BytesIO(data))
    assert doc.tables[0].rows[0].cells[0].text == "附："
    assert doc.tables[0].rows[0].cells[1].text == "说明"


def test_render_docx_caps_absurd_rowspan():
    """rowspan 远超表格行数时夹到剩余行数，不撑出上千行的空表。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "T", "group": "tech"}]}
    chapters = {"t1": '<table><tr><td rowspan="999">合并到底</td><td>x</td></tr>'
                      '<tr><td>y</td></tr></table>'}
    data = render_docx(outline, chapters)
    doc = Document(io.BytesIO(data))
    assert len(doc.tables[0].rows) == 2            # 只有两行 tr，就该是两行


def test_scope_tech_adds_volume_suffix_and_drops_group_tags():
    """分册(spec 2026-08-08-export-scope):技术册封面/页脚带「·技术标部分」,
    章标题不再带「（技术标）」尾巴——整册同组,逐章带尾巴是噪音。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "技术方案", "group": "tech"}]}
    data = render_docx(outline, {"t1": "<p>正文</p>"},
                       meta={"name": "XX项目投标文件"}, scope="tech")
    text = _doc_text(data)
    assert "XX项目投标文件·技术标部分" in text
    assert "（技术标）" not in text
    assert "投标人承诺与签章" in text          # 签章页每册都要(已拍板:独立提交物)


def test_scope_business_suffix_and_signature_kept():
    outline = {"chapters": [{"id": "b1", "no": "第一章", "title": "报价说明", "group": "business"}]}
    data = render_docx(outline, {"b1": "<p>正文</p>"},
                       meta={"name": "XX项目投标文件"}, scope="business")
    text = _doc_text(data)
    assert "·商务标部分" in text and "（商务标）" not in text
    assert "投标人承诺与签章" in text


def test_scope_full_output_is_byte_identical_to_today():
    """缺省/显式 full 与改动前逐字节一致——老调用方零感知(Global Constraints)。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "技术方案", "group": "tech"}]}
    chapters = {"t1": "<p>正文</p>"}
    meta = {"name": "XX项目投标文件"}
    assert render_docx(outline, chapters, meta=meta) == render_docx(outline, chapters, meta=meta, scope="full")
    text = _doc_text(render_docx(outline, chapters, meta=meta, scope="full"))
    assert "（技术标）" not in text and "·技术标部分" not in text   # 2026-08-15 起 full 也不带组尾巴


def test_render_strips_legacy_disclaimers_from_stored_chapters():
    """T6 回放实证（联通 08-12 存量项目）：旧提示词年代生成的章节里带着「本表格式与招标
    文件模板可能存在差异」，只在生成/改写时清，存量项目一导出就原样漏出——渲染层必须兜底。"""
    import io
    from docx import Document
    from agent.agents.bidding_agent.render.docx import render_docx

    outline = {"chapters": [{"id": "b1", "no": "第一章", "title": "报价函", "group": "business"}]}
    chapters = {"b1": ("<p><strong>提示：本表格式与招标文件模板可能存在差异，"
                       "请对照招标原文核对后使用。</strong></p><p>正文实质内容留下。</p>")}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "可能存在差异" not in text
    assert "正文实质内容留下" in text


def _page_break_precedes(doc, heading_text: str) -> bool:
    """标题段的前一段里是否有分页符（材料章分页断言专用）。"""
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text == heading_text and p.style.name.startswith("Heading"):
            return i > 0 and 'type="page"' in paras[i - 1]._p.xml
    raise AssertionError(f"标题未找到: {heading_text}")


def test_material_chapter_sections_each_start_new_page():
    """2026-08-15 用户实测：资格文件章里财务状况/信用中国/声明小节全挤在营业执照图后
    同一页，贴扫描件没版面。材料章（含证照占位图或「待补充」行的顶级小节 ≥2 个）
    每个顶级小节各起一页；首节紧跟章标题不加。"""
    outline = {"chapters": [{"id": "b6", "no": "第六章", "title": "要求的资格文件", "group": "business"}]}
    chapters = {"b6": (
        '<h3>一、营业执照及资质证书</h3>'
        '<p><img data-file-id="f1" data-object-key="k1" alt="营业执照|营业执照" /></p>'
        '<h3>二、财务状况证明材料</h3><p>（待补充：财务状况证明材料）</p>'
        '<h3>三、信用中国截图及无重大违法记录声明</h3><p>（待补充：信用中国截图及无重大违法记录声明）</p>'
        '<h3>四、独立承担民事责任能力及商业信誉声明</h3><p>供应商须出具声明，承诺：</p>'
    )}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    assert not _page_break_precedes(doc, "一、营业执照及资质证书")
    assert _page_break_precedes(doc, "二、财务状况证明材料")
    assert _page_break_precedes(doc, "三、信用中国截图及无重大违法记录声明")
    assert _page_break_precedes(doc, "四、独立承担民事责任能力及商业信誉声明")


def test_plain_chapter_sections_keep_flowing():
    """普通正文章不受材料分页影响：编辑器手贴的 data: 图不算证照材料（无 data-object-key）；
    只有一个材料小节的章也不整章分页。"""
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "技术方案", "group": "tech"}]}
    data_img = f'<p><img src="data:image/png;base64,{__import__("base64").b64encode(_TINY_PNG).decode()}" /></p>'
    chapters = {"t1": (
        f'<h3>1.1 架构图</h3>{data_img}'
        f'<h3>1.2 部署图</h3>{data_img}'
        '<h3>1.3 资质说明</h3><p>（待补充：安全服务资质）</p>'
    )}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    # 封面/目录/签章页自带分页，只断言小节标题前不许有
    assert not _page_break_precedes(doc, "1.2 部署图")
    assert not _page_break_precedes(doc, "1.3 资质说明")


def test_chapter_headings_carry_no_group_tag():
    """2026-08-15 用户拍板：正文章标题不带（技术标）/（商务标）尾巴——招标原文的章名
    就没有这种尾巴，逐章带上是噪音；目录条目取自标题，自动跟随。"""
    outline = {"chapters": [
        {"id": "t1", "no": "第一章", "title": "技术方案", "group": "tech"},
        {"id": "b1", "no": "第二章", "title": "响应函", "group": "business"},
    ]}
    data = render_docx(outline, {"t1": "<p>正文</p>", "b1": "<p>表单</p>"})
    doc = Document(io.BytesIO(data))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "（技术标）" not in texts and "（商务标）" not in texts
    assert "第一章 技术方案" in texts and "第二章 响应函" in texts


def _styles_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        return z.read("word/styles.xml").decode("utf-8")


def test_heading_styles_carry_no_theme_font_attrs():
    """2026-08-15 用户实测（(9).docx）：导出章节标题字体显示 ＭＳ ゴシック。python-docx
    默认 Heading 样式的 rFonts 带主题属性（asciiTheme/eastAsiaTheme），OOXML 规则主题
    属性**优先于**显式 ascii/eastAsia——查看器顺着 majorEastAsia 的日文脚本映射把标题
    解析成 MS Gothic，显式设的黑体形同虚设。设字体必须同时摘掉主题属性。"""
    import re as _re
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "方案", "group": "tech"}]}
    body = {"t1": "<h3>1.1 节</h3><p>正文</p><h4>1.1.1 小节</h4><h5>a</h5><h6>b</h6>"}
    styles = _styles_xml(render_docx(outline, body))
    for sid in ("Heading1", "Heading2", "Heading3", "Heading4", "Heading5"):
        frag = _re.search(rf'<w:style [^>]*w:styleId="{sid}".*?</w:style>', styles, _re.S).group(0)
        fonts = _re.search(r"<w:rFonts[^/]*/>", frag).group(0)
        assert "Theme" not in fonts and "theme" not in fonts, f"{sid} 仍带主题字体属性: {fonts}"
        assert 'w:eastAsia="黑体"' in fonts, f"{sid} 中文字体未显式落地: {fonts}"


def test_custom_format_heading_font_wins_over_theme():
    """spec330 自定义格式路径同病同治：用户配了标题宋体，主题属性不摘照样被 MS Gothic 顶掉。"""
    import re as _re
    outline = {"chapters": [{"id": "t1", "no": "第一章", "title": "方案", "group": "tech"}]}
    styles = _styles_xml(render_docx(outline, {"t1": "<h3>1.1 节</h3><p>正文</p>"},
                                     fmt={"heading_font": "宋体"}))
    frag = _re.search(r'<w:style [^>]*w:styleId="Heading1".*?</w:style>', styles, _re.S).group(0)
    fonts = _re.search(r"<w:rFonts[^/]*/>", frag).group(0)
    assert "Theme" not in fonts and "theme" not in fonts, fonts
    assert 'w:eastAsia="宋体"' in fonts, fonts


def test_fill_placeholders_do_not_fake_a_material_chapter():
    """评审 F1 CONFIRMED：填空占位「（待补充：____）」是写手/填空引擎给普通正文留的空，
    两个带填空的散文小节不得把整章误判成材料章强制一节一页；证照「待补充」（冒号后
    是证照名）照常算数。"""
    outline = {"chapters": [{"id": "b8", "no": "第八章", "title": "商务响应", "group": "business"}]}
    chapters = {"b8": (
        '<h3>一、报价说明</h3><p>付款方式：（待补充：____）</p>'
        '<h3>二、售后服务</h3><p>响应时限 2 小时。</p>'
        '<h3>三、联系方式</h3><p>联系人：（待补充：____）</p>'
    )}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    assert not _page_break_precedes(doc, "二、售后服务")
    assert not _page_break_precedes(doc, "三、联系方式")


def test_stray_h2_does_not_disable_material_breaks():
    """评审 F4：模型跑偏吐出的防御位 h2 与 h3 同落 Word 2 级——顶级按落级算，
    材料小节的分页不得因一个杂散 h2 整章失效。"""
    outline = {"chapters": [{"id": "b6", "no": "第六章", "title": "资格文件", "group": "business"}]}
    chapters = {"b6": (
        '<h2>资格文件总述</h2>'
        '<h3>一、营业执照</h3><p><img data-file-id="f1" data-object-key="k1" alt="营业执照|证" /></p>'
        '<h3>二、财务状况证明材料</h3><p>（待补充：财务状况证明材料）</p>'
        '<h3>三、信用中国截图</h3><p>（待补充：信用中国截图）</p>'
    )}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    assert _page_break_precedes(doc, "二、财务状况证明材料")
    assert _page_break_precedes(doc, "三、信用中国截图")


def test_copier_chapter_trailing_blank_paras_are_trimmed():
    """2026-08-15 (10).docx 实测第 11 页整页空白：报价一览表模板尾部带 14 个空段——
    招标原文用这串空行把下一份表单顶到新页，我们的章与章之间另有分页符，空行被顶成
    一整页空白。章间分页符落下前裁掉**章尾**连续空段；表单内部空行不在章尾，不动。"""
    from docx.oxml import parse_xml
    ns = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

    def p(text=""):
        run = f"<w:r><w:t>{text}</w:t></w:r>" if text else ""
        return parse_xml(f"<w:p {ns}>{run}</w:p>")

    nodes = [p("合计（大写）："), p(), p("表内空行后还有内容")] + [p() for _ in range(14)]
    outline = {"chapters": [
        {"id": "b1", "no": "第一章", "title": "报价一览表", "group": "business"},
        {"id": "b2", "no": "第二章", "title": "服务承诺", "group": "business"}]}
    doc = Document(io.BytesIO(render_docx(outline, {"b2": "<p>正文</p>"},
                                          copier_nodes={"b1": nodes})))
    paras = [(pp.text, 'type="page"' in pp._p.xml) for pp in doc.paragraphs]
    i = next(idx for idx, (t, _) in enumerate(paras) if t == "表内空行后还有内容")
    # 「第二章」先出现在目录缓存里——必须从 i 之后找正文里的那份,否则切片倒置空转
    j = next(idx for idx, (t, _) in enumerate(paras) if idx > i and t == "第二章 服务承诺")
    assert j > i
    empties = [t for t, br in paras[i + 1:j] if not t.strip() and not br]
    assert len(empties) == 0, f"章尾空段没裁干净: {paras[i + 1:j]}"
    k = next(idx for idx, (t, _) in enumerate(paras) if t == "合计（大写）：")
    assert paras[k + 1][0] == "" and paras[k + 2][0] == "表内空行后还有内容", "表单内部空行被误裁"


def test_nested_pending_material_sections_get_their_own_page():
    """2026-08-16 用户实测：资格文件章只有一个顶级 h3、材料小节在 h4 层（3.财务状况/
    5.信用中国 待补充）——顶级规则整章失效，待补充小节挤在一页。材料章判定改为
    **任意层级**材料段 ≥2；达标后嵌套材料段也各占一页，声明类嵌套段不动。"""
    outline = {"chapters": [{"id": "b6", "no": "第六章", "title": "资格文件", "group": "business"}]}
    chapters = {"b6": (
        '<h3>一、资格审查材料</h3><p>按招标要求提供以下材料。</p>'
        '<h4>1. 营业执照</h4><p><img data-file-id="f1" data-object-key="k1" alt="营业执照|证" /></p>'
        '<h4>3. 财务状况证明材料</h4><p>（待补充：财务状况证明材料）</p>'
        '<h4>4. 无重大违法记录声明</h4><p>我方郑重声明无重大违法记录。</p>'
        '<h4>5. 信用中国截图</h4><p>（待补充：信用中国截图）</p>'
    )}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    assert not _page_break_precedes(doc, "一、资格审查材料")   # 首标题紧跟章标题
    assert _page_break_precedes(doc, "1. 营业执照")
    assert _page_break_precedes(doc, "3. 财务状况证明材料")
    assert _page_break_precedes(doc, "5. 信用中国截图")
    assert not _page_break_precedes(doc, "4. 无重大违法记录声明"), "声明类嵌套段不该单开页"


def test_prose_chapter_with_two_cert_mentions_is_not_fully_paged():
    """评审：材料章判定放宽到任意层级后，散文章里两处证照「待补充」会把**整章顶级小节**
    全部分页，技术方案被拆散。顶级全分页只在**顶级材料段 ≥2** 时才开（原上线行为）；
    其余情况只给材料段本身单开页。"""
    outline = {"chapters": [{"id": "t2", "no": "第二章", "title": "整体服务方案", "group": "tech"}]}
    chapters = {"t2": (
        '<h3>一、项目理解</h3><p>需求分析……</p>'
        '<h4>1. 资质说明</h4><p>（待补充：安全服务资质）</p>'
        '<h3>二、技术方案</h3><p>架构设计……</p>'
        '<h4>2. 认证资质</h4><p>（待补充：信息安全认证）</p>'
        '<h3>三、实施计划</h3><p>分三阶段推进……</p>'
    )}
    doc = Document(io.BytesIO(render_docx(outline, chapters)))
    assert not _page_break_precedes(doc, "二、技术方案"), "散文章顶级小节被误分页"
    assert not _page_break_precedes(doc, "三、实施计划")
    assert _page_break_precedes(doc, "1. 资质说明"), "材料段本身仍单开页"
    assert _page_break_precedes(doc, "2. 认证资质")
