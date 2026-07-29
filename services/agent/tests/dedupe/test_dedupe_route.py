"""spec315b 契约 2：POST /dedupe 路由（monkeypatch read_bytes，内存 docx，hermetic）。"""
import io

from agent.routes import dedupe as dedupe_mod
from agent.routes.dedupe import DedupeBody, DedupeFile, dedupe

_A = ["本项目采用微服务架构设计，保障系统高可用与弹性扩展能力",
      "我方承诺在合同签订后三十个日历天内完成全部系统部署与上线工作",
      "项目经理具备十年以上同类系统集成项目管理经验"]
_B = ["本项目采用微服务架构设计，确保系统高可用与弹性扩展能力",
      "我方承诺在合同签订后三十个日历天内完成全部系统部署与上线任务",
      "售后服务团队提供七乘二十四小时热线响应服务"]


def _docx(paras: list[str], author: str | None = None) -> bytes:
    from docx import Document
    d = Document()
    if author:
        d.core_properties.author = author
        d.core_properties.last_modified_by = author
    for p in paras:
        d.add_paragraph(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()



def _fake_store(monkeypatch, files: dict[str, bytes]):
    """monkeypatch 路由模块的 read_bytes：按 key 从内存字典取，不碰 MinIO。"""
    monkeypatch.setattr(dedupe_mod, "read_bytes", lambda key: files[key])


async def test_dedupe_two_similar_files(monkeypatch):
    """大部分相同 + 同作者的两份 docx → 高分、text/meta 命中齐全、契约形状完整。"""
    _fake_store(monkeypatch, {"u/a.docx": _docx(_A, author="张三"),
                              "u/b.docx": _docx(_B, author="张三")})
    res = await dedupe(DedupeBody(files=[DedupeFile(key="u/a.docx", label="A 公司标书"),
                                         DedupeFile(key="u/b.docx", label="B 公司标书")],
                                  dims=["text", "image", "meta"]))
    assert res["dims_run"] == ["text", "image", "meta"]
    assert len(res["pairs"]) == 1
    p = res["pairs"][0]
    assert p["a"] == "A 公司标书" and p["b"] == "B 公司标书"
    assert p["score"] >= 40 and p["tone"] in ("warning", "destructive")
    dims_hit = {h["dim"] for h in p["hits"]}
    assert "text" in dims_hit and "meta" in dims_hit          # 相似句 + 同作者都命中
    text_hit = next(h for h in p["hits"] if h["dim"] == "text")
    assert text_hit["a_text"] and text_hit["b_text"] and "相似度" in text_hit["detail"]
    assert res["overall"]["max_score"] == p["score"]
    assert "；" in p["note"]                                   # 可解释中文 note


async def test_dedupe_baseline_strips_tender_sentences(monkeypatch):
    """基线开且给 tender_key：与招标文件相同的句子先剔除，得分显著下降。"""
    files = {"u/a.docx": _docx(_A), "u/b.docx": _docx(_B),
             "u/tender.docx": _docx(_A[:2])}                  # 招标文件含前两句（即相似来源）
    _fake_store(monkeypatch, files)
    body = lambda dims, tk=None: DedupeBody(  # noqa: E731
        files=[DedupeFile(key="u/a.docx", label="A"), DedupeFile(key="u/b.docx", label="B")],
        dims=dims, tender_key=tk)
    plain = await dedupe(body(["text"]))
    with_bl = await dedupe(body(["text", "baseline"], tk="u/tender.docx"))
    assert with_bl["dims_run"] == ["text", "baseline"]
    assert with_bl["pairs"][0]["score"] < plain["pairs"][0]["score"]
    assert "剔除" in with_bl["pairs"][0]["note"]


async def test_dedupe_three_files_three_pairs(monkeypatch):
    _fake_store(monkeypatch, {"u/a.docx": _docx(_A), "u/b.docx": _docx(_B),
                              "u/c.docx": _docx(["完全无关的第三份文件内容用于凑数比较"])})
    res = await dedupe(DedupeBody(files=[DedupeFile(key="u/a.docx", label="A"),
                                         DedupeFile(key="u/b.docx", label="B"),
                                         DedupeFile(key="u/c.docx", label="C")],
                                  dims=["text"]))
    assert len(res["pairs"]) == 3                              # 3 份 → C(3,2)=3 对
    assert {(p["a"], p["b"]) for p in res["pairs"]} == {("A", "B"), ("A", "C"), ("B", "C")}


async def test_dedupe_strategy_tone_thresholds_shift(monkeypatch):
    """同样的内容，tone 阈值随 strategy 平移：得分 ~72 在 fast（destructive≥80）只算
    warning，在 strict（destructive≥60）升级为 destructive。"""
    shared = "本项目严格按照招标文件要求组建实施团队并制定详细的项目进度计划与质量保障体系"
    _fake_store(monkeypatch, {"u/a.docx": _docx([shared, "我方拥有完备的售后服务网络体系"]),
                              "u/b.docx": _docx([shared, "本地化服务团队可两小时内到场"])})
    body = lambda s: DedupeBody(  # noqa: E731
        files=[DedupeFile(key="u/a.docx", label="A"), DedupeFile(key="u/b.docx", label="B")],
        dims=["text"], strategy=s)
    fast = (await dedupe(body("fast")))["pairs"][0]
    strict = (await dedupe(body("strict")))["pairs"][0]
    assert fast["tone"] == "warning" and strict["tone"] == "destructive"


async def test_dedupe_strategy_k_and_threshold_effective(monkeypatch):
    """同义改写句对：fast（k=8, th=0.50）不命中得 0 分；strict（k=3, th=0.45）命中高分——
    验证 k 与句对门槛按档生效。"""
    p1 = "本项目采用微服务架构设计保障系统高可用与弹性扩展能力满足业务持续增长需求"
    p2 = "本项目采用微服务架构设计确保系统高可用与弹性扩展能力满足业务快速增长需求"
    _fake_store(monkeypatch, {"u/a.docx": _docx([p1]), "u/b.docx": _docx([p2])})
    body = lambda s: DedupeBody(  # noqa: E731
        files=[DedupeFile(key="u/a.docx", label="A"), DedupeFile(key="u/b.docx", label="B")],
        dims=["text"], strategy=s)
    fast = (await dedupe(body("fast")))["pairs"][0]
    strict = (await dedupe(body("strict")))["pairs"][0]
    assert fast["score"] == 0 and not fast["hits"]
    assert strict["score"] > fast["score"] and strict["hits"][0]["dim"] == "text"


async def test_dedupe_image_dim_hits(monkeypatch):
    """image 维打分链：两份 docx 内嵌同一张图片 → 图片指纹命中并计入得分/说明。"""
    from PIL import Image
    from docx.shared import Inches

    img = Image.new("L", (64, 64))
    img.putdata([x * 4 for _ in range(64) for x in range(64)])   # 确定性渐变图
    png = io.BytesIO()
    img.save(png, format="PNG")

    def _docx_with_pic(paras: list[str]) -> bytes:
        from docx import Document
        d = Document()
        for p in paras:
            d.add_paragraph(p)
        d.add_picture(io.BytesIO(png.getvalue()), width=Inches(2))
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    _fake_store(monkeypatch, {"u/a.docx": _docx_with_pic(["甲方文件的完全不同正文内容"]),
                              "u/b.docx": _docx_with_pic(["乙方文件的另一套无关正文表述"])})
    res = await dedupe(DedupeBody(files=[DedupeFile(key="u/a.docx", label="A"),
                                         DedupeFile(key="u/b.docx", label="B")],
                                  dims=["image"]))
    p = res["pairs"][0]
    image_hits = [h for h in p["hits"] if h["dim"] == "image"]
    assert len(image_hits) == 1 and "汉明距离" in image_hits[0]["detail"]
    assert p["score"] == 6                                       # 1 处命中 × IMAGE_BONUS
    assert "图片指纹命中 1 处" in p["note"]


async def test_dedupe_file_count_400():
    res = await dedupe(DedupeBody(files=[DedupeFile(key="u/a.docx", label="A")]))
    assert res.status_code == 400


async def test_dedupe_parse_failure_422(monkeypatch):
    """取不到/解析不了某文件 → 422 且指明是哪份。"""
    _fake_store(monkeypatch, {"u/a.docx": _docx(_A)})          # b 缺失 → read_bytes KeyError
    res = await dedupe(DedupeBody(files=[DedupeFile(key="u/a.docx", label="A"),
                                         DedupeFile(key="u/b.docx", label="B 公司标书")]))
    assert res.status_code == 422
    body = res.body.decode()
    assert "B 公司标书" in body
    assert "请确认为文本可读" in body       # 只给稳定文案
    assert "KeyError" not in body           # 原始异常细节不外泄


async def test_dedupe_matches_same_content_across_docx_and_pdf(monkeypatch):
    """用户反馈「两份实际上传的是同一份文件内容，查重显示 0」：同一内容一份 docx 一份 PDF，
    PDF 经 pypdf 逐视觉行提取（每行一条 clause），旧实现把换行当句号切 → 两侧句子对不上判 0%。
    这里按 kind 造出两种形态的解析结果，验证 PDF 侧粘回硬换行后判为高度相同。"""
    from agent.parsing.parsers import ParsedDoc

    # 用真实标书的句读形态（句末有句号）：PDF 只有靠句号才知道段落到哪结束
    paras = ["本项目采用微服务架构设计，保障系统高可用与弹性扩展能力。",
             "我方承诺在合同签订后三十个日历天内完成全部系统部署与上线工作。",
             "售后服务团队提供七乘二十四小时热线响应服务并按时限闭环。"]

    def fake_parse(data: bytes, key: str) -> ParsedDoc:
        if key.endswith(".pdf"):   # PDF：每 20 字一条 clause，模拟逐视觉行
            lines = [para[i:i + 20] for para in paras for i in range(0, len(para), 20)]
            return ParsedDoc(text="\n".join(lines), kind="pdf",
                             clauses=[{"id": f"sec-1-c{i}", "text": t} for i, t in enumerate(lines)])
        return ParsedDoc(text="\n".join(paras), kind="docx",
                         clauses=[{"id": f"sec-1-c{i}", "text": t} for i, t in enumerate(paras)])

    _fake_store(monkeypatch, {"u/a.docx": b"x", "u/b.pdf": b"y"})
    monkeypatch.setattr(dedupe_mod, "parse_bytes", fake_parse)
    res = await dedupe(DedupeBody(files=[DedupeFile(key="u/a.docx", label="应答文件.docx"),
                                         DedupeFile(key="u/b.pdf", label="投标书.pdf")],
                                  dims=["text"]))
    p = res["pairs"][0]
    assert p["score"] >= 70 and p["tone"] == "destructive", p["note"]


async def test_dedupe_flags_documents_with_no_extractable_text(monkeypatch):
    """扫描件/图片版 PDF 提不出文字时，绝不能报「未见明显围标特征」——那是假放行，
    用户会据此认定两份标书没问题。必须显式说明本次文本比对不成立。"""
    _fake_store(monkeypatch, {"u/a.docx": _docx(_A), "u/b.docx": _docx([])})
    res = await dedupe(DedupeBody(files=[DedupeFile(key="u/a.docx", label="A.docx"),
                                         DedupeFile(key="u/b.docx", label="扫描件.docx")],
                                  dims=["text"]))
    note = res["pairs"][0]["note"]
    assert "未见明显围标特征" not in note
    assert "扫描件.docx" in note and "未提取到可比对文本" in note
