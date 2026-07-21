"""spec315b 契约 3：POST /render/checklist（mock storage，断言返回 key 且 upload 被调）。"""
import io
import uuid

from agent.routes import checklist as checklist_mod
from agent.routes.checklist import (ChecklistBody, ChecklistGroup, ChecklistItem,
                                    render_checklist)

_BODY = ChecklistBody(
    title="投标文件终极审核表",
    project_name="智慧园区综合管理平台项目",
    groups=[ChecklistGroup(id="g1", title="资质文件", items=[
        ChecklistItem(text="营业执照副本已盖章", status="pass", owner="张三"),
        ChecklistItem(text="授权委托书原件", status="risk", owner="李四",
                      note="缺法人签字", library_hit="缺失 · 授权委托书"),
        ChecklistItem(text="业绩合同复印件", status="pending"),
    ])])


class _Storage:
    def __init__(self):
        self.calls: list[tuple[str, bytes, str]] = []

    async def put_bytes(self, key, data, content_type=None):
        self.calls.append((key, data, content_type))


async def test_render_checklist_uploads_and_returns_key(monkeypatch):
    store = _Storage()
    monkeypatch.setattr(checklist_mod, "storage", store)
    res = await render_checklist(_BODY)
    assert len(store.calls) == 1                               # upload 被调且只调一次
    key, data, ct = store.calls[0]
    assert res == {"key": key}
    assert key.startswith("artifacts/checklist/") and key.endswith(".docx")
    uuid.UUID(key.removeprefix("artifacts/checklist/").removesuffix(".docx"))  # 中段是合法 uuid
    assert "wordprocessingml" in ct and len(data) > 0


async def test_render_checklist_docx_content(monkeypatch):
    """回读 docx：状态中文映射（✓/⚠/待办）、知识库列（字符串/空→—）、签字栏都在。"""
    store = _Storage()
    monkeypatch.setattr(checklist_mod, "storage", store)
    await render_checklist(_BODY)
    from docx import Document
    doc = Document(io.BytesIO(store.calls[0][1]))
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "检查项" in cells and "营业执照副本已盖章" in cells
    assert "✓ 通过" in cells and "⚠ 风险" in cells and "待办" in cells
    assert "知识库" in cells                                    # 知识库独立一列
    assert "缺失 · 授权委托书" in cells                          # 命中字符串原样输出
    assert "—" in cells and "缺法人签字" in cells                # 未命中显示 —；备注照常
    paras = [p.text for p in doc.paragraphs]
    assert any("智慧园区" in p for p in paras)
    assert any("签字" in p for p in paras)                     # 末尾签字/日期栏


_REPORT = checklist_mod.RiskReportBody(
    project_name="智慧园区项目·包件一",
    score=82, high=1, mid=2, passed=9,
    items=[checklist_mod.ReportItem(level="高", title="缺少★ISO27001 认证", chapter="资质文件", advice="补充认证复印件并加盖公章")],
    passed_items=["投标函格式符合要求"],
)


async def test_render_risk_report_docx(monkeypatch):
    """体检报告 docx：上传一次、key 形如 artifacts/report/<uuid>.docx、format 如实=docx；
    回读含评分总览/风险表格/已通过项/AI 声明。"""
    store = _Storage()
    monkeypatch.setattr(checklist_mod, "storage", store)
    res = await checklist_mod.render_risk_report(_REPORT)
    key, data, ct = store.calls[0]
    assert res == {"key": key, "format": "docx"}
    assert key.startswith("artifacts/report/") and key.endswith(".docx")
    assert "wordprocessingml" in ct
    from docx import Document
    doc = Document(io.BytesIO(data))
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "风险等级" in cells and "缺少★ISO27001 认证" in cells and "补充认证复印件并加盖公章" in cells
    paras = [p.text for p in doc.paragraphs]
    assert any("82 分" in p for p in paras)                    # 评分总览
    assert any("投标函格式符合要求" in p for p in paras)         # 已通过项
    assert any("AI 辅助生成" in p for p in paras)               # 声明


async def test_render_risk_report_pdf_falls_back_to_docx(monkeypatch):
    """format=pdf 但转换失败（docx_to_pdf 返回 None）→ 回落 docx，format 如实返回。"""
    store = _Storage()
    monkeypatch.setattr(checklist_mod, "storage", store)
    monkeypatch.setattr(checklist_mod, "docx_to_pdf", lambda _b: None)
    res = await checklist_mod.render_risk_report(_REPORT.model_copy(update={"format": "pdf"}))
    assert res["format"] == "docx" and res["key"].endswith(".docx")


async def test_render_risk_report_pdf_success(monkeypatch):
    """format=pdf 转换成功 → 存 pdf 字节、content-type application/pdf、format=pdf。"""
    store = _Storage()
    monkeypatch.setattr(checklist_mod, "storage", store)
    monkeypatch.setattr(checklist_mod, "docx_to_pdf", lambda _b: b"%PDF-fake")
    res = await checklist_mod.render_risk_report(_REPORT.model_copy(update={"format": "pdf"}))
    key, data, ct = store.calls[0]
    assert res == {"key": key, "format": "pdf"}
    assert key.endswith(".pdf") and data == b"%PDF-fake" and ct == "application/pdf"
