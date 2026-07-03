"""spec207 e2e 种子（一次性）：造一份小招标 docx 上传 MinIO，打印 key。"""
import asyncio
import io

from docx import Document

from agent.parsing.storage_read import storage

KEY = "uploads/e2e207/tender.docx"


def build() -> bytes:
    d = Document()
    d.add_heading("某市政务云平台运维服务项目 招标文件", level=0)
    d.add_heading("第一章 招标公告", level=1)
    d.add_paragraph("项目名称：某市政务云平台运维服务项目；招标编号：ZWY-2026-001；采购人：某市大数据管理局。预算金额：人民币 300 万元。服务期：1 年。")
    d.add_heading("第二章 投标人资格要求", level=1)
    d.add_paragraph("投标人须具备有效的 ISO27001 信息安全管理体系认证（★不可偏离，缺失即废标）。")
    d.add_paragraph("投标人须具备 ISO9001 质量管理体系认证；近三年至少 2 个同类政务云运维业绩。")
    d.add_heading("第三章 评分办法", level=1)
    d.add_paragraph("技术方案 50 分：运维服务体系 20 分、应急预案 15 分、人员配置 15 分。")
    d.add_paragraph("商务条款 30 分：同类业绩 15 分、资质证书 15 分。投标报价 20 分：低价优先。")
    d.add_heading("第四章 技术需求", level=1)
    d.add_paragraph("平台整体可用性不低于 99.9%（★）；7×24 小时值守；重大故障 30 分钟内响应、2 小时内到场。")
    d.add_paragraph("须满足网络安全等级保护 2.0 三级要求；提供分级 SLA 承诺与月度服务报告。")
    d.add_heading("第五章 商务要求", level=1)
    d.add_paragraph("投标保证金 5 万元；报价不得超过预算金额，超出即废标；付款方式：按季度支付。")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


async def main():
    await storage.put_bytes(KEY, build(),
                            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    print("seeded", KEY)


if __name__ == "__main__":
    asyncio.run(main())
